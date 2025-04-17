import mimetypes
import traceback
import uuid
import comtypes.client
from markupsafe import Markup
import pythoncom
from unidecode import unidecode
import bcrypt, json, os, sys, re, tempfile, shutil, psutil, time, logging
from PIL import Image
from datetime import datetime, date,  timedelta
from flask import Flask, abort, g, render_template, request, redirect, url_for, flash, jsonify, send_file, session as flask_session, make_response
from docx import Document
from docx.shared import Inches
from functools import wraps
from sqlalchemy.exc import IntegrityError, SQLAlchemyError
from sqlalchemy.orm.exc import StaleDataError
from itsdangerous import URLSafeTimedSerializer
from sqlalchemy import String, asc, cast, desc, extract, or_, func
from sqlalchemy.orm import joinedload
from sqlalchemy.orm import aliased
from flask_cors import CORS 
from werkzeug.utils import secure_filename
from flask import send_from_directory
from authlib.integrations.flask_client import OAuth
from mimetypes import MimeTypes
from Banco_Dados.Cria_Sessao import session
from Banco_Dados.models import Registro, Manutencao, LocalizacaoUsuario, Equipamento, Usuario, ChecklistItem, EntradaLog, StatusUsuario, FotoRegistro, AnexoManutencao, Setores, TiposEquipamentos, ChecklistTipo

# Adiciona o diretório principal ao sys.path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

app = Flask(__name__)
CORS(app)
app.secret_key = 'xdxdxd' 
serializer = URLSafeTimedSerializer(app.secret_key)
oauth = OAuth(app)

@app.errorhandler(404)
def not_found_error(error):
    return jsonify({'error': 'Endpoint não encontrado'}), 404

@app.errorhandler(500)
def internal_error(error):
    return jsonify({'error': 'Erro interno no servidor'}), 500

# Métodos de Login e Segurança ###################################################################

# Função para gerar token
def generate_token(user_id):
    return serializer.dumps(user_id, salt=app.secret_key)

# Função para verificar token
def verify_token(token, expiration=3600):
    try:
        user_id = serializer.loads(token, salt=app.secret_key, max_age=expiration)
    except Exception:
        return None
    return user_id

def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        token = flask_session.get('auth_token')
        if not token or not verify_token(token):
            # Registro de tentativa de acesso não autorizado
            data_atual = datetime.now()
            data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')
            
            mensagem_log = (
                f"(WEB) Tentativa de acesso não autorizado em {data_formatada}.\n"
                f"Token: {'Presente' if token else 'Ausente'}"
            )
            log_entry = EntradaLog(
                mensagem=mensagem_log,
                usuario="Desconhecido",
                nivel_acesso="Desconhecido",
                tipo_log="WEB - Segurança - Acesso Negado",
                data_geracao=data_formatada
            )
            session.add(log_entry)
            session.commit()
            
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def require_password_reset(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'auth_token' in flask_session:
            user_id = verify_token(flask_session['auth_token'])
            user = session.query(Usuario).filter_by(id=user_id).first()
            if user and user.precisa_redefinir_senha:
                return redirect(url_for('reset_senha'), code=302)
        return f(*args, **kwargs)
    return decorated_function

# Métodos de Configurações ###################################################################

# Configurações padrão com metadados para validação
DEFAULT_CONFIG = {
    'PERMISSIONS': {
        'value': {
            'OPERATOR_PERMISSIONS': [
                'inspecao',
                'inspecoes_cadastrados',
                'pendencias_inspecoes',
                'novo_equipamento'
                'equipamentos_cadastrados',
                'novo_setor',
                'setores_cadastrados',
                'relatorio_equipamentos',
                'relatorio_inspecoes',
                'notificacoes',
                'notificacoes_equipamentos'
            ],
            'READONLY_PERMISSIONS': [
                'inspecoes_cadastrados',
                'pendencias_inspecoes',
                'visualizar_dashboards'
            ]
        },
        'type': dict,
        'description': 'Configurações de permissões por nível de acesso'
    },
    'UPLOAD_FOLDER': {
        'value': os.path.join('static', 'images', 'Inspecoes_Realizadas'),
        'type': str,
        'description': 'Diretório para upload de arquivos de inspeções'
    },
    'FOTOS_EQP_FOLDER': {
        'value': os.path.join('static', 'uploads', 'equipamentos'),
        'type': str,
        'description': 'Diretório para upload de arquivos de inspeções'
    },
    'PERFIL_FOLDER': {
        'value': os.path.join('static', 'images', 'FotosPerfil'),
        'type': str,
        'description': 'Diretório para fotos de perfil de usuários'
    },
    'PDF_TEMPLATE_PATH': {
        'value': os.path.join('static', 'modelos', 'Checklist_Modelo.docx'),
        'type': str,
        'description': 'Caminho para o modelo de documento Word usado para gerar PDFs'
    },
    'PDF_OUTPUT_FOLDER': {
        'value': os.path.join('static', 'modelos', 'inspecoes_geradas'),
        'type': str,
        'description': 'Diretório onde os PDFs gerados serão salvos'
    },
    'PDF_FILENAME_PREFIX': {
        'value': 'Checklist_Inspecao_',
        'type': str,
        'description': 'Prefixo para os nomes dos arquivos PDF gerados'
    },    
    'MAX_CONTENT_LENGTH': {
        'value': 16,
        'type': int,
        'min': 1,
        'max': 100,
        'description': 'Tamanho máximo total de upload em MB'
    },
    'MAX_FILE_SIZE': {
        'value': 10,
        'type': int,
        'min': 1,
        'max': 100,
        'description': 'Tamanho máximo por arquivo em MB'
    },
    'PDF_CONVERSION_TIMEOUT': {
        'value': 120,
        'type': int,
        'min': 10,
        'max': 600,
        'description': 'Tempo limite para conversão de PDF em segundos'
    },
    'MAX_RETRIES': {
        'value': 2,
        'type': int,
        'min': 1,
        'max': 5,
        'description': 'Número máximo de tentativas para processamento'
    },
    'ALLOWED_EXTENSIONS': {
        'value': ['.png', '.jpg', '.jpeg', '.gif', '.mp4', '.avi', '.mov', '.mkv', '.wmv', '.pdf'],
        'type': list,
        'description': 'Extensões de arquivo permitidas'
    },
    'INSPECTION_PREFIX': {
        'value': 'INS-',
        'type': str,
        'description': 'Prefixo para números de inspeção'
    },    
    'SERVER_IP': {
        'value': '0.0.0.0',
        'type': str,
        'description': 'Endereço IP do servidor'
    },
    'SERVER_PORT': {
        'value': 5000,
        'type': int,
        'min': 1024,
        'max': 65535,
        'description': 'Porta do servidor'
    },
    'SERVER_PROTOCOL': {
        'value': 'http',
        'type': str,
        'options': ['http', 'https'],
        'description': 'Protocolo do servidor'
    },
    'AUTO_CREATE_FOLDERS': {
        'value': True,
        'type': bool,
        'description': 'Criar diretórios automaticamente se não existirem'
    },
    'CONFIG_VERSION': {
        'value': '1.0',
        'type': str,
        'description': 'Versão do esquema de configuração'
    },
    'OAUTH_PROVIDERS': {
        'value': {
            'azure': {
                'enabled': False,
                'client_id': '',
                'client_secret': '',
                'tenant_id': '',
                'scopes': 'openid email profile',
                'authorize_params': {'response_type': 'code'}
            }
        },
        'type': dict,
        'description': 'Configurações dos provedores OAuth'
    }
    
}

CONFIG_FILE = 'system_config.json'
BACKUP_FOLDER = 'config_backups'

def create_backup():
    """Criar backup da configuração atual"""
    try:
        if not os.path.exists(BACKUP_FOLDER):
            os.makedirs(BACKUP_FOLDER)
        
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        backup_file = os.path.join(BACKUP_FOLDER, f'config_backup_{timestamp}.json')
        
        with open(CONFIG_FILE, 'r') as original, open(backup_file, 'w') as backup:
            backup.write(original.read())
            
        logging.info(f'Backup criado: {backup_file}')
        return True
    except Exception as e:
        logging.error(f'Erro ao criar backup: {str(e)}')
        return False

def init_config():
    print("[DEBUG] Entrou em init_config")
    if not os.path.exists(CONFIG_FILE):
        print("[DEBUG] Arquivo de config não existe, salvando padrão")
        save_config(get_default_values())
        create_directories()

def get_default_values():
    """Retorna apenas os valores padrão sem os metadados"""
    result = {}
    for key, data in DEFAULT_CONFIG.items():
        if 'value' in data:
            result[key] = data['value']
        else:
            # Para chaves que não seguem o formato padrão (como OAUTH_PROVIDERS)
            result[key] = data
    return result

def create_directories():
    """Criar diretórios necessários"""
    config = load_config()
    if config.get('AUTO_CREATE_FOLDERS', True):
        try:
            os.makedirs(config['UPLOAD_FOLDER'], exist_ok=True)
            os.makedirs(config['PERFIL_FOLDER'], exist_ok=True)
            logging.info('Diretórios criados com sucesso')
        except Exception as e:
            logging.error(f'Erro ao criar diretórios: {str(e)}')

def load_config():
    """Carregar configurações do arquivo JSON com validação"""
    try:
        with open(CONFIG_FILE, 'r') as f:
            config = json.load(f)
            
            # Validação básica do esquema
            if not isinstance(config, dict):
                raise ValueError("Formato inválido do arquivo de configuração")
                
            # Garantir que todas as chaves padrão existam
            for key, data in DEFAULT_CONFIG.items():
                if key not in config:
                    config[key] = data['value']
                    logging.warning(f'Chave {key} faltando no arquivo de configuração. Usando valor padrão.')
            
            return config
    except (FileNotFoundError, json.JSONDecodeError) as e:
        logging.error(f'Erro ao carregar configurações: {str(e)}. Usando valores padrão.')
        return get_default_values()
    except Exception as e:
        logging.error(f'Erro inesperado ao carregar configurações: {str(e)}')
        return get_default_values()

def validate_config(config):
    """Validar configurações contra o esquema definido"""
    errors = []
    for key, value in config.items():
        if key in DEFAULT_CONFIG:
            expected_type = DEFAULT_CONFIG[key]['type']
            
            # Verificação de tipo
            if not isinstance(value, expected_type):
                errors.append(f'Tipo inválido para {key}. Esperado {expected_type}, obtido {type(value)}')
                continue
                
            # Verificação de valores para números
            if expected_type in (int, float):
                if 'min' in DEFAULT_CONFIG[key] and value < DEFAULT_CONFIG[key]['min']:
                    errors.append(f'Valor para {key} menor que o mínimo permitido ({DEFAULT_CONFIG[key]["min"]})')
                if 'max' in DEFAULT_CONFIG[key] and value > DEFAULT_CONFIG[key]['max']:
                    errors.append(f'Valor para {key} maior que o máximo permitido ({DEFAULT_CONFIG[key]["max"]})')
            
            # Verificação de opções para strings
            if expected_type == str and 'options' in DEFAULT_CONFIG[key]:
                if value not in DEFAULT_CONFIG[key]['options']:
                    errors.append(f'Valor inválido para {key}. Opções válidas: {", ".join(DEFAULT_CONFIG[key]["options"])}')
    
    return errors if errors else None

def save_config(config):
    """Salvar configurações no arquivo JSON com validação"""
    try:
        # Validar antes de salvar
        if errors := validate_config(config):
            raise ValueError(" | ".join(errors))
            
        # Criar backup antes de salvar
        create_backup()
        
        with open(CONFIG_FILE, 'w') as f:
            json.dump(config, f, indent=4)
        
        # Atualizar configurações do app Flask
        update_flask_config(config)
        
        logging.info('Configurações salvas com sucesso')
        return True
    except Exception as e:
        logging.error(f'Erro ao salvar configurações: {str(e)}')
        raise

def update_flask_config(config):
    """Atualizar configurações do Flask com base no arquivo"""

    def get_config_value(campo):
        valor = config.get(campo)
        if isinstance(valor, dict) and 'value' in valor:
            return valor['value']
        return valor

    # Configurações essenciais
    app.config['UPLOAD_FOLDER'] = os.path.abspath(get_config_value('UPLOAD_FOLDER'))
    app.config['FOTOS_EQP_FOLDER'] = os.path.abspath(get_config_value('FOTOS_EQP_FOLDER'))
    app.config['ALLOWED_EXTENSIONS'] = get_config_value('ALLOWED_EXTENSIONS')  # provavelmente uma lista
    app.config['PERFIL_FOLDER'] = os.path.abspath(get_config_value('PERFIL_FOLDER'))
    app.config['MAX_CONTENT_LENGTH'] = get_config_value('MAX_CONTENT_LENGTH') * 1024 * 1024

    # Configurações adicionais
    app.config['PDF_TEMPLATE_PATH'] = os.path.abspath(get_config_value('PDF_TEMPLATE_PATH'))
    app.config['PDF_OUTPUT_FOLDER'] = os.path.abspath(get_config_value('PDF_OUTPUT_FOLDER'))

    # Criar diretórios se necessário
    if get_config_value('AUTO_CREATE_FOLDERS') is not False:
        create_directories()

@app.route('/api/system/config', methods=['GET'])
@login_required
def get_system_config():
    try:
        usuario_logado, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

        config = load_config()
        response = {
            'status': 'success',
            'config': config,
            'metadata': {key: {
                'description': DEFAULT_CONFIG[key]['description'],
                'type': DEFAULT_CONFIG[key]['type'].__name__,
                **({k: v for k, v in DEFAULT_CONFIG[key].items() if k not in ['value', 'type', 'description']})
            } for key in DEFAULT_CONFIG if key in config}
        }

        # Registro de log
        mensagem_log = (
            f"(WEB) O Usuário {usuario_logado} - {nivel_acesso} acessou as configurações do sistema em {data_formatada}."
        )
        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario_logado,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Configurações - Acesso",
            data_geracao=data_formatada
        )
        session.add(log_entry)
        session.commit()

        return jsonify(response)
    except Exception as e:
        logging.error(f'Erro no endpoint get_system_config: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': f'Erro ao carregar configurações: {str(e)}'
        }), 500

@app.route('/api/system/config', methods=['POST'])
@login_required
def update_system_config():
    """Endpoint para atualizar configurações do sistema com validação"""
    try:
        new_config = request.get_json()
        if not new_config:
            raise ValueError("Nenhum dado de configuração fornecido")
            
        current_config = load_config()
        
        # Validar e mesclar configurações
        for key, value in new_config.items():
            if key in DEFAULT_CONFIG:
                current_config[key] = value
        
        # Validar e salvar
        save_config(current_config)
        
        return jsonify({
            'status': 'success',
            'message': 'Configurações atualizadas com sucesso',
            'config': current_config
        })
    except Exception as e:
        logging.error(f'Erro no endpoint update_system_config: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': f'Erro ao atualizar configurações: {str(e)}'
        }), 400

@app.route('/api/system/config/reset', methods=['POST'])
@login_required
def reset_system_config():
    """Endpoint para resetar configurações para os valores padrão"""
    try:
        save_config(get_default_values())
        return jsonify({
            'status': 'success',
            'message': 'Configurações resetadas para os valores padrão',
            'config': get_default_values()
        })
    except Exception as e:
        logging.error(f'Erro no endpoint reset_system_config: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': f'Erro ao resetar configurações: {str(e)}'
        }), 500

@app.route('/api/system/config/backups', methods=['GET'])
@login_required
def list_backups():
    """Endpoint para listar backups disponíveis"""
    try:
        backups = []
        if os.path.exists(BACKUP_FOLDER):
            for file in sorted(os.listdir(BACKUP_FOLDER)):
                if file.endswith('.json'):
                    file_path = os.path.join(BACKUP_FOLDER, file)
                    backups.append({
                        'name': file,
                        'path': file_path,
                        'size': os.path.getsize(file_path),
                        'modified': os.path.getmtime(file_path)
                    })
        
        return jsonify({
            'status': 'success',
            'backups': backups
        })
    except Exception as e:
        logging.error(f'Erro no endpoint list_backups: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': f'Erro ao listar backups: {str(e)}'
        }), 500

@app.route('/api/system/config/restore', methods=['POST'])
@login_required
def restore_backup():
    """Endpoint para restaurar a partir de um backup"""
    try:
        backup_file = request.json.get('backup_file')
        if not backup_file or not os.path.exists(backup_file):
            raise ValueError("Arquivo de backup inválido ou não encontrado")
            
        with open(backup_file, 'r') as f:
            backup_config = json.load(f)
            
        save_config(backup_config)
        
        return jsonify({
            'status': 'success',
            'message': 'Configurações restauradas com sucesso',
            'config': backup_config
        })
    except Exception as e:
        logging.error(f'Erro no endpoint restore_backup: {str(e)}')
        return jsonify({
            'status': 'error',
            'message': f'Erro ao restaurar backup: {str(e)}'
        }), 400

def get_allowed_menus_for_user(user):
    if user is None:
        return []  # ou alguma resposta padrão
    
    if user.nivel_acesso == 'Administrador':
        return {
            'inspecao',
            'inspecoes_cadastrados',
            'pendencias_inspecoes',
            'mapas_inspecoes',
            'excluir_inspecao',
            'imprimir_inspecao',
            'alterar_status_inspecao',
            'ver_anexos',
            'desativa_localizacao',

            'equipamentos_cadastrados',
            'novo_equipamento',
            'excluir_equipamento',
            'editar_equipamento',
            'alterar_status_equipamento',
            'gerar_codigo',

            'usuarios_cadastrados',
            'desbloquear_usuario',
            'bloquear_usuario',
            'excluir_usuario',
            'editar_usuario',
            'novo_usuario',

            'setores_cadastrados',
            'novo_setor',
            'editar_setor',
            'excluir_setor',

            'tipos_equipamentos_cadastrados',
            'novo_tipo_equipamento',
            'editar_tipo_equipamento',
            'gerenciar_tipos_equipamentos',

            'relatorio_equipamentos',
            'relatorio_inspecoes',

            'notificacoes',
            'notificacoes_equipamentos',

            'logs_auditoria',
            'visualizar_dashboards',
            'configuracoes',
            'alterar_senha'
        }

    config = load_config()
    permissions = config.get('PERMISSIONS', {})
    level_permissions = {
        'Operador': set(permissions.get('OPERATOR_PERMISSIONS', [])),
        'Somente Leitura': set(permissions.get('READONLY_PERMISSIONS', [])),
    }

    return level_permissions.get(user.nivel_acesso, set())

def permission_required(menu_id):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            current_user = get_current_user(flask_session, session)

            # Administrador tem acesso total
            if current_user.nivel_acesso == 'Administrador':
                return f(*args, **kwargs)

            # Carrega as permissões do sistema
            config = load_config()
            permissions = config.get('PERMISSIONS', {})

            # Define os menus permitidos por nível de acesso
            level_permissions = {
                'Operador': set(permissions.get('OPERATOR_PERMISSIONS', [])),
                'Somente Leitura': set(permissions.get('READONLY_PERMISSIONS', [])),
            }

            allowed_menus = level_permissions.get(current_user.nivel_acesso, set())

            if menu_id in allowed_menus:
                return f(*args, **kwargs)
            
            # Usuário não tem permissão
            return render_template('403.html'), 403

        return decorated_function
    return decorator

def init_oauth(app):
    """Inicializa o sistema OAuth com as configurações atuais"""
    config = load_config()
    oauth_providers = config.get('OAUTH_PROVIDERS', {})
    
    if not oauth_providers.get('azure', {}).get('enabled', False):
        return None
    
    oauth = OAuth(app)
    
    azure_config = oauth_providers['azure']
    oauth.register(
        name='azure',
        client_id=azure_config['client_id'],
        client_secret=azure_config['client_secret'],
        authorize_url=f"https://login.microsoftonline.com/{azure_config['tenant_id']}/oauth2/v2.0/authorize",
        authorize_params=azure_config['authorize_params'],
        access_token_url=f"https://login.microsoftonline.com/{azure_config['tenant_id']}/oauth2/v2.0/token",
        access_token_params=None,
        client_kwargs={'scope': azure_config['scopes']}
    )
    
    return oauth

@app.route('/auth/azure')
def auth_azure():
    if 'oauth' not in g:
        abort(403, description="Azure OAuth não está configurado")
    return g.oauth.azure.authorize_redirect(url_for('auth_azure_callback', _external=True))

@app.route('/auth/azure/callback')
def auth_azure_callback():
    if 'oauth' not in g:
        abort(403, description="Azure OAuth não está configurado")
    
    token = g.oauth.azure.authorize_access_token()
    user_info = g.oauth.azure.parse_id_token(token)
    session['user'] = user_info
    return redirect(url_for('index'))

@app.route('/admin/configuracoes')
@login_required
@permission_required('configuracoes')
def system_config_page():
    """Rota para renderizar a página de configurações do sistema"""
    current_user = get_current_user(flask_session, session) 
    return render_template('configuracoes.html', current_user=current_user)  

# Métodos de Login ###################################################################

@app.route('/', methods=['GET', 'POST'])
@app.route('/login', methods=['GET', 'POST'])
def login():
    user = None  # Inicializa a variável user como None

    # Verifica se há um token válido na sessão
    if 'auth_token' in flask_session and verify_token(flask_session['auth_token']):
        user_id = verify_token(flask_session['auth_token'])
        user = session.query(Usuario).filter_by(id=user_id).first()

        if user and user.precisa_redefinir_senha:
            return redirect(url_for('reset_senha'), code=302)

        return redirect(url_for('index', token=flask_session['auth_token']), code=302)

    # Se o método for POST, processa o formulário de login
    if request.method == 'POST':
        login_input = request.form.get('usuario', '').strip().lower()
        senha = request.form.get('senha', '').strip()

        if not login_input or not senha:
            return jsonify({"error": "Por favor, preencha todos os campos."}), 400

        # Buscar o usuário pelo nome de usuário ou e-mail
        user = session.query(Usuario).filter(
            or_(Usuario.usuario == login_input, Usuario.email == login_input)
        ).first()

        if user:
            if bcrypt.checkpw(senha.encode('utf-8'), user.senha):
                if user.status and any(status.status == "Acesso Bloqueado" for status in user.status):
                    return jsonify({"error": "Sua conta está bloqueada. Entre em contato com o administrador."}), 403

                if user.precisa_redefinir_senha:
                    flask_session['auth_token'] = generate_token(user.id)
                    return jsonify({"redirect": url_for('reset_senha')}), 200

                flask_session['auth_token'] = generate_token(user.id)
                token = flask_session['auth_token']
                user.tentativas_login = 0

                # Adicionar o registro de log
                data_atual = datetime.now()
                data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')
                nivel_acesso = user.nivel_acesso  # Ajuste para capturar o nível de acesso correto

                mensagem_log = (
                    f"(WEB) O Usuário {user.usuario} ({nivel_acesso}) acessou o InspekApp em {data_formatada}.\n"
                )
                log_entry = EntradaLog(
                    mensagem=mensagem_log,
                    usuario=user.usuario,
                    nivel_acesso=nivel_acesso,
                    tipo_log="WEB - Sistema - Login - Web",
                    data_geracao=data_formatada
                )
                session.add(log_entry)
                session.commit()

                # Retorna o nome e a foto de perfil do usuário
                foto_perfil = user.foto_perfil if user.foto_perfil else 'default_profile.png'
                return jsonify({
                    "redirect": url_for('index', token=token),
                    "nome": user.nome,
                    "foto_perfil": url_for('static', filename=foto_perfil)
                }), 200
            else:
                user.tentativas_login += 1
                if user.tentativas_login >= 3:
                    status_entrada = session.query(StatusUsuario).filter_by(usuario_id=user.id).first()
                    if not status_entrada:
                        status_entrada = StatusUsuario(status="Acesso Bloqueado", usuario_id=user.id)
                        session.add(status_entrada)
                    else:
                        status_entrada.status = "Acesso Bloqueado"
                    session.commit()
                    return jsonify({"error": "Conta bloqueada devido a várias tentativas de login falhas."}), 403
                else:
                    session.commit()
                    return jsonify({"error": "Senha incorreta. Tente novamente."}), 401
        else:
            return jsonify({"error": "Usuário ou e-mail não encontrado."}), 404

    return render_template('login/login.html')

@app.route('/logout', methods=['POST'])
@login_required
def logout():
    try:
        usuario, nivel_acesso, foto_perfil_url = get_current_user_name_and_access_level(flask_session, session)
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

        # Adicionar o registro de log
        mensagem_log = (
            f"(WEB) O Usuário {usuario} ({nivel_acesso}) saiu do InspekApp em {data_formatada}.\n"
        )
        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Sistema - Logout - Web",
            data_geracao=data_formatada
        )
        session.add(log_entry)
        session.commit()
    except SQLAlchemyError as e:
        session.rollback()
        print(f"Erro ao registrar log de logout: {e}")

    # Remover o token de autenticação
    flask_session.pop('auth_token', None)
    
    # Redirecionar para a página de login
    response = make_response(redirect(url_for('login'), code=302))
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, post-check=0, pre-check=0, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '-1'

    return response

@app.route('/novo_usuario_login')
def form_novo_usuario_login():
    return render_template('login/novo_usuario_login.html')

@app.route('/api/novo_usuario_cadastro', methods=['POST'])
def novo_usuario_login():
    try:
        # Obter dados do formulário
        primeiro_nome = request.form.get('primeiro_nome', '').strip()
        nome_meio = request.form.get('nome_meio', '').strip()
        sobrenome = request.form.get('sobrenome', '').strip()
        nome = request.form.get('nome', '').strip()
        email = request.form.get('email', '').strip()
        matricula = request.form.get('matricula', '').strip()
        setor = request.form.get('setor', '').strip()
        usuario_nome = request.form.get('usuario', '').strip()
        senha = request.form.get('senha', '').strip()
        senha_confirmacao = request.form.get('confirme_senha', '').strip()
        pergunta_seguranca = request.form.get('pergunta_seguranca', '').strip()
        resposta_seguranca = request.form.get('resposta_seguranca', '').strip()
        nivel_acesso = request.form.get('nivel_acesso', '').strip()

        # Validações básicas
        required_fields = {
            'primeiro_nome': primeiro_nome,
            'sobrenome': sobrenome,
            'nome': nome,
            'email': email,
            'matricula': matricula,
            'setor': setor,
            'usuario': usuario_nome,
            'senha': senha,
            'confirme_senha': senha_confirmacao,
            'pergunta_seguranca': pergunta_seguranca,
            'resposta_seguranca': resposta_seguranca,
            'nivel_acesso': nivel_acesso
        }

        missing_fields = [field for field, value in required_fields.items() if not value]
        if missing_fields:
            return jsonify({
                'error': 'Campos obrigatórios faltando',
                'missing_fields': missing_fields
            }), 400

        if senha != senha_confirmacao:
            return jsonify({'error': 'As senhas não correspondem.'}), 400

        # Validação da política de senhas
        if len(senha) < 8:
            return jsonify({'error': 'A senha deve ter pelo menos 8 caracteres.'}), 400

        if not re.search(r'[A-Z]', senha):
            return jsonify({'error': 'A senha deve conter pelo menos uma letra maiúscula.'}), 400

        if not re.search(r'[a-z]', senha):
            return jsonify({'error': 'A senha deve conter pelo menos uma letra minúscula.'}), 400

        if not re.search(r'[!@#$%^&*(),.?":{}|<>]', senha):
            return jsonify({'error': 'A senha deve conter pelo menos um símbolo.'}), 400

        if session.query(Usuario).filter_by(email=email).first() or session.query(Usuario).filter_by(usuario=usuario_nome).first():
            return jsonify({'error': 'Email ou usuário já estão em uso.'}), 400

        senha_criptografada = bcrypt.hashpw(senha.encode('utf-8'), bcrypt.gensalt())

        novo_usuario = Usuario(
            primeiro_nome=primeiro_nome,
            nome_meio=nome_meio,
            sobrenome=sobrenome,
            nome=nome,
            email=email, 
            matricula=matricula, 
            setor=setor,                 
            usuario=usuario_nome, 
            senha=senha_criptografada, 
            nivel_acesso=nivel_acesso,
            pergunta_seguranca=pergunta_seguranca, 
            resposta_seguranca=resposta_seguranca,
            usuario_geracao="Tela Login",
        )

        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

        mensagem_log = (
            f"(WEB) O Usuário atravez da Tela de Login cadastrou um novo usuário em {data_atual.strftime('%d-%m-%Y %H:%M:%S')}.\n"
            f"Nome: {nome}\n"
            f"Email: {email}\n"
            f"Matricula: {matricula}\n"  
            f"Setor: {setor}\n"             
            f"Usuário: {usuario_nome}\n"
            f"Nível de Acesso: {nivel_acesso}\n"
        )

        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario="Tela Login",
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Usuário - Novo Cadastro",
            data_geracao=data_formatada
        )

        session.add(log_entry)
        session.add(novo_usuario)
        session.commit()

        return jsonify({'success': 'Usuário cadastrado com sucesso.'})

    except Exception as e:
        # Log do erro no servidor
        print(f"Erro ao cadastrar usuário: {str(e)}")
        # Retorna uma resposta JSON mesmo em caso de erro
        return jsonify({
            'error': 'Erro interno no servidor',
            'details': str(e)
        }), 500
 
@app.route('/api/alterar_foto_perfil', methods=['GET', 'POST'])
@login_required
def alterar_foto_perfil():
    usuario = None
    nivel_acesso = None
    foto_perfil_url = None

    try:
        print("Iniciando processamento da foto de perfil...")

        # Garante que o PERFIL_FOLDER está carregado
        if 'PERFIL_FOLDER' not in app.config:
            init_config()
            update_flask_config(load_config())

        perfil_folder = os.path.abspath(app.config['PERFIL_FOLDER'])

        if request.method != 'POST':
            return jsonify({'error': 'Método não permitido. Use POST.'}), 405

        if 'foto_perfil' not in request.files or request.files['foto_perfil'].filename == '':
            return jsonify({'error': 'Arquivo inválido.'}), 400

        foto_perfil = request.files['foto_perfil']
        allowed_extensions = {'png', 'jpg', 'jpeg', 'gif', 'mp4', 'avi', 'mov', 'mkv', 'wmv', 'pdf'}
        file_extension = foto_perfil.filename.split('.')[-1].lower()

        if '.' not in foto_perfil.filename or file_extension not in allowed_extensions:
            return jsonify({'error': 'Tipo de arquivo não suportado.'}), 400

        # Obter usuário atual
        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        if not usuario:
            return jsonify({'error': 'Usuário não encontrado.'}), 404

        # Nome do arquivo
        filename = f"usuario_{usuario}.{file_extension}"
        foto_path = os.path.join(perfil_folder, filename)

        # Salvar arquivo
        os.makedirs(perfil_folder, exist_ok=True)
        foto_perfil.save(foto_path)

        # Criar URL relativa ao /static
        static_folder_abs = os.path.abspath(os.path.join(app.root_path, 'static'))
        foto_perfil_url = os.path.relpath(foto_path, static_folder_abs).replace("\\", "/")

        # Atualizar no banco de dados
        usuario_atual = session.query(Usuario).filter_by(usuario=usuario).first()
        if usuario_atual:
            usuario_atual.foto_perfil = foto_perfil_url
            session.commit()
        else:
            return jsonify({'error': 'Erro ao atualizar foto no banco de dados.'}), 500

        # Registrar log
        data_formatada = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
        log_entry = EntradaLog(
            mensagem=(
                f"(WEB) O Usuário {usuario} ({nivel_acesso}) alterou a foto de perfil em {data_formatada}.\n"
                f"Nova Foto: {foto_perfil_url}\n"
            ),
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Usuário - Alteração de Foto de Perfil",
            data_geracao=data_formatada
        )
        session.add(log_entry)
        session.commit()

        return jsonify({'success': 'Foto de perfil atualizada com sucesso.', 'foto_perfil_url': foto_perfil_url})

    except Exception as e:
        data_formatada = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
        mensagem_log = (
            f"(WEB) Erro ao alterar a foto de perfil do usuário {usuario or '[Desconhecido]'} ({nivel_acesso or '[Desconhecido]'}) em {data_formatada}.\n"
            f"Erro: {str(e)}\n"
        )

        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario or 'Desconhecido',
            nivel_acesso=nivel_acesso or 'Desconhecido',
            tipo_log="WEB - Usuário - Erro ao Alterar Foto de Perfil",
            data_geracao=data_formatada
        )
        session.add(log_entry)
        session.commit()

        return jsonify({'error': 'Erro interno ao processar a foto de perfil.'}), 500
    
# Métodos de Reset de Senha ###################################################################

@app.route('/senhas/reset_password', methods=['GET', 'POST'])
def reset_password():
    if request.method == 'POST':
        login_input = request.form.get('usuario', '').strip().lower()
        pergunta_seguranca = request.form.get('pergunta_seguranca', '').strip()
        resposta_seguranca = request.form.get('resposta_seguranca', '').strip()
        nova_senha = request.form.get('nova_senha', '').strip()
        confirme_senha = request.form.get('confirme_senha', '').strip()

        if not login_input or not nova_senha or not confirme_senha:
            flash("Todos os campos são obrigatórios.", "danger")
            return redirect(url_for('reset_password'), code=302)

        if nova_senha != confirme_senha:
            flash("As senhas não coincidem.", "danger")
            return redirect(url_for('reset_password'), code=302)
        
        # Validação da política de senhas
        if len(nova_senha) < 8:
            flash("A senha deve ter pelo menos 8 caracteres.", "danger")
            return redirect(url_for('reset_password'), code=302)

        if not re.search(r'[A-Z]', nova_senha):
            flash("A senha deve conter pelo menos uma letra maiúscula.", "danger")
            return redirect(url_for('reset_password'), code=302)
        
        if not re.search(r'[a-z]', nova_senha):
            flash("A senha deve conter pelo menos uma letra minúscula.", "danger")
            return redirect(url_for('reset_password'), code=302)
        
        if not re.search(r'[!@#$%^&*(),.?":{}|<>]', nova_senha):
            flash("A senha deve conter pelo menos um símbolo.", "danger")
            return redirect(url_for('reset_password'), code=302)
        
        # Busca o usuário pelo nome de usuário ou pelo e-mail
        user = session.query(Usuario).filter(
            or_(Usuario.usuario == login_input, Usuario.email == login_input)
        ).first()

        if not user:
            flash("Usuário ou e-mail não encontrado.", "danger")
            return redirect(url_for('reset_password'), code=302)

        # Verifica a pergunta de segurança se foi preenchida
        if pergunta_seguranca and resposta_seguranca:
            if user.pergunta_seguranca != pergunta_seguranca or user.resposta_seguranca.lower() != resposta_seguranca.lower():
                flash("A pergunta ou resposta de segurança está incorreta.", "danger")
                return redirect(url_for('reset_password'), code=302)

        # Gera um novo hash de senha
        hashed_password = bcrypt.hashpw(nova_senha.encode('utf-8'), bcrypt.gensalt())
        user.senha = hashed_password
        user.precisa_redefinir_senha = False
        session.commit()

        flash("Senha redefinida com sucesso. Você pode fazer login agora.", "success")
        return redirect(url_for('login'), code=302)

    return render_template('login/senhas/reset_password.html')

@app.route('/senhas/reset_senha', methods=['GET', 'POST'])
@login_required
def reset_senha():
    if request.method == 'POST':
        login_input = request.form.get('usuario', '').strip().lower()
        nova_senha = request.form.get('nova_senha', '').strip()
        confirme_senha = request.form.get('confirme_senha', '').strip()

        if nova_senha != confirme_senha:
            flash("As senhas não coincidem.", "danger")
            return redirect(url_for('reset_senha'), code=302)
        
        # Validação da política de senhas
        if len(nova_senha) < 8:
            flash("A senha deve ter pelo menos 8 caracteres.", "danger")
            return redirect(url_for('reset_password'), code=302)

        if not re.search(r'[A-Z]', nova_senha):
            flash("A senha deve conter pelo menos uma letra maiúscula.", "danger")
            return redirect(url_for('reset_password'), code=302)
        
        if not re.search(r'[a-z]', nova_senha):
            flash("A senha deve conter pelo menos uma letra minúscula.", "danger")
            return redirect(url_for('reset_password'), code=302)
        
        if not re.search(r'[!@#$%^&*(),.?":{}|<>]', nova_senha):
            flash("A senha deve conter pelo menos um símbolo.", "danger")
            return redirect(url_for('reset_password'), code=302)
        
        # Busca o usuário pelo nome de usuário ou pelo e-mail
        user = session.query(Usuario).filter(
            or_(Usuario.usuario == login_input, Usuario.email == login_input)
        ).first()

        if not user:
            flash("Usuário ou e-mail não encontrado.", "danger")
            return redirect(url_for('reset_senha'), code=302)

        hashed_password = bcrypt.hashpw(nova_senha.encode('utf-8'), bcrypt.gensalt())
        user.senha = hashed_password
        user.precisa_redefinir_senha = False
        session.commit()

        flash("Senha redefinida com sucesso. Você pode continuar agora.", "success")
        return redirect(url_for('login'), code=302)

    return render_template('login/senhas/reset_senha.html')

# Métodos de Pegar Usuario Logado ###################################################################

def get_current_user(flask_session, session):
    user_id = verify_token(flask_session.get('auth_token'))
    if user_id:
        return session.query(Usuario).filter_by(id=user_id).first()
    return None

def get_current_user_name_and_access_level(flask_session, session):
    user = get_current_user(flask_session, session)
    if user:
        access_level = user.nivel_acesso  # Supondo que `nivel_acesso` é um atributo em `Usuario`
        foto_perfil = user.foto_perfil  # Supondo que `nivel_acesso` é um atributo em `Usuario`        
        return user.usuario, access_level, foto_perfil
    return "Desconhecido", "Desconhecido"

# Métodos de Tela Inicial ###################################################################

@app.route('/index')
@login_required
def index():
    token = request.args.get('token')
    user_id = verify_token(token)
    current_user = get_current_user(flask_session, session)
    allowed_menus = get_allowed_menus_for_user(current_user)
    return render_template('menu_inicial/home.html', current_user=current_user, token=token, allowed_menus=allowed_menus)

# Métodos de Realizar Inspeção ###################################################################

mime = MimeTypes()

@app.route('/api/equipamento/<codigo_barras>')
@login_required
def get_equipamento(codigo_barras):
    try:
        equipamento = session.query(Equipamento).filter_by(codigo_barras=codigo_barras).first()
        if equipamento:
            # Busca o tipo correspondente no banco de dados
            tipo = session.query(TiposEquipamentos).filter_by(nome=equipamento.tipo).first()
            
            if not tipo:
                return jsonify({
                    'error': f'Tipo de equipamento "{equipamento.tipo}" não encontrado na tabela de tipos',
                    'id': equipamento.id,
                    'tipo': equipamento.tipo,
                    'classe': equipamento.classe,
                    'localizacao': equipamento.localizacao,
                    'status': equipamento.status,
                    'codigo_barras': equipamento.codigo_barras
                }), 404
                
            return jsonify({
                'id': equipamento.id,
                'tipo': equipamento.tipo,
                'tipo_id': tipo.id,  # ID do tipo encontrado
                'classe': equipamento.classe,
                'localizacao': equipamento.localizacao,
                'status': equipamento.status,
                'codigo_barras': equipamento.codigo_barras
            })
        return jsonify({'error': 'Equipamento não encontrado'}), 404
    except SQLAlchemyError as e:
        return jsonify({'error': str(e)}), 500
    finally:
        session.close()

@app.route('/api/gerar_numero_inspecao', methods=['GET'])
@login_required
def gerar_numero_inspecao():
    config = load_config()
    prefixo = config.get('INSPECTION_PREFIX', 'INS-')
    
    ultimo_numero = session.query(Registro).order_by(Registro.numero_inspecao.desc()).first()

    if ultimo_numero:
        try:
            # Extrai apenas a parte numérica do último número
            ultimo_numero = int(ultimo_numero.numero_inspecao.split('-')[1])
            novo_numero = f"{prefixo}{str(ultimo_numero + 1).zfill(5)}"
        except (IndexError, ValueError):
            # Caso o formato esteja incorreto, começa do 1 com o prefixo atual
            novo_numero = f"{prefixo}00001"
    else:
        novo_numero = f"{prefixo}00001"
    
    return jsonify({'numero_inspecao': novo_numero})

@app.route('/inspecoes/nova_inspecao', methods=['GET', 'POST'])
@login_required
@permission_required('inspecao')
def inspecao():
    if request.method == 'POST':
        try:
            usuario_logado, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)          
            config = load_config()
            upload_folder = config['UPLOAD_FOLDER']

            if 'fotos' not in request.files:
                return jsonify({'error': 'Nenhuma foto foi enviada.'}), 400

            fotos = request.files.getlist('fotos')
            for foto in fotos:
                if foto.filename == '':
                    return jsonify({'error': 'Nome de arquivo vazio.'}), 400
            
            # Process location data
            location_data = None
            if 'localizacao' in request.form:
                try:
                    location_data = json.loads(request.form['localizacao'])
                    if not all(key in location_data for key in ['latitude', 'longitude']):
                        location_data = None
                    else:
                        # Add additional metadata
                        location_data['usuario'] = usuario_logado
                        location_data['data_registro'] = datetime.now().isoformat()
                except json.JSONDecodeError:
                    location_data = None

            data = request.form
            equipamento_id = data.get('equipamento_id')
            equipamento_codigo_barras = data.get('equipamento_codigo_barras')
            tipo_equipamento = data.get('tipo_equipamento')
            classe_equipamento = data.get('classe_equipamento')
            localizacao_equipamento = data.get('localizacao_equipamento')
            status_equipamento = data.get('status_equipamento')
            criticidade_inspecao = data.get('criticidade_inspecao')            
            responsavel = data.get('responsavel')
            observacoes = data.get('observacoes', '')
            numero_inspecao = data.get('numero_inspecao')
            status_inspecao = request.form.get('status_inspecao', 'Pendente')
            motivo = request.form.get('motivo', '')
            tipo_equipamento_id = data.get('tipo_equipamento_id')

            if not numero_inspecao:
                return jsonify({'error': 'Número de inspeção não foi gerado.'}), 400
            
            if not all([equipamento_id, equipamento_codigo_barras, tipo_equipamento, 
                        localizacao_equipamento, status_equipamento, responsavel, numero_inspecao]):
                return jsonify({'error': 'Todos os campos obrigatórios devem ser preenchidos.'}), 400

            equipamento = session.get(Equipamento, equipamento_id)
            if not equipamento:
                return jsonify({'error': 'Equipamento não encontrado.'}), 404

            # Check for pending records
            registros_existente = session.query(Registro).filter(
                Registro.equipamento_id == equipamento_id,
                Registro.status_inspecao == "Pendente"
            ).order_by(Registro.data_registro.desc()).limit(15).all()

            # Salvamos a localização APENAS se não houver registros pendentes
            location_id = None
            if location_data and not registros_existente:
                loc = LocalizacaoUsuario(
                    latitude=location_data['latitude'],
                    longitude=location_data['longitude'],
                    usuario=usuario_logado,
                    data_registro=datetime.now(),
                    metadados={
                        'accuracy': location_data.get('accuracy'),
                        'device_info': location_data.get('device_info')
                    }
                )
                session.add(loc)
                session.flush()
                location_id = loc.id

            if registros_existente:
                if len(registros_existente) > 1:
                    return jsonify({
                        'error': 'Este equipamento possui mais de uma inspeção pendente.',
                        'registros_pendentes': len(registros_existente),
                        'localizacao_salva': False,  # Não salvamos ainda
                        'location_data': location_data  # Enviamos os dados para usar depois
                    }), 400
                
                ultimo_registro = registros_existente[0]
                itens_pendentes = session.query(ChecklistItem.description).filter(
                    ChecklistItem.registro_id == ultimo_registro.id,
                    ChecklistItem.checked == False
                ).all()
                
                itens_pendentes_lista = [item[0] for item in itens_pendentes]
                
                registros_json = [{
                    'numero_inspecao': registro.numero_inspecao,
                    'tipo': registro.tipo,
                    'classe': registro.classe,
                    'localizacao': registro.localizacao,
                    'data_inspecao': registro.data_inspecao.strftime("%d-%m-%Y") if registro.data_inspecao else None,
                    'data_validade_inspecao': registro.data_validade_inspecao.strftime("%d-%m-%Y") if registro.data_validade_inspecao else None,
                    'itens_pendentes': itens_pendentes_lista if registro.id == ultimo_registro.id else [],
                    'localizacao_salva': False,  # Indica que ainda não foi salva
                    'location_data': location_data  # Envia os dados brutos
                } for registro in registros_existente]

                return jsonify({
                    'registros_existente': registros_json,
                    'itens_pendentes': itens_pendentes_lista,
                    'localizacao_data': location_data  # Mantemos os dados para uso posterior
                })

            # Create new inspection if no pending records
            if not tipo_equipamento_id:
                return jsonify({'error': 'Tipo de equipamento não identificado.'}), 400
            
            inspecao_data = get_checklist_items(tipo_equipamento_id, request.form)
            status_inspecao, motivo, data_encerramento_inspecao = verificar_status_inspecao(tipo_equipamento_id, inspecao_data)

            registro = criar_registro(
                equipamento_id=equipamento_id,
                tipo_equipamento=tipo_equipamento,
                classe_equipamento=classe_equipamento,
                localizacao_equipamento=localizacao_equipamento,
                status_equipamento=status_equipamento,
                data_encerramento_inspecao=data_encerramento_inspecao,
                criticidade_inspecao=criticidade_inspecao,                
                responsavel=responsavel,
                observacoes=observacoes,
                inspecao_data=inspecao_data,
                numero_inspecao=numero_inspecao,
                status_inspecao=status_inspecao,
                motivo=motivo
            )
            
            if not registro:
                return jsonify({'error': 'Falha ao criar registro.'}), 500
            
            # Process photos
            inspecao_folder = os.path.join(
                upload_folder,
                f"{registro.numero_inspecao}-{registro.data_inspecao.strftime('%d-%m-%Y')}"
            )
            os.makedirs(inspecao_folder, exist_ok=True)

            for foto in fotos:
                filename = secure_filename(foto.filename)
                filepath = os.path.join(inspecao_folder, filename)
                foto.save(filepath)
                session.add(FotoRegistro(registro_id=registro.id, caminho=filepath))

            # Update equipment
            equipamento.data_ultima_inspecao = datetime.now().date()
            equipamento.data_proxima_inspecao = datetime.now().date() + timedelta(days=30)
            
            # Link location to inspection if available (só acontece se não houver registros pendentes)
            if location_data and location_id:
                loc = session.query(LocalizacaoUsuario).get(location_id)
                if loc:
                    loc.registro_id = registro.id
            
            session.commit()

            return jsonify({
                'success': True,
                'message': f'Inspeção {registro.numero_inspecao} registrada com sucesso!',
                'registro_id': registro.id,
                'localizacao_salva': bool(location_data),
                'location_id': location_id if location_data else None
            })
 
        except Exception as e:
            session.rollback()
            return jsonify({'error': str(e)}), 500
        finally:
            session.close()

    current_user = get_current_user(flask_session, session)
    allowed_menus = get_allowed_menus_for_user(current_user)      
    return render_template('inspecoes/nova_inspecao.html', current_user=current_user, allowed_menus=allowed_menus)

@app.route('/inspecoes/acao', methods=['POST'])
@login_required
def inspecao_acao():
    try:
        # Process location data
        location_data = None
        location_id = request.form.get('location_id')  # Recebe o ID da localização já existente
        
        if 'localizacao' in request.form and not location_id:
            try:
                location_data = json.loads(request.form['localizacao'])
                if not all(key in location_data for key in ['latitude', 'longitude']):
                    location_data = None
                else:
                    usuario_logado = get_current_user_name_and_access_level(flask_session, session)[0]
                    location_data['usuario'] = usuario_logado
                    location_data['data_registro'] = datetime.now().isoformat()
            except json.JSONDecodeError:
                location_data = None


        # Form data
        equipamento_id = request.form.get('equipamento_id')
        acao = request.form.get('acao')
        tipo_equipamento = request.form.get('tipo_equipamento')
        classe_equipamento = request.form.get('classe_equipamento')
        localizacao_equipamento = request.form.get('localizacao_equipamento')
        status_equipamento = request.form.get('status_equipamento')
        criticidade_inspecao = request.form.get('criticidade_inspecao')        
        responsavel = request.form.get('responsavel')
        observacoes = request.form.get('observacoes', '')
        numero_inspecao = request.form.get('numero_inspecao')
        status_inspecao = request.form.get('status_inspecao', 'Pendente')
        motivo = request.form.get('motivo', '')

        config = load_config()
        upload_folder = config['UPLOAD_FOLDER']
                
        if not numero_inspecao:
            return jsonify({'error': 'Número de inspeção não foi gerado.'}), 400
        
        # Process checklist
        inspecao_data = {}
        for key in request.form:
            if key.startswith('inspecao[') and key.endswith(']'):
                item_name = key[9:-1]
                inspecao_data[item_name] = request.form[key] == '1'

        status_inspecao, motivo, data_encerramento_inspecao = verificar_status_inspecao(tipo_equipamento, inspecao_data)

        if not equipamento_id or not acao:
            return jsonify({'error': 'Dados inválidos.'}), 400

        fotos = request.files.getlist('fotos')
        usuario_logado = get_current_user_name_and_access_level(flask_session, session)[0]

        if acao == 'sobrescrever':
            try:
                # Log
                usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
                log_entry = EntradaLog(
                    mensagem=f"(WEB) Usuário {usuario} iniciou sobrescrita para equipamento {equipamento_id}",
                    usuario=usuario,
                    nivel_acesso=nivel_acesso,
                    tipo_log="WEB - inspeção - Sobrescrever",
                    data_geracao=datetime.now()
                )
                session.add(log_entry)

                # Update pending records
                registros_pendentes = session.query(Registro).filter(
                    Registro.equipamento_id == equipamento_id,
                    Registro.status_inspecao == "Pendente"
                ).all()

                for registro in registros_pendentes:
                    registro.status_inspecao = "Finalizada"
                    registro.motivo_acao = f"Finalizada pela inspeção {numero_inspecao}"
                    registro.data_validade_inspecao = datetime.now().date() + timedelta(days=30)
                    registro.data_encerramento_inspecao = data_encerramento_inspecao

                # Create new record
                novo_registro = criar_registro(
                    equipamento_id=equipamento_id,
                    tipo_equipamento=tipo_equipamento,
                    classe_equipamento=classe_equipamento,
                    localizacao_equipamento=localizacao_equipamento,
                    status_equipamento=status_equipamento,
                    criticidade_inspecao=criticidade_inspecao,                    
                    responsavel=responsavel,
                    observacoes=observacoes,
                    inspecao_data=inspecao_data,
                    numero_inspecao=numero_inspecao,
                    data_encerramento_inspecao=data_encerramento_inspecao,
                    status_inspecao=status_inspecao,
                    motivo=motivo
                )

                # Link location to new record if available
                if location_data:
                    session.add(LocalizacaoUsuario(
                        registro_id=novo_registro.id,
                        latitude=location_data['latitude'],
                        longitude=location_data['longitude'],
                        usuario=usuario_logado,
                        data_registro=datetime.now(),
                        metadados={
                            'accuracy': location_data.get('accuracy'),
                            'device_info': location_data.get('device_info')
                        }
                    ))

                # Process photos
                if fotos:
                    inspecao_folder = os.path.join(
                        upload_folder,
                        f"{novo_registro.numero_inspecao}-{datetime.now().date().strftime('%d-%m-%Y')}"
                    )
                    os.makedirs(inspecao_folder, exist_ok=True)
                    
                    for foto in fotos:
                        if foto.filename:
                            filename = secure_filename(foto.filename)
                            foto.save(os.path.join(inspecao_folder, filename))
                            session.add(FotoRegistro(
                                registro_id=novo_registro.id,
                                caminho=os.path.join(inspecao_folder, filename)
                            ))

                # Update equipment
                equipamento = session.get(Equipamento, equipamento_id)
                if equipamento:
                    equipamento.data_ultima_inspecao = datetime.now().date()
                    equipamento.data_proxima_inspecao = datetime.now().date() + timedelta(days=30)

                session.commit()

                return jsonify({
                    'success': True,
                    'message': f'Inspeção {novo_registro.numero_inspecao} registrada com sucesso!',
                    'registro_id': novo_registro.id,
                    'localizacao_salva': bool(location_data),
                    'location_id': location_id
                })

            except Exception as e:
                session.rollback()
                return jsonify({'error': str(e)}), 500

        elif acao == 'ignorar':
            try:
                usuario, nivel_acesso, foto_perfil_url = get_current_user_name_and_access_level(flask_session, session)
                
                # Log before creating new record
                data_atual_log = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
                mensagem_log = (
                    f"(WEB) O Usuário {usuario} ({nivel_acesso}) optou por ignorar registros pendentes "
                    f"e criar um novo registro para o equipamento ID {equipamento_id} em {data_atual_log}.\n"
                )
                log_entry = EntradaLog(
                    mensagem=mensagem_log,
                    usuario=usuario,
                    nivel_acesso=nivel_acesso,
                    tipo_log="WEB - inspeção - Ignorar e Criar Registro",
                    data_geracao=datetime.now()
                )
                session.add(log_entry)
                session.commit()

                # Create new record
                novo_registro = criar_registro(
                    equipamento_id=equipamento_id,
                    tipo_equipamento=tipo_equipamento,
                    classe_equipamento=classe_equipamento,
                    localizacao_equipamento=localizacao_equipamento,
                    status_equipamento=status_equipamento,
                    criticidade_inspecao=criticidade_inspecao,    
                    data_encerramento_inspecao=data_encerramento_inspecao,               
                    responsavel=responsavel,
                    observacoes=observacoes,
                    inspecao_data=inspecao_data,
                    numero_inspecao=numero_inspecao,
                    status_inspecao=status_inspecao,
                    motivo=motivo
                )

                # Link location to new record if available
                if location_data:
                    session.add(LocalizacaoUsuario(
                        registro_id=novo_registro.id,
                        latitude=location_data['latitude'],
                        longitude=location_data['longitude'],
                        usuario=usuario,
                        data_registro=datetime.now(),
                        metadados={
                            'accuracy': location_data.get('accuracy'),
                            'device_info': location_data.get('device_info')
                        }
                    ))

                if not novo_registro:
                    mensagem_log += "  - Falha ao criar novo registro.\n"
                    log_entry.mensagem = mensagem_log
                    session.commit()
                    return jsonify({'error': 'Falha ao criar registro.'}), 500

                # Update equipment
                equipamento = session.get(Equipamento, equipamento_id)
                if equipamento:
                    equipamento.data_ultima_inspecao = datetime.now().date()
                    equipamento.data_proxima_inspecao = datetime.now().date() + timedelta(days=30)
                    session.commit()

                # Save photos if any
                if fotos:
                    data_inspecao = datetime.now().date()
                    inspecao_folder = os.path.join(
                        upload_folder,
                        f"{novo_registro.numero_inspecao}-{data_inspecao.strftime('%d-%m-%Y')}"
                    )
                    os.makedirs(inspecao_folder, exist_ok=True)

                    for foto in fotos:
                        if foto and foto.filename != '':
                            filename = secure_filename(foto.filename)
                            filepath = os.path.join(inspecao_folder, filename)
                            foto.save(filepath)
                            foto_registro = FotoRegistro(
                                registro_id=novo_registro.id,
                                caminho=filepath
                            )
                            session.add(foto_registro)
                    session.commit()

                # Success log
                mensagem_log += (
                    f"  - Novo registro criado com ID: {novo_registro.id}\n"
                    f"  - Número de inspeção: {novo_registro.numero_inspecao}\n"
                )
                log_entry.mensagem = mensagem_log
                session.commit()

                return jsonify({
                    'success': True,
                    'message': f'Inspeção {novo_registro.numero_inspecao} registrada com sucesso!',
                    'registro_id': novo_registro.id,
                    'localizacao_salva': bool(location_data),
                    'location_id': location_id
                })

            except Exception as e:
                session.rollback()
                return jsonify({'error': str(e)}), 500

        elif acao == 'cancelar':
            # Log cancellation
            usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
            data_atual_log = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
            mensagem_log = (
                f"(WEB) O Usuário {usuario} ({nivel_acesso}) cancelou a operação de inspeção "
                f"para o equipamento ID {equipamento_id} em {data_atual_log}.\n"
            )
            log_entry = EntradaLog(
                mensagem=mensagem_log,
                usuario=usuario,
                nivel_acesso=nivel_acesso,
                tipo_log="WEB - inspeção - Cancelar Operação",
                data_geracao=datetime.now()
            )
            session.add(log_entry)
            session.commit()

            return jsonify({
                'success': True, 
                'message': 'Operação cancelada.',
                'localizacao_salva': bool(location_data),
                'location_id': location_id
            })

        else:
            return jsonify({'error': 'Ação inválida.'}), 400

    except Exception as e:
        session.rollback()
        return jsonify({'error': str(e)}), 500
    
def criar_registro(equipamento_id, tipo_equipamento, classe_equipamento, localizacao_equipamento,  
                   status_equipamento, responsavel, observacoes, inspecao_data, status_inspecao, criticidade_inspecao, motivo, data_encerramento_inspecao, fotos=None, numero_inspecao=None):
    try:
        config = load_config()
        data_atual = datetime.now().date()

        # Busca validade da inspeção
        tipo_info = session.query(TiposEquipamentos).filter_by(nome=tipo_equipamento).first()
        validade_dias = tipo_info.validade_inspecao if tipo_info and tipo_info.validade_inspecao else 30

        # Cria registro
        registro = Registro(
            equipamento_id=equipamento_id,
            tipo=tipo_equipamento,
            classe=classe_equipamento,
            localizacao=localizacao_equipamento,
            status_equipamento=status_equipamento,
            criticidade_inspecao=criticidade_inspecao,            
            responsavel=responsavel,
            observacoes=observacoes,
            data_registro=data_atual,
            data_inspecao=data_atual,
            data_validade_inspecao=data_atual + timedelta(days=validade_dias),
            status_inspecao=status_inspecao,
            motivo_acao=motivo,
            usuario_geracao=get_current_user_name_and_access_level(flask_session, session)[0],
            numero_inspecao=numero_inspecao,
            data_encerramento_inspecao=data_encerramento_inspecao  # Novo campo adicionado
        )

        session.add(registro)
        session.flush()  # Para obter o ID antes do commit

        # Adiciona itens do checklist
        for item_name, checked in inspecao_data.items():
            session.add(ChecklistItem(
                registro_id=registro.id,
                description=item_name,
                checked=checked
            ))

        return registro

    except SQLAlchemyError as e:
        session.rollback()
        raise

@app.route('/api/checklist_tipo/<int:tipo_id>')
@login_required
def get_checklist_tipo(tipo_id):
    try:
        # Verifica se o tipo existe antes de buscar os itens
        tipo_existe = session.query(TiposEquipamentos).filter_by(id=tipo_id).first()
        if not tipo_existe:
            return jsonify({'error': 'Tipo de equipamento não encontrado'}), 404
            
        checklist_items = session.query(ChecklistTipo).filter(
            ChecklistTipo.tipo_id == tipo_id
        ).order_by(ChecklistTipo.ordem).all()
        
        return jsonify([{
            'id': item.id,
            'nome_item': item.nome_item,
            'ordem': item.ordem,
            'obrigatorio': item.obrigatorio,
            'tipo_id': item.tipo_id
        } for item in checklist_items])
    except Exception as e:
        return jsonify({'error': str(e)}), 500

def normalizar_nome_item(nome):
    return unidecode(nome.strip().replace(' ', '_').lower())

def get_checklist_items(tipo_equipamento_id, data):
    try:

        # Corrige tipo de ID vindo como string
        tipo_id = None
        if isinstance(tipo_equipamento_id, str):
            if tipo_equipamento_id.isdigit():
                tipo_id = int(tipo_equipamento_id)
            else:
                # Busca pelo nome do tipo
                tipo_obj = session.query(TiposEquipamentos).filter(
                    func.lower(TiposEquipamentos.nome) == func.lower(tipo_equipamento_id.strip())
                ).first()
                
                if tipo_obj:
                    tipo_id = tipo_obj.id
                else:
                    return {}
        else:
            tipo_id = tipo_equipamento_id

        if tipo_id is None:
            return {}

        # Busca os itens do checklist no banco de dados
        checklist_items = session.query(ChecklistTipo).filter(
            ChecklistTipo.tipo_id == tipo_id
        ).all()

        if hasattr(data, 'getlist'):
            data_dict = {k: v for k, v in data.items()}
        else:
            data_dict = data

        # Debug: Mostrar os dados recebidos
        for key, value in data_dict.items():
            if key.startswith('inspecao_'):
                print(f"  {key}: {value}")

        items = {
            item.nome_item: data_dict.get(f'inspecao_{normalizar_nome_item(item.nome_item)}', '0') == '1'
            for item in checklist_items
        }

        return items

    except Exception as e:
        print(f"Erro ao obter itens do checklist: {str(e)}", exc_info=True)
        return {}
    
def verificar_status_inspecao(tipo_equipamento_id, inspecao_data):
    """
    Verifica o status da inspeção com base nos itens verificados.
    Retorna o status, o motivo e a data de encerramento (quando finalizada)
    """
    # Se for string, tenta converter ou buscar pelo nome
    if isinstance(tipo_equipamento_id, str):
        # Se for número em string, converte direto
        if tipo_equipamento_id.isdigit():
            tipo_equipamento_id = int(tipo_equipamento_id)
        else:
            tipo_obj = session.query(TiposEquipamentos).filter_by(nome=tipo_equipamento_id).first()
            if tipo_obj:
                tipo_equipamento_id = tipo_obj.id
            else:
                return "Pendente", f"Tipo de equipamento '{tipo_equipamento_id}' não encontrado.", None

    # Obter todos os itens do checklist para este tipo de equipamento
    itens_checklist = session.query(ChecklistTipo.nome_item, ChecklistTipo.obrigatorio)\
                             .filter(ChecklistTipo.tipo_id == tipo_equipamento_id)\
                             .all()
    
    if not itens_checklist:
        return "Finalizada", "Nenhum item de checklist cadastrado para este tipo de equipamento", datetime.now().date()
    
    itens_pendentes = []
    itens_obrigatorios_pendentes = []

    for item in itens_checklist:
        item_name = item.nome_item
        is_obrigatorio = item.obrigatorio
        
        # Verifica se o item está marcado no inspecao_data
        is_checked = inspecao_data.get(item_name, False)
        
        if not is_checked:
            itens_pendentes.append(item_name)
            if is_obrigatorio:
                itens_obrigatorios_pendentes.append(item_name)

    # A inspeção só é finalizada se TODOS os itens estiverem marcados
    if not itens_pendentes:
        return "Finalizada", "Todos os itens do checklist foram verificados", datetime.now().date()
    else:
        # Se houver itens obrigatórios pendentes, destaca isso no motivo
        if itens_obrigatorios_pendentes:
            if len(itens_obrigatorios_pendentes) > 3:
                motivo = f"{len(itens_obrigatorios_pendentes)} itens obrigatórios pendentes"
            else:
                motivo = f"Itens obrigatórios pendentes: {', '.join(itens_obrigatorios_pendentes)}"
        else:
            if len(itens_pendentes) > 3:
                motivo = f"{len(itens_pendentes)} itens pendentes"
            else:
                motivo = f"Itens pendentes: {', '.join(itens_pendentes)}"
        
        return "Pendente", motivo, None
     
@app.route('/api/inspecoes_pendentes/<int:equipamento_id>')
@login_required
def obter_pendencias(equipamento_id):
    try:
        # Busca o último registro pendente
        ultimo_registro = session.query(Registro).filter(
            Registro.equipamento_id == equipamento_id,
            Registro.status_inspecao == "Pendente"
        ).order_by(Registro.data_registro.desc()).first()

        if not ultimo_registro:
            return jsonify({'pendencias': []})

        # Busca os itens pendentes diretamente do banco
        itens_pendentes = session.query(ChecklistItem.description).filter(
            ChecklistItem.registro_id == ultimo_registro.id,
            ChecklistItem.checked == False
        ).all()

        # Extrai os nomes diretamente sem formatação adicional
        pendencias_formatadas = [item[0] for item in itens_pendentes]

        return jsonify({'pendencias': pendencias_formatadas})

    except Exception as e:
        return jsonify({'error': str(e)}), 500
 
# Métodos de Minhas Pendencias ###################################################################

@app.route('/inspecoes/minhas_pendencias')
@login_required
@permission_required('pendencias_inspecoes')
def pendencias_inspecoes():
    current_user = get_current_user(flask_session, session)
    allowed_menus = get_allowed_menus_for_user(current_user)
    return render_template('inspecoes/pendencias_inspecoes.html', current_user=current_user, allowed_menus=allowed_menus)

@app.route('/api/minhas_inspecoes', methods=['GET'])
@login_required
def get_minhas_inspecoes():
    try:
        # Obtém o usuário logado
        usuario_logado, nivel_acesso, foto_perfil_url = get_current_user_name_and_access_level(flask_session, session)
        
        # Busca o usuário no banco de dados
        usuario = session.query(Usuario).filter_by(usuario=usuario_logado).first()

        if not usuario:
            return jsonify({'error': 'Usuário não encontrado'}), 404
                
        termo_pesquisa = request.args.get('search', '').strip()
        status_filter = request.args.get('status', '').strip()
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 1500))

        if page < 1 or limit < 1:
            return jsonify({'error': 'Parâmetros de paginação inválidos.'}), 400

        # Filtra apenas as inspeções do usuário logado usando o nome completo
        query = (
            session.query(Registro, Equipamento)
            .join(Equipamento, Registro.equipamento_id == Equipamento.id)
            .filter(Registro.responsavel == usuario.nome)
        )

        # Aplica filtros adicionais
        if termo_pesquisa:
            query = query.filter(
                or_(
                    Equipamento.tag_equipamento.ilike(f'%{termo_pesquisa}%'),
                    Registro.localizacao.ilike(f'%{termo_pesquisa}%'),
                    Registro.responsavel.ilike(f'%{termo_pesquisa}%')
                )
            )

        if status_filter:
            if status_filter not in ['Pendente', 'Finalizada', 'Cancelada']:
                return jsonify({'error': 'Status de inspeção inválido.'}), 400
            query = query.filter(Registro.status_inspecao == status_filter)

        total_inspecoes = query.count()
        resultados = query.offset((page - 1) * limit).limit(limit).all()

        # Formata os dados para retorno
        inspecoes_data = [
            {
                'id': registro.id,
                'numero_inspecao': registro.numero_inspecao,
                'equipamento_id': registro.equipamento_id,
                'equipamento': f"{equipamento.tag_equipamento} - {registro.localizacao}",
                'tag_equipamento': equipamento.tag_equipamento,
                'tipo_equipamento': equipamento.tipo,
                'localizacao': equipamento.localizacao,
                'responsavel': registro.responsavel,
                'criticidade_inspecao': registro.criticidade_inspecao,                  
                'data_registro': registro.data_registro.strftime('%d-%m-%Y'),
                'data_validade_inspecao': registro.data_validade_inspecao.strftime('%d-%m-%Y'),
                'status_inspecao': registro.status_inspecao,
                'status_equipamento': registro.status_equipamento,
                'observacoes': registro.observacoes
            }
            for registro, equipamento in resultados
        ]

        return jsonify({
            'data': inspecoes_data,
            'total': total_inspecoes,
            'page': page,
            'limit': limit
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Métodos de Dashboards ###################################################################

@app.route('/relatorios/visualizar_dashboards')
@login_required
@permission_required('visualizar_dashboards')
def visualizar_dashboards():
    current_user = get_current_user(flask_session, session)
    return render_template('relatorios/visualizar_dashboards.html', current_user=current_user)

# Métodos de Inspeções Realizadas ###################################################################

def is_file_locked(filepath):
    """Verifica se um arquivo está em uso por outro processo"""
    for proc in psutil.process_iter(['pid', 'name', 'open_files']):
        try:
            for item in proc.info['open_files'] or []:
                if filepath.lower() == item.path.lower():
                    return True
        except (psutil.NoSuchProcess, psutil.AccessDenied):
            continue
    return False

def wait_for_file_unlock(filepath, timeout=10, check_interval=0.5):
    """Aguarda até que o arquivo esteja desbloqueado ou timeout seja atingido"""
    start_time = time.time()
    while time.time() - start_time < timeout:
        if not is_file_locked(filepath):
            return True
        time.sleep(check_interval)
    return False

def validate_file(file_path):
    """Valida se o arquivo é válido e dentro dos limites de tamanho"""
    try:
        if not os.path.exists(file_path):
            return False, "Arquivo não encontrado"
        
        # Obter configurações do sistema
        config = load_config()
        max_file_size = config['MAX_FILE_SIZE'] * 1024 * 1024  # Converter MB para bytes
        allowed_extensions = config['ALLOWED_EXTENSIONS']
        
        if os.path.getsize(file_path) > max_file_size:
            return False, f"Arquivo muito grande (limite: {config['MAX_FILE_SIZE']}MB)"
        
        ext = os.path.splitext(file_path)[1].lower()
        
        # Se for imagem, valida como imagem
        if ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:  # Extensões de imagem fixas
            try:
                with Image.open(file_path) as img:
                    img.verify()
                return True, "OK_IMAGE"
            except Exception as e:
                return False, f"Erro na validação de imagem: {str(e)}"
        else:
            # Para outros tipos de arquivo, verifica na lista de extensões permitidas
            if ext in allowed_extensions:
                return True, "OK_FILE"
            else:
                return False, f"Extensão de arquivo não suportada ({ext})"
                
    except Exception as e:
        return False, f"Erro na validação: {str(e)}"

def processar_arquivo(paragraph, file_path):
    """Processa um arquivo para inserção no documento com tratamento robusto de erros"""
    temp_dir = tempfile.mkdtemp()
    try:
        # Validação inicial do arquivo
        is_valid, validation_msg = validate_file(file_path)
        if not is_valid:
            paragraph.add_run(f"[Arquivo inválido: {validation_msg}]")
            return False

        ext = os.path.splitext(file_path)[1].lower()
        file_name = os.path.basename(file_path)
        
        # Se for imagem, processa como imagem
        if "OK_IMAGE" in validation_msg:
            temp_file_path = os.path.join(temp_dir, file_name)
            shutil.copy2(file_path, temp_file_path)

            # Inserção no documento
            run = paragraph.add_run()
            try:
                run.add_picture(temp_file_path, width=Inches(4))
                paragraph.add_run().add_break()
                paragraph.add_run(file_name)
                paragraph.add_run().add_break()
                return True
            except Exception as img_error:
                print(f"Erro ao inserir imagem {temp_file_path}: {img_error}")
                paragraph.add_run(f"[Erro ao processar imagem: {file_name}]")
                return False
        else:
            # Para outros tipos de arquivo, insere apenas o nome como link
            paragraph.add_run().add_break()
            run = paragraph.add_run(f"[Arquivo anexado: {file_name}]")
            # Aqui você pode adicionar lógica para criar um hyperlink se necessário
            return True
            
    except Exception as e:
        print(f"Erro ao processar arquivo {file_path}: {e}")
        paragraph.add_run(f"[Erro ao processar: {os.path.basename(file_path)}]")
        return False
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

def substituir_ids_e_gerar_docx(caminho_docx, caminho_saida, dados_substituicao, checklist_items=None, arquivos=[]):
    """Gera documento DOCX com substituição de marcadores e tratamento robusto de erros"""
    try:
        # Converter todos os valores para string
        dados_substituicao = {k: str(v) if v is not None else "" for k, v in dados_substituicao.items()}
        
        # Verificar se o modelo existe
        if not os.path.exists(caminho_docx):
            raise FileNotFoundError(f"Modelo DOCX não encontrado em {caminho_docx}")
            
        # Verificar se o arquivo de saída está bloqueado
        if os.path.exists(caminho_saida) and is_file_locked(caminho_saida):
            if not wait_for_file_unlock(caminho_saida):
                raise IOError(f"Arquivo de saída está bloqueado por outro processo: {caminho_saida}")
        
        doc = Document(caminho_docx)
        
        # Processar parágrafos
        for paragraph in doc.paragraphs:
            print("Dados de substituição recebidos:")
            for chave, valor in dados_substituicao.items():
                if chave in paragraph.text:
                    if chave == "{ARQUIVOS}":
                        for run in paragraph.runs:
                            run.text = run.text.replace(chave, "")
                        if arquivos:
                            for arquivo in arquivos:
                                processar_arquivo(paragraph, arquivo)
                    else:
                        try:
                            run.text = run.text.replace(chave, valor)
                        except Exception as e:
                            print(f"Erro ao substituir {chave} por {valor}: {e}")
                            run.text = run.text.replace(chave, "[VALOR INVÁLIDO]")

        # Criar tabela de checklist no lugar do marcador {TABELAS}
        if checklist_items:
            for tabela in doc.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for par in celula.paragraphs:
                            if "{TABELAS}" in par.text:
                                # Encontrou a célula que precisa da tabela
                                
                                # 1. Limpar a célula
                                celula.text = ""

                                # 2. Criar tabela DENTRO da célula
                                nova_tabela = celula.add_table(rows=1, cols=3)
                                nova_tabela.style = 'Table Grid'
                                hdr_cells = nova_tabela.rows[0].cells
                                hdr_cells[0].text = "Item"
                                hdr_cells[1].text = "Conforme"
                                hdr_cells[2].text = "Não Conforme"

                                for item in checklist_items:
                                    row_cells = nova_tabela.add_row().cells
                                    row_cells[0].text = getattr(item, 'description', '') or getattr(item, 'descricao', 'Item')
                                    row_cells[1].text = '✓' if getattr(item, 'checked', False) else ''
                                    row_cells[2].text = 'X' if not getattr(item, 'checked', False) else ''

                                break

        # Processar tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for chave, valor in dados_substituicao.items():
                            if chave in paragraph.text:
                                if chave == "{FOTOS}":
                                    paragraph.text = paragraph.text.replace(chave, "")
                                    if arquivos:
                                        for foto in arquivos:
                                            processar_arquivo(paragraph, foto)
                                else:
                                    paragraph.text = paragraph.text.replace(chave, valor)
                    
                    # Processar tabelas aninhadas
                    for nested_table in cell.tables:
                        for nested_row in nested_table.rows:
                            for nested_cell in nested_row.cells:
                                for nested_paragraph in nested_cell.paragraphs:
                                    for chave, valor in dados_substituicao.items():
                                        if chave in nested_paragraph.text:
                                            nested_paragraph.text = nested_paragraph.text.replace(chave, valor)
        
        # Garantir que o diretório de saída existe
        os.makedirs(os.path.dirname(caminho_saida), exist_ok=True)
        
        # Tentar salvar com tratamento de erros de permissão
        try:
            doc.save(caminho_saida)
        except PermissionError:
            # Tentar novamente após pequeno delay
            time.sleep(0.5)
            doc.save(caminho_saida)
        
        # Verificar se o arquivo foi criado
        if not os.path.exists(caminho_saida):
            raise IOError(f"Falha ao salvar o documento em {caminho_saida}")
            
        return True
    except Exception as erro:
        print(f"Erro crítico ao gerar documento: {erro}")
        raise

def converter_docx_para_pdf(caminho_docx, caminho_pdf):
    """Converte DOCX para PDF com tratamento robusto de erros e timeout"""
    try:
        pythoncom.CoInitialize()
        
        # Obter configurações do sistema
        config = load_config()
        pdf_timeout = config['PDF_CONVERSION_TIMEOUT']
        
        # Verificar se o arquivo DOCX existe e não está bloqueado
        if not os.path.exists(caminho_docx):
            raise FileNotFoundError(f"Arquivo DOCX não encontrado em {caminho_docx}")
        
        if is_file_locked(caminho_docx):
            if not wait_for_file_unlock(caminho_docx):
                raise IOError(f"Arquivo DOCX está bloqueado por outro processo: {caminho_docx}")
        
        # Configurar timeout para a conversão
        start_time = datetime.now()
        
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        try:
            doc = word.Documents.Open(os.path.abspath(caminho_docx))
            
            # Verificar timeout periodicamente
            if (datetime.now() - start_time).seconds > pdf_timeout:
                raise TimeoutError(f"Tempo limite excedido na conversão para PDF ({pdf_timeout} segundos)")
                
            # Verificar se o arquivo PDF está bloqueado
            if os.path.exists(caminho_pdf) and is_file_locked(caminho_pdf):
                if not wait_for_file_unlock(caminho_pdf):
                    raise IOError(f"Arquivo PDF de destino está bloqueado: {caminho_pdf}")
            
            doc.SaveAs(os.path.abspath(caminho_pdf), FileFormat=17)
            doc.Close()
            
            # Verificar se o PDF foi criado
            if not os.path.exists(caminho_pdf):
                raise IOError(f"Arquivo PDF não foi gerado em {caminho_pdf}")
                
            return True
        finally:
            word.Quit()
            pythoncom.CoUninitialize()
            
    except Exception as e:
        print(f"Erro na conversão para PDF: {e}")
        raise

def encontrar_caminho_real(caminho_armazenado, base_dir):
    """
    Tenta encontrar o caminho real da foto a partir do caminho armazenado no banco.
    """
    # Normalizar e garantir que o caminho está correto
    caminho_testar = os.path.normpath(os.path.join(base_dir, caminho_armazenado))

    # Verificar se o arquivo existe diretamente
    if os.path.exists(caminho_testar):
        return caminho_testar

    print(f"AVISO: Foto não encontrada no caminho esperado - {caminho_testar}")
    
    # Procurar na pasta base por um arquivo com o mesmo nome
    nome_arquivo = os.path.basename(caminho_testar)
    for root, _, files in os.walk(base_dir):
        if nome_arquivo in files:
            caminho_candidato = os.path.join(root, nome_arquivo)
            print(f"Encontrado caminho alternativo: {caminho_candidato}")
            return caminho_candidato

    print(f"ERRO: Arquivo {nome_arquivo} não encontrado em nenhum local correspondente.")
    return None

@app.route('/api/gerar_pdf/<int:inspecao_id>', methods=['GET'])
@login_required
def gerar_pdf(inspecao_id):
    try:
        usuario_logado, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        
        # Adicionar log de acesso
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')
        
        mensagem_log = (
            f"(WEB) O Usuário {usuario_logado} - {nivel_acesso} gerou o PDF da inspeção ID {inspecao_id} em {data_formatada}."
        )

        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario_logado,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Inspeção - Geração PDF",
            data_geracao=data_formatada
        )
        session.add(log_entry)
        session.commit()

        if not inspecao_id or inspecao_id <= 0:
            return jsonify({'error': 'ID inválido'}), 400

        config = load_config()
        max_retries = config.get('MAX_RETRIES', 3)
        base_dir = os.path.dirname(os.path.abspath(__file__))

        last_error = None
        for attempt in range(max_retries):
            try:
                query = (
                    session.query(Registro, Equipamento)
                    .join(Equipamento, Registro.equipamento_id == Equipamento.id)
                    .filter(Registro.id == inspecao_id)
                    .first()
                )

                if not query:
                    return jsonify({'error': 'Inspeção não encontrada'}), 404

                registro, equipamento = query

                # Buscar fotos
                fotos_query = session.query(FotoRegistro.caminho).filter(
                    FotoRegistro.registro_id == inspecao_id
                ).all()

                caminhos_fotos = []
                for foto in fotos_query:
                    caminho_foto = encontrar_caminho_real(foto[0], base_dir)
                    if caminho_foto:
                        caminhos_fotos.append(caminho_foto)

                checklist_items = session.query(ChecklistItem).filter(
                    ChecklistItem.registro_id == inspecao_id
                ).all()

                def format_date(date):
                    return date.strftime('%d/%m/%Y') if date else "N/A"

                dados_fixos = {
                    "{NUMERO_INSPECAO}": str(registro.numero_inspecao) if registro.numero_inspecao else "N/A",
                    "{RESPONSAVEL}": str(registro.responsavel) if registro.responsavel else "N/A",
                    "{DATA_REGISTRO}": format_date(registro.data_registro),
                    "{LOCALIZACAO}": str(registro.localizacao) if registro.localizacao else "N/A",
                    "{STATUS_INSPECAO}": str(registro.status_inspecao) if registro.status_inspecao else "N/A",
                    "{MOTIVO}": str(registro.motivo_acao) if registro.motivo_acao else "N/A",            
                    "{TIPO}": str(equipamento.tipo) if equipamento.tipo else "N/A",
                    "{TAG_EQUIPAMENTO}": str(equipamento.tag_equipamento) if equipamento.tag_equipamento else "N/A",
                    "{STATUS_EQUIPAMENTO}": str(registro.status_equipamento) if registro.status_equipamento else "N/A",
                    "{CLASSE}": str(getattr(equipamento, 'classe', 'N/A')),
                    "{DATA_FABRICACAO}": format_date(getattr(equipamento, 'data_fabricacao', None)),
                    "{DATA_VALIDADE}": format_date(getattr(equipamento, 'data_validade', None)),
                    "{PESO}": str(getattr(equipamento, 'peso', 'N/A')),
                    "{data_ultima_inspecao}": format_date(getattr(equipamento, 'data_ultima_inspecao', None)),
                    "{data_proxima_inspecao}": format_date(getattr(equipamento, 'data_proxima_inspecao', None)),
                    "{DATA_INSPECAO}": format_date(registro.data_registro),
                    "{CRITICIDADE}": str(registro.criticidade_inspecao) if registro.criticidade_inspecao else "N/A",               
                    "{DATA_FINALIZACAO}": format_date(registro.data_encerramento_inspecao) if registro.data_encerramento_inspecao else "N/A",                    
                    "{OBSERVACOES}": str(registro.observacoes) if registro.observacoes else "N/A",
                    "{FOTOS}": "(Nenhuma foto anexada)" if not caminhos_fotos else ""
                }

                output_folder = os.path.join(
                    base_dir,
                    config['PDF_OUTPUT_FOLDER'],
                    f"{registro.numero_inspecao}-{registro.data_inspecao.strftime('%d-%m-%Y')}"
                )
                os.makedirs(output_folder, exist_ok=True)

                caminho_modelo = os.path.join(base_dir, config['PDF_TEMPLATE_PATH'])
                filename = f"{config['PDF_FILENAME_PREFIX']}{registro.numero_inspecao}"
                caminho_docx = os.path.join(output_folder, f"{filename}.docx")
                caminho_pdf = caminho_docx.replace(".docx", ".pdf")

                if not substituir_ids_e_gerar_docx(
                    caminho_modelo, caminho_docx, dados_fixos, checklist_items, caminhos_fotos
                ):
                    raise Exception("Erro ao gerar DOCX.")

                if not converter_docx_para_pdf(caminho_docx, caminho_pdf):
                    raise Exception("Erro na conversão para PDF.")

                return send_file(
                    caminho_pdf,
                    as_attachment=True,
                    mimetype='application/pdf',
                    download_name=f"{filename}.pdf"
                )

            except Exception as e:
                last_error = e
                print(f"Tentativa {attempt + 1} falhou: {str(e)}")
                if attempt < max_retries - 1:
                    continue
                raise

    except Exception as e:
        return jsonify({
            'error': 'Falha ao gerar relatório',
            'details': str(e),
            'suggestion': 'Tente novamente ou contate o suporte.'
        }), 500

@app.route('/inspecoes/inspecoes_cadastrados')
@login_required
@permission_required('inspecoes_cadastrados')
def inspecoes_cadastrados():
    current_user = get_current_user(flask_session, session)
    allowed_menus = get_allowed_menus_for_user(current_user)    
    return render_template('inspecoes/inspecao_registrados.html', current_user=current_user, allowed_menus=allowed_menus)

@app.route('/inspecoes/mapa_inspecoes')
@login_required
@permission_required('mapas_inspecoes')
def mapa_inspecoes():
    current_user = get_current_user(flask_session, session)
    allowed_menus = get_allowed_menus_for_user(current_user)    
    return render_template('inspecoes/mapas_inspecoes.html', current_user=current_user, allowed_menus=allowed_menus)

@app.route('/api/inspecoes', methods=['GET'])
@login_required
def get_inspecoes():
    try:
        # Parâmetros do DataTables
        draw = request.args.get('draw', type=int)
        start = request.args.get('start', type=int)
        length = request.args.get('length', type=int)
        search_value = request.args.get('search[value]', '').strip()
        
        # Filtros adicionais
        status_filter = request.args.get('status', '').strip()
        tipo_equipamento_filter = request.args.get('tipo_equipamento', '').strip()
        data_inicio_filter = request.args.get('data_inicio', '').strip()
        data_fim_filter = request.args.get('data_fim', '').strip()
        responsavel_filter = request.args.get('responsavel', '').strip()
        localizacao_filter = request.args.get('localizacao', '').strip()
        numero_inspecao_filter = request.args.get('numero_inspecao', '').strip()
        tag_equipamento_filter = request.args.get('tag_equipamento', '').strip()
        
        # Ordenação
        order_column_index = request.args.get('order[0][column]', default=0, type=int)
        order_direction = request.args.get('order[0][dir]', 'asc')
        order_column_name = request.args.get(f'columns[{order_column_index}][data]', 'numero_inspecao')
        
        # Query base
        query = (
            session.query(Registro, Equipamento)
            .join(Equipamento, Registro.equipamento_id == Equipamento.id)
        )

        # Aplicar filtro de pesquisa global
        if search_value:
            query = query.filter(
                or_(
                    Equipamento.tag_equipamento.ilike(f'%{search_value}%'),
                    Registro.localizacao.ilike(f'%{search_value}%'),
                    Registro.responsavel.ilike(f'%{search_value}%'),
                    Registro.numero_inspecao.ilike(f'%{search_value}%'),
                    cast(Registro.data_registro, String).ilike(f'%{search_value}%'),
                    cast(Registro.data_validade_inspecao, String).ilike(f'%{search_value}%'),
                    Equipamento.tipo.ilike(f'%{search_value}%')
                )
            )

        # Aplicar filtros individuais
        if status_filter and status_filter != 'all':
            query = query.filter(Registro.status_inspecao == status_filter)
            
        if tipo_equipamento_filter and tipo_equipamento_filter != 'all':
            query = query.filter(Equipamento.tipo == tipo_equipamento_filter)
            
        if responsavel_filter:
            query = query.filter(Registro.responsavel.ilike(f'%{responsavel_filter}%'))
            
        if localizacao_filter:
            query = query.filter(Registro.localizacao.ilike(f'%{localizacao_filter}%'))
            
        if numero_inspecao_filter:
            query = query.filter(Registro.numero_inspecao.ilike(f'%{numero_inspecao_filter}%'))
            
        if tag_equipamento_filter:
            query = query.filter(Equipamento.tag_equipamento.ilike(f'%{tag_equipamento_filter}%'))

        if data_inicio_filter and data_fim_filter:
            try:
                data_inicio = datetime.strptime(data_inicio_filter, '%d-%m-%Y').date()
                data_fim = datetime.strptime(data_fim_filter, '%d-%m-%Y').date()
                query = query.filter(Registro.data_registro.between(data_inicio, data_fim))
            except ValueError:
                return jsonify({'error': 'Formato de data inválido. Use DD-MM-AAAA.'}), 400

        # Contagem total sem filtros
        total_records = session.query(func.count(Registro.id)).scalar()

        # Contagem com filtros aplicados
        records_filtered = query.count()
        
        # Aplicar ordenação
        if order_column_name in ['tag_equipamento', 'tipo_equipamento']:
            order_expr = getattr(Equipamento, order_column_name.replace('tipo_equipamento', 'tipo'))
        else:
            order_expr = getattr(Registro, order_column_name)
            
        if order_direction == 'desc':
            order_expr = order_expr.desc()
            
        query = query.order_by(order_expr)

        # Paginação
        resultados = query.offset(start).limit(length).all()

        # Preparar dados para resposta
        inspecoes_data = [
            {
                'id': registro.id,
                'numero_inspecao': registro.numero_inspecao,
                'equipamento_id': registro.equipamento_id,
                'equipamento': f"{equipamento.tag_equipamento} - {registro.localizacao}",
                'tag_equipamento': equipamento.tag_equipamento,
                'tipo_equipamento': equipamento.tipo,
                'localizacao': registro.localizacao,
                'responsavel': registro.responsavel,
                'criticidade_inspecao': registro.criticidade_inspecao,                
                'data_registro': registro.data_registro.strftime('%d-%m-%Y'),
                'data_validade_inspecao': registro.data_validade_inspecao.strftime('%d-%m-%Y'),
                'data_encerramento_inspecao': registro.data_encerramento_inspecao.strftime('%d-%m-%Y') if registro.data_encerramento_inspecao else None,               
                'status_inspecao': registro.status_inspecao,
                'status_equipamento': registro.status_equipamento,
                'observacoes': registro.observacoes,
                'motivo_acao': registro.motivo_acao,
            }
            for registro, equipamento in resultados
        ]

        return jsonify({
            'draw': draw,
            'recordsTotal': total_records,
            'recordsFiltered': records_filtered,
            'data': inspecoes_data
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/inspecoes/vencidas_mensal', methods=['GET'])
@login_required
def contagem_inspecoes_vencidas_mensal():
    try:
        hoje = datetime.now().date().isoformat()  # exemplo: '2025-04-13'

        # Contar todos os registros vencidos
        total_vencidas = session.query(func.count(Registro.id)).filter(
            Registro.status_inspecao == 'Pendente',
            func.date(Registro.data_validade_inspecao) < func.date(hoje)
        ).scalar()

        return jsonify({
            'status': 'Vencida',
            'total': total_vencidas
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/inspecoes/mapa', methods=['GET'])
@login_required
def get_inspecoes_mapa():
    try:
        # Parâmetros de filtro
        status_filter = request.args.get('status', '').strip()
        tipo_equipamento_filter = request.args.get('tipo_equipamento', '').strip()
        data_inicio_filter = request.args.get('data_inicio', '').strip()
        data_fim_filter = request.args.get('data_fim', '').strip()
        responsavel_filter = request.args.get('responsavel', '').strip()
        localizacao_filter = request.args.get('localizacao', '').strip()
        numero_inspecao_filter = request.args.get('numero_inspecao', '').strip()
        tag_equipamento_filter = request.args.get('tag_equipamento', '').strip()

        # Alias para pegar a última localização por registro
        ultima_localizacao = (
            session.query(LocalizacaoUsuario)
            .filter(LocalizacaoUsuario.latitude.isnot(None), LocalizacaoUsuario.longitude.isnot(None))
            .order_by(desc(LocalizacaoUsuario.data_registro))
            .distinct(LocalizacaoUsuario.registro_id)
            .subquery()
        )

        LocalizacaoAlias = aliased(LocalizacaoUsuario, ultima_localizacao)

        # Query base com joins
        query = (
            session.query(Registro, Equipamento, LocalizacaoAlias)
            .join(Equipamento, Registro.equipamento_id == Equipamento.id)
            .join(LocalizacaoAlias, LocalizacaoAlias.registro_id == Registro.id)
        )

        # Aplicar filtros
        if status_filter and status_filter != 'all':
            query = query.filter(Registro.status_inspecao == status_filter)

        if tipo_equipamento_filter and tipo_equipamento_filter != 'all':
            query = query.filter(Equipamento.tipo == tipo_equipamento_filter)

        if responsavel_filter:
            query = query.filter(Registro.responsavel.ilike(f'%{responsavel_filter}%'))

        if localizacao_filter:
            query = query.filter(Registro.localizacao.ilike(f'%{localizacao_filter}%'))

        if numero_inspecao_filter:
            query = query.filter(Registro.numero_inspecao.ilike(f'%{numero_inspecao_filter}%'))

        if tag_equipamento_filter:
            query = query.filter(Equipamento.tag_equipamento.ilike(f'%{tag_equipamento_filter}%'))

        if data_inicio_filter and data_fim_filter:
            try:
                data_inicio = datetime.strptime(data_inicio_filter, '%d-%m-%Y').date()
                data_fim = datetime.strptime(data_fim_filter, '%d-%m-%Y').date()
                query = query.filter(Registro.data_registro.between(data_inicio, data_fim))
            except ValueError:
                return jsonify({'error': 'Formato de data inválido. Use DD-MM-AAAA.'}), 400

        resultados = query.all()

        # Preparar dados para resposta
        inspecoes_data = [
            {
                'id': registro.id,
                'numero_inspecao': registro.numero_inspecao,
                'tag_equipamento': equipamento.tag_equipamento,
                'tipo_equipamento': equipamento.tipo,
                'localizacao': registro.localizacao,
                'responsavel': registro.responsavel,
                'criticidade_inspecao': registro.criticidade_inspecao,                  
                'data_registro': registro.data_registro.strftime('%d-%m-%Y'),
                'data_validade_inspecao': registro.data_validade_inspecao.strftime('%d-%m-%Y') if registro.data_validade_inspecao else None,
                'data_validade_inspecao': registro.data_validade_inspecao.strftime('%d-%m-%Y') if registro.data_validade_inspecao else None,                
                'status_inspecao': registro.status_inspecao,
                'latitude': localizacao.latitude,
                'longitude': localizacao.longitude
            }
            for registro, equipamento, localizacao in resultados
        ]

        return jsonify(inspecoes_data)

    except Exception as e:
        return jsonify({'error': str(e)}), 500
        
@app.route('/api/inspecoes/<int:inspecao_id>', methods=['GET'])
@login_required
def get_inspecao_detalhes(inspecao_id):
    try:
        resultado = (
            session.query(Registro, Equipamento)
            .join(Equipamento, Registro.equipamento_id == Equipamento.id)
            .filter(Registro.id == inspecao_id)
            .first()
        ) 

        if not resultado:
            return jsonify({'error': 'Inspeção não encontrada'}), 404

        registro, equipamento = resultado

        inspecao_data = {
            'id': registro.id,
            'numero_inspecao': registro.numero_inspecao,
            'equipamento_id': registro.equipamento_id,
            'equipamento': f"{equipamento.tag_equipamento} - {registro.tipo} - {registro.localizacao}",
            'tag_equipamento': equipamento.tag_equipamento,
            'tipo': equipamento.tipo,
            'localizacao': equipamento.localizacao,
            'responsavel': registro.responsavel,
            'criticidade_inspecao': registro.criticidade_inspecao,              
            'status_inspecao': registro.status_inspecao,
            'motivo_acao': registro.motivo_acao,
            'data_validade_inspecao': registro.data_validade_inspecao.strftime('%d-%m-%Y'),           
            'data_registro': registro.data_registro.strftime('%d-%m-%Y'),
            'data_encerramento_inspecao': registro.data_encerramento_inspecao.strftime('%d-%m-%Y')if registro.data_encerramento_inspecao else None,   

            'status_equipamento': registro.status_equipamento,
            'observacoes': registro.observacoes
        }

        return jsonify(inspecao_data)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/inspecoes/<int:inspecao_id>/localizacao', methods=['GET'])
@login_required
def get_inspecao_localizacao(inspecao_id):
    try:
        localizacao = session.query(LocalizacaoUsuario).filter(
            LocalizacaoUsuario.registro_id == inspecao_id
        ).order_by(LocalizacaoUsuario.data_registro.desc()).first()

        if not localizacao:
            return jsonify({'error': 'Localização não encontrada para esta inspeção'}), 404

        localizacao_data = {
            'latitude': localizacao.latitude,
            'longitude': localizacao.longitude,
            'data_registro': localizacao.data_registro.strftime('%d-%m-%Y %H:%M:%S'),
            'usuario': localizacao.usuario,
            'metadados': localizacao.metadados
        }

        return jsonify(localizacao_data)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/inspecoes/<int:inspecao_id>/inspecao', methods=['GET'])
@login_required
def get_checklist_inspecao(inspecao_id):
    try:
        inspecao = session.query(ChecklistItem).filter_by(registro_id=inspecao_id).all()
        if not inspecao:
            return jsonify({'error': 'Checklist não encontrado para esta inspeção'}), 404

        inspecao_data = [
            {
                'id': item.id,
                'description': item.description,
                'checked': item.checked
            }
            for item in inspecao
        ]

        return jsonify(inspecao_data)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/inspecoes/<int:inspecao_id>', methods=['DELETE'])
@login_required
@permission_required('excluir_inspecao')
def delete_inspecao(inspecao_id):
    try:
        # Encontra a inspeção
        inspecao = session.query(Registro).filter_by(id=inspecao_id).first()
        if not inspecao:
            return jsonify({'error': 'Inspeção não encontrada'}), 404

        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

        # Primeiro deleta todos os itens de checklist associados
        session.query(ChecklistItem).filter_by(registro_id=inspecao_id).delete()
        
        # Depois deleta todas as fotos associadas
        session.query(FotoRegistro).filter_by(registro_id=inspecao_id).delete()

        mensagem_log = (
            f"(WEB) O Usuário {usuario} ({nivel_acesso}) excluiu a inspeção {inspecao.numero_inspecao} "
            f"e todos os seus itens relacionados em {data_formatada}.\n"
        )

        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Inspeção - Exclusão",
            data_geracao=data_formatada
        )

        session.add(log_entry)
        session.delete(inspecao)
        session.commit()

        return jsonify({'message': 'Inspeção e todos os seus itens relacionados foram apagados com sucesso'})

    except Exception as e:
        session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/inspecoes/<int:inspecao_id>/baixar', methods=['POST'])
@login_required
def baixar_inspecao(inspecao_id):
    try:
        data = request.get_json()
        motivo_acao = data.get('motivo_acao', '').strip()

        if not motivo_acao:
            return jsonify({'error': 'O motivo da ação é obrigatório.'}), 400

        inspecao = session.query(Registro).filter_by(id=inspecao_id).first()
        if not inspecao:
            return jsonify({'error': 'Inspeção não encontrada'}), 404

        if inspecao.status_inspecao == 'Finalizada':
            return jsonify({'error': 'Não é possível finalizar uma inspeção já finalizada.'}), 400
        if inspecao.status_inspecao == 'Cancelada':
            return jsonify({'error': 'Não é possível finalizar uma inspeção cancelada.'}), 400
        
        data_atual = datetime.now()

        inspecao.status_inspecao = 'Finalizada'
        inspecao.motivo_acao = motivo_acao
        inspecao.data_encerramento_inspecao = data_atual

        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)

        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

        mensagem_log = (
            f"(WEB) O Usuário {usuario} ({nivel_acesso}) finalizou a inspeção {inspecao.numero_inspecao} em {data_formatada}.\n"
            f"Motivo: {motivo_acao}\n"
        )

        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Inspeção - Finalização",
            data_geracao=data_formatada
        )

        session.add(log_entry)
        session.commit()

        return jsonify({'message': 'Inspeção concluída com sucesso!'})

    except Exception as e:
        session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/inspecoes/<int:inspecao_id>/cancelar', methods=['POST'])
@login_required
def cancelar_inspecao(inspecao_id):
    try:
        data = request.get_json()
        motivo_acao = data.get('motivo_acao', '').strip()

        if not motivo_acao:
            return jsonify({'error': 'O motivo da ação é obrigatório.'}), 400

        inspecao = session.query(Registro).filter_by(id=inspecao_id).first()
        if not inspecao:
            return jsonify({'error': 'Inspeção não encontrada'}), 404

        if inspecao.status_inspecao == 'Cancelada':
            return jsonify({'error': 'Não é possível cancelar uma inspeção já cancelada.'}), 400

        inspecao.status_inspecao = 'Cancelada'
        inspecao.motivo_acao = motivo_acao

        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

        mensagem_log = (
            f"(WEB) O Usuário {usuario} ({nivel_acesso}) cancelou a inspeção {inspecao.numero_inspecao} em {data_formatada}.\n"
            f"Motivo: {motivo_acao}\n"
        )

        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Inspeção - Cancelamento",
            data_geracao=data_formatada
        )

        session.add(log_entry)
        session.commit()

        return jsonify({'message': 'Inspeção cancelada com sucesso!'})

    except Exception as e:
        session.rollback()
        return jsonify({'error': str(e)}), 500

def ensure_config_loaded():
    if 'UPLOAD_FOLDER' not in app.config:
        update_flask_config(load_config())

@app.route('/api/inspecoes/<int:inspecao_id>/fotos', methods=['GET'])
@login_required
def get_fotos(inspecao_id):
    try:
        ensure_config_loaded()

        upload_folder = os.path.abspath(app.config['UPLOAD_FOLDER'])

        fotos = session.query(FotoRegistro).filter(FotoRegistro.registro_id == inspecao_id).all()
        if not fotos:
            return jsonify([]), 200

        fotos_urls = []
        for foto in fotos:
            rel_path = os.path.relpath(foto.caminho, upload_folder)
            web_path = rel_path.replace(os.sep, '/')

            fotos_urls.append({
                "url": f"/uploads/{web_path}",
                "path": web_path,
                "nome": os.path.basename(foto.caminho)
            })

        return jsonify(fotos_urls)
    except Exception as e:
        app.logger.error(f'Erro ao buscar anexos da inspeção {inspecao_id}: {str(e)}')
        return jsonify({'erro': 'Erro ao buscar anexos'}), 500

@app.route('/uploads/<path:filename>')
@login_required
def serve_uploaded_file(filename):
    try:
        ensure_config_loaded()

        upload_folder = app.config['UPLOAD_FOLDER']
        file_path = os.path.join(upload_folder, filename)

        if not os.path.exists(file_path):
            return jsonify({"error": "Arquivo não encontrado"}), 404

        return send_from_directory(upload_folder, filename)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Métodos de Equipamento ###################################################################

@app.route('/api/selecao_equipamentos', methods=['GET'])
@login_required
def selecao_equipamentos():
    try:
        termo_pesquisa = request.args.get('search', '').strip()
        query = (
            session.query(Equipamento)
            .filter(
                or_(
                    Equipamento.localizacao.ilike(f'%{termo_pesquisa}%'),
                    Equipamento.tipo.ilike(f'%{termo_pesquisa}%'),
                    Equipamento.tag_equipamento.ilike(f'%{termo_pesquisa}%'),
                    Equipamento.codigo_barras.ilike(f'%{termo_pesquisa}%')
                )
            )
            .filter(Equipamento.status == 'Ativo')  # Filtro adicional para status "Ativo"
        )
        equipamentos = query.all()

        equipamentos_data = [
            {
                'id': equipamento.id,
                'tag_equipamento': equipamento.tag_equipamento,
                'localizacao': equipamento.localizacao,
                'tipo': equipamento.tipo,
                'classe': equipamento.classe,
                'data_fabricacao': equipamento.data_fabricacao.strftime('%d-%m-%Y'),
                'peso': equipamento.peso,
                'teste_hidrostatico_n2': equipamento.teste_hidrostatico_n2.strftime('%d-%m-%Y') if equipamento.teste_hidrostatico_n2 else None,
                'teste_hidrostatico_n3': equipamento.teste_hidrostatico_n3.strftime('%d-%m-%Y') if equipamento.teste_hidrostatico_n3 else None,
                'data_validade': equipamento.data_validade.strftime('%d-%m-%Y'),
                'status': equipamento.status,
                'codigo_barras': equipamento.codigo_barras,
                'data_ultima_inspecao': equipamento.data_ultima_inspecao.strftime('%d-%m-%Y') if equipamento.data_ultima_inspecao else None,
                'data_proxima_inspecao': equipamento.data_proxima_inspecao.strftime('%d-%m-%Y') if equipamento.data_proxima_inspecao else None
            }
            for equipamento in equipamentos
        ]

        return jsonify(equipamentos_data)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/equipamentos/equipamentos_cadastrados')
@login_required
@permission_required('equipamentos_cadastrados')
def equipamentos_cadastrados():
    current_user = get_current_user(flask_session, session)
    allowed_menus = get_allowed_menus_for_user(current_user)    
    return render_template('equipamento/equipamentos_cadastrados.html', current_user=current_user, allowed_menus=allowed_menus)
 
@app.route('/api/equipamentos_cadastrados', methods=['GET'])
@login_required 
def get_equipamentos_cadastrados():
    try:
        # Parâmetros do DataTables
        draw = request.args.get('draw', type=int)
        start = request.args.get('start', type=int)
        length = request.args.get('length', type=int)
        search_value = request.args.get('search[value]', '').strip()
        
        # Filtros adicionais
        status_filter = request.args.get('status', '').strip()
        tipo_filter = request.args.get('tipo', '').strip()
        localizacao_filter = request.args.get('localizacao', '').strip()
        data_inicio_filter = request.args.get('data_inicio', '').strip()
        data_fim_filter = request.args.get('data_fim', '').strip()

        # Ordenação
        order_column_index = request.args.get('order[0][column]', default=0, type=int)
        order_direction = request.args.get('order[0][dir]', 'asc')
        
        # Mapeamento de colunas
        columns = [
            'tag_equipamento',
            'localizacao',
            'tipo',
            'status',
            'id'  # Coluna de ações
        ]
        order_column = columns[order_column_index] if order_column_index < len(columns) else 'id'

        # Construção da query
        query = session.query(Equipamento)

        # Aplicar filtros
        if search_value:
            query = query.filter(
                or_(
                    Equipamento.tag_equipamento.ilike(f'%{search_value}%'),
                    Equipamento.localizacao.ilike(f'%{search_value}%'),
                    Equipamento.tipo.ilike(f'%{search_value}%'),
                    Equipamento.codigo_barras.ilike(f'%{search_value}%')
                )
            )

        if status_filter:
            if status_filter not in ['Ativo', 'Manutenção', 'Descartado', 'Reserva']:
                return jsonify({'error': 'Status de equipamento inválido.'}), 400
            query = query.filter(Equipamento.status == status_filter)
            
        if tipo_filter:
            query = query.filter(Equipamento.tipo.ilike(f'%{tipo_filter}%'))
            
        if localizacao_filter:
            query = query.filter(Equipamento.localizacao.ilike(f'%{localizacao_filter}%'))
            
        if data_inicio_filter and data_fim_filter:
            try:
                data_inicio = datetime.strptime(data_inicio_filter, '%Y-%m-%d').date()
                data_fim = datetime.strptime(data_fim_filter, '%Y-%m-%d').date()
                query = query.filter(Equipamento.data_fabricacao.between(data_inicio, data_fim))
            except ValueError:
                return jsonify({'error': 'Formato de data inválido. Use YYYY-MM-DD.'}), 400

        # Ordenação
        if order_direction == 'desc':
            query = query.order_by(desc(getattr(Equipamento, order_column)))
        else:
            query = query.order_by(getattr(Equipamento, order_column))

        # Paginação
        total_records = query.count()
        equipamentos = query.offset(start).limit(length).all()

        # Formatar dados
        equipamentos_data = [
            {
                'tag_equipamento': equipamento.tag_equipamento,
                'localizacao': equipamento.localizacao,
                'tipo': equipamento.tipo,
                'criticidade': equipamento.criticidade,                
                'status': equipamento.status,
                'id': equipamento.id,
                'DT_RowId': f'equipamento_{equipamento.id}'  # Para referência no frontend
            }
            for equipamento in equipamentos
        ]

        return jsonify({
            'draw': draw,
            'recordsTotal': total_records,
            'recordsFiltered': total_records,
            'data': equipamentos_data
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/equipamentos/contagem_mensal', methods=['GET'])
@login_required
def contagem_equipamentos_mensal():
    try:
        hoje = datetime.now()
        mes_atual = hoje.month
        ano_atual = hoje.year
        mes_anterior = (hoje.replace(day=1) - timedelta(days=1)).month
        ano_anterior = hoje.year if mes_anterior != 12 else hoje.year - 1

        def contagem(mes, ano):
            return session.query(func.count(Equipamento.id)).filter(
                extract('month', Equipamento.criado_em) == mes,
                extract('year', Equipamento.criado_em) == ano
            ).scalar()

        atual = contagem(mes_atual, ano_atual)
        anterior = contagem(mes_anterior, ano_anterior)
        variacao = round(((atual - anterior) / anterior) * 100, 2) if anterior else 100.0

        return jsonify({
            'atual': atual,
            'anterior': anterior,
            'variacao': variacao
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/equipamentos/ativos', methods=['GET'])
@login_required
def contagem_equipamentos_ativos_mensal():
    try:
        # Contar todos os equipamentos ativos, sem considerar a data
        def contagem_ativos():
            return session.query(func.count(Equipamento.id)).filter(
                Equipamento.status == 'Ativo'  # Apenas equipamentos ativos
            ).scalar()

        # Contagem total de equipamentos ativos
        total_ativos = contagem_ativos()

        # Para variação, podemos calcular com base no total de equipamentos ativos do mês passado
        # ou simplesmente retornar a contagem total sem comparar com o mês anterior, se preferir
        anterior = 0  # Para exemplo, assumimos que no mês anterior a contagem era 0
        variacao = round(((total_ativos - anterior) / anterior) * 100, 2) if anterior else 100.0

        return jsonify({
            'total_ativos': total_ativos,
            'variacao': variacao
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/equipamentos/vencidos_mensal', methods=['GET'])
@login_required
def contagem_equipamentos_vencidos_mensal():
    try:
        hoje = datetime.now()
        primeiro_dia_atual = hoje.replace(day=1)
        primeiro_dia_anterior = primeiro_dia_atual - timedelta(days=1)
        primeiro_dia_anterior = primeiro_dia_anterior.replace(day=1)

        def contagem(inicio, fim):
            return session.query(func.count(Equipamento.id)).filter(
                Equipamento.data_validade < fim,
                Equipamento.data_validade >= inicio,
                Equipamento.status != 'Descartado'
            ).scalar()

        atual = contagem(primeiro_dia_atual, hoje)
        anterior = contagem(primeiro_dia_anterior, primeiro_dia_atual)
        variacao = round(((atual - anterior) / anterior) * 100, 2) if anterior else 100.0

        return jsonify({
            'atual': atual,
            'anterior': anterior,
            'variacao': variacao
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500
           
@app.route('/api/equipamentos/<int:equipamento_id>', methods=['GET'])
@login_required
def get_equipamento_detalhes(equipamento_id):
    try:
        equipamento = session.query(Equipamento).filter_by(id=equipamento_id).first()
        if not equipamento:
            return jsonify({'error': 'Equipamento não encontrado'}), 404

        equipamento_data = {
            'id': equipamento.id,
            'tag_equipamento': equipamento.tag_equipamento,
            'localizacao': equipamento.localizacao,
            'tipo': equipamento.tipo,
            'criticidade': equipamento.criticidade,            
            'classe': equipamento.classe,
            'fabricante': equipamento.fabricante,
            'modelo': equipamento.modelo,
            'numero_serie': equipamento.numero_serie,
            'data_fabricacao': equipamento.data_fabricacao.strftime('%d-%m-%Y') if equipamento.data_fabricacao else None,
            'peso': equipamento.peso,
            'teste_hidrostatico_n2': equipamento.teste_hidrostatico_n2.strftime('%d-%m-%Y') if equipamento.teste_hidrostatico_n2 else None,
            'teste_hidrostatico_n3': equipamento.teste_hidrostatico_n3.strftime('%d-%m-%Y') if equipamento.teste_hidrostatico_n3 else None,
            'data_validade': equipamento.data_validade.strftime('%d-%m-%Y') if equipamento.data_validade else None,
            'status': equipamento.status,
            'codigo_barras': equipamento.codigo_barras,
            'patrimonio': equipamento.patrimonio,
            'foto': equipamento.foto,
            'observacoes': equipamento.observacoes,
            'data_ultima_inspecao': equipamento.data_ultima_inspecao.strftime('%d-%m-%Y') if equipamento.data_ultima_inspecao else None,
            'data_proxima_inspecao': equipamento.data_proxima_inspecao.strftime('%d-%m-%Y') if equipamento.data_proxima_inspecao else None,
            'usuario_geracao': equipamento.usuario_geracao,
            'usuario_atualizacao': equipamento.usuario_atualizacao,
            'criado_em': equipamento.criado_em.strftime('%d-%m-%Y %H:%M:%S') if equipamento.criado_em else None,
            'atualizado_em': equipamento.atualizado_em.strftime('%d-%m-%Y %H:%M:%S') if equipamento.atualizado_em else None
        }

        return jsonify(equipamento_data)

    except Exception as e:
        app.logger.error(f"Erro ao buscar equipamento: {str(e)}", exc_info=True)
        return jsonify({'error': 'Erro interno ao buscar equipamento'}), 500
  
@app.route('/api/equipamentos/<equipamentoId>/inspecoes', methods=['GET'])
@login_required
def get_historico_inspecoes_usuario(equipamentoId):
    equipamento = session.query(Equipamento).filter_by(id=equipamentoId).first()
    if not equipamento:
        return jsonify({'error': 'Equipamento não encontrado'}), 404

    inspecoes = session.query(Registro).filter_by(equipamento_id=equipamento.id).order_by(Registro.data_inspecao.desc()).all()

    inspecoes_data = [
        {
            'id': inspecao.id,
            'numero_inspecao': inspecao.numero_inspecao,            
            'equipamento': inspecao.equipamento_id,
            'tipo': inspecao.tipo,
            'classe': inspecao.classe,
            'localizacao': inspecao.localizacao,
            'status': inspecao.status_equipamento, 
            'responsavel': inspecao.responsavel,
            'data_inspecao': inspecao.data_inspecao.strftime('%d/%m/%Y') if inspecao.data_inspecao else '',
            'validade': inspecao.data_validade_inspecao.strftime('%d/%m/%Y') if inspecao.data_validade_inspecao else '',
            'observacoes': inspecao.observacoes,
        }
        for inspecao in inspecoes
    ]

    return jsonify(inspecoes_data)

@app.route('/api/equipamentos/<int:equipamento_id>', methods=['DELETE'])
@login_required
@permission_required('excluir_equipamento')
def delete_equipamento(equipamento_id):
    try:
        equipamento = session.query(Equipamento).filter_by(id=equipamento_id).first()
        if not equipamento:
            return jsonify({'error': 'Equipamento não encontrado'}), 404

        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

        mensagem_log = (
            f"(WEB) O Usuário {usuario} ({nivel_acesso}) excluiu o equipamento {equipamento.tag_equipamento} em {data_formatada}.\n"
        )

        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Equipamento - Exclusão",
            data_geracao=data_formatada
        )

        session.add(log_entry)
        session.delete(equipamento)
        session.commit()

        return jsonify({'message': 'Equipamento excluído com sucesso'})

    except Exception as e:
        session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/equipamentos/<int:equipamento_id>/manutencao', methods=['POST'])
@login_required
def manutencao_equipamento(equipamento_id):
    try:
        equipamento = session.query(Equipamento).filter_by(id=equipamento_id).first()
        if not equipamento:
            return jsonify({'error': 'Equipamento não encontrado'}), 404

        if equipamento.status == 'Manutenção':
            return jsonify({'error': 'O equipamento já está em manutenção.'}), 400

        equipamento.status = 'Manutenção'

        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

        mensagem_log = (
            f"(WEB) O Usuário {usuario} ({nivel_acesso}) alterou o status do equipamento {equipamento.tag_equipamento} para 'Manutenção' em {data_formatada}.\n"
        )
 
        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Equipamento - Em Manutenção",
            data_geracao=data_formatada
        )

        session.add(log_entry)
        session.commit()

        return jsonify({'message': 'Status do equipamento alterado para "Manutenção".'})

    except Exception as e:
        session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/equipamentos/<int:equipamento_id>/ativo', methods=['POST'])
@login_required
def ativo_equipamento(equipamento_id):
    try:
        equipamento = session.query(Equipamento).filter_by(id=equipamento_id).first()
        if not equipamento:
            return jsonify({'error': 'Equipamento não encontrado'}), 404

        if equipamento.status == 'Ativo':
            return jsonify({'error': 'O equipamento já está ativo.'}), 400

        equipamento.status = 'Ativo'

        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

        mensagem_log = (
            f"(WEB) O Usuário {usuario} ({nivel_acesso}) alterou o status do equipamento {equipamento.tag_equipamento} para 'Ativo' em {data_formatada}.\n"
        )

        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Equipamento - Ativo",
            data_geracao=data_formatada
        )

        session.add(log_entry)
        session.commit()

        return jsonify({'message': 'Status do equipamento alterado para "Ativo".'})

    except Exception as e:
        session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/equipamentos/<int:equipamento_id>/descarte', methods=['POST'])
@login_required
def descarte_equipamento(equipamento_id):
    try:
        equipamento = session.query(Equipamento).filter_by(id=equipamento_id).first()
        if not equipamento:
            return jsonify({'error': 'Equipamento não encontrado'}), 404

        if equipamento.status == 'Descartado':
            return jsonify({'error': 'O equipamento já está descartado.'}), 400

        equipamento.status = 'Descartado'

        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

        mensagem_log = (
            f"(WEB) O Usuário {usuario} ({nivel_acesso}) alterou o status do equipamento {equipamento.tag_equipamento} para 'Descartado' em {data_formatada}.\n"
        )

        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Equipamento - Descartado",
            data_geracao=data_formatada
        )

        session.add(log_entry)
        session.commit()

        return jsonify({'message': 'Status do equipamento alterado para "Descartado".'})

    except Exception as e:
        session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/equipamentos', methods=['POST'])
@login_required
def registra_novo_equipamento():
    try:
        # Inicializar dados
        data = {}

        if request.content_type.startswith('multipart/form-data'):
            data = request.form.to_dict()
        else:
            data = request.get_json() or {}

        # Processar foto, se enviada
        if 'foto' in request.files:
            file = request.files['foto']

            tag_equipamento = secure_filename(data.get('tag_equipamento', 'default'))

            upload_folder = os.path.join('static', 'images', 'Fotos_Equipamento', tag_equipamento)
            os.makedirs(upload_folder, exist_ok=True)

            ext = os.path.splitext(secure_filename(file.filename))[1].lower()
            nome_base = f"{tag_equipamento}{ext}"
            file_path = os.path.join(upload_folder, nome_base)
            file.save(file_path)

            # Caminho relativo para salvar no banco
            data['foto'] = os.path.join('images', 'Fotos_Equipamento', tag_equipamento, nome_base).replace('\\', '/')

        # Validar campos obrigatórios
        campos_obrigatorios = ['localizacao', 'tipo', 'data_fabricacao', 'codigo_barras', 'data_validade', 'status']
        for campo in campos_obrigatorios:
            if campo not in data or not data[campo]:
                return jsonify({"message": f"Campo obrigatório '{campo}' não foi fornecido."}), 400

        # Verificar duplicidade
        if session.query(Equipamento).filter_by(codigo_barras=data['codigo_barras']).first():
            return jsonify({"message": "Código de barras já cadastrado."}), 400

        if 'tag_equipamento' in data and session.query(Equipamento).filter_by(tag_equipamento=data['tag_equipamento']).first():
            return jsonify({"message": "TAG do equipamento já existe."}), 400

        # Validar datas
        try:
            data_fabricacao = datetime.strptime(data['data_fabricacao'], "%Y-%m-%d").date()
            data_validade = datetime.strptime(data['data_validade'], "%Y-%m-%d").date()
        except ValueError:
            return jsonify({"message": "Formato de data inválido. Use o formato 'YYYY-MM-DD'."}), 400

        if data_validade < data_fabricacao:
            return jsonify({"message": "A data de validade não pode ser anterior à data de fabricação."}), 400

        # Validar status
        status_permitidos = ['Ativo', 'Manutenção', 'Descartado', 'Reserva']
        if data['status'] not in status_permitidos:
            return jsonify({"message": f"Status inválido. Os status permitidos são: {', '.join(status_permitidos)}."}), 400

        # Converter datas de teste hidrostático (opcionais)
        def converter_data_mm_aaaa_para_iso(data_str):
            if data_str and len(data_str) == 7 and data_str[2] == '/':
                mes, ano = data_str.split('/')
                return datetime.strptime(f"{ano}-{mes}-01", "%Y-%m-%d").date()
            return None

        teste_hidrostatico_n2 = converter_data_mm_aaaa_para_iso(data.get('teste_hidrostatico_n2'))
        teste_hidrostatico_n3 = converter_data_mm_aaaa_para_iso(data.get('teste_hidrostatico_n3'))

        # Obter usuário atual
        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)

        equipamento = Equipamento(
            tag_equipamento=data.get('tag_equipamento'),
            criticidade=data.get('criticidade'),            
            localizacao=data['localizacao'],
            tipo=data['tipo'],
            classe=data.get('classe', 'N/A'),
            fabricante=data.get('fabricante'),
            modelo=data.get('modelo'),
            numero_serie=data.get('numero_serie'),
            data_fabricacao=data_fabricacao,
            peso=data['peso'],
            teste_hidrostatico_n2=teste_hidrostatico_n2,
            teste_hidrostatico_n3=teste_hidrostatico_n3,
            data_validade=data_validade,
            status=data['status'],
            codigo_barras=data['codigo_barras'],
            patrimonio=data.get('patrimonio'),
            foto=data.get('foto'),
            observacoes=data.get('observacoes'),
            usuario_geracao=usuario,
            usuario_atualizacao=usuario,
        )

        session.add(equipamento)
        session.commit()

        data_atual = datetime.now().strftime('%d-%m-%Y %H:%M:%S')

        mensagem_log = (
            f"(WEB) O Usuário {usuario} ({nivel_acesso}) cadastrou um novo equipamento em {data_atual}.\n"
            f"Detalhes do equipamento:\n"
            f"  - Tag do Equipamento: {equipamento.tag_equipamento}\n"
            f"  - Localização: {equipamento.localizacao}\n"
            f"  - Tipo: {equipamento.tipo}\n"
            f"  - Classe: {equipamento.classe}\n"
            f"  - Data de Fabricação: {data_fabricacao.strftime('%d-%m-%Y')}\n"
            f"  - Peso: {equipamento.peso}\n"
            f"  - Criticidade: {equipamento.criticidade}\n"           
            f"  - Teste Hidrostático N2: {teste_hidrostatico_n2.strftime('%d-%m-%Y') if teste_hidrostatico_n2 else 'N/A'}\n"
            f"  - Teste Hidrostático N3: {teste_hidrostatico_n3.strftime('%d-%m-%Y') if teste_hidrostatico_n3 else 'N/A'}\n"
            f"  - Data de Validade: {data_validade.strftime('%d-%m-%Y')}\n"
            f"  - Status: {equipamento.status}\n"
            f"  - Código de Barras: {equipamento.codigo_barras}\n"
        )

        session.add(EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Cadastro de Equipamento - Novo",
            data_geracao=data_atual
        ))

        session.commit()

        return jsonify({"message": "Equipamento cadastrado com sucesso!"}), 200

    except IntegrityError as ie:
        session.rollback()
        return jsonify({"message": f"Erro de integridade ao cadastrar equipamento: {str(ie)}"}), 500

    except Exception as e:
        session.rollback()
        return jsonify({"message": f"Erro ao cadastrar equipamento: {str(e)}"}), 500
    
@app.route('/api/upload_foto_equipamento', methods=['POST'])
@login_required
def upload_foto_equipamento():
    try:

        # Verifica se foi enviado um arquivo
        if 'foto' not in request.files:
            return jsonify({"message": "Nenhum arquivo enviado."}), 400

        foto = request.files['foto']

        # Verifica se o arquivo tem um nome válido
        if foto.filename == '':
            return jsonify({"message": "Nome de arquivo inválido."}), 400

        # Obtém a TAG do equipamento
        tag_equipamento = request.form.get('tag_equipamento')
        if not tag_equipamento:
            return jsonify({"message": "Tag do equipamento não informada."}), 400

        tag_equipamento = secure_filename(tag_equipamento)

        # Cria a pasta de destino
        upload_folder = os.path.join('static', 'images', 'Fotos_Equipamento', tag_equipamento)
        os.makedirs(upload_folder, exist_ok=True)

        # Define o nome do arquivo
        ext = os.path.splitext(secure_filename(foto.filename))[1].lower()
        nome_arquivo = f"{tag_equipamento}{ext}"
        caminho_completo = os.path.join(upload_folder, nome_arquivo)
        foto.save(caminho_completo)

        # Caminho relativo para uso no frontend
        caminho_relativo = os.path.join('images', 'Fotos_Equipamento', tag_equipamento, nome_arquivo).replace('\\', '/')

        return jsonify({
            "message": "Foto enviada com sucesso!",
            "caminho_foto": caminho_relativo
        }), 200

    except Exception as e:
        return jsonify({"message": f"Erro ao enviar foto: {str(e)}"}), 500

@app.route('/equipamentos/edit_equipamento/<int:equipamento_id>', methods=['GET', 'POST'])
@login_required
@permission_required('editar_equipamento')
def edit_equipamento(equipamento_id):
    equipamento = session.query(Equipamento).filter_by(id=equipamento_id).first()
    if not equipamento:
        return jsonify({'error': 'Equipamento não encontrado'}), 404

    usuario, nivel_acesso, foto_perfil_url = get_current_user_name_and_access_level(flask_session, session)

    if request.method == 'POST':
        try:
            # Atualiza campos principais
            equipamento.fabricante = request.form.get('fabricante')
            equipamento.modelo = request.form.get('modelo')
            equipamento.numero_serie = request.form.get('numero_serie')
            equipamento.patrimonio = request.form.get('patrimonio')
            equipamento.tag_equipamento = request.form.get('tag_equipamento')
            equipamento.criticidade = request.form.get('criticidade')            
            equipamento.localizacao = request.form.get('localizacao')
            equipamento.tipo = request.form.get('tipo')
            equipamento.classe = request.form.get('classe') if equipamento.tipo == 'Extintor' else None
            equipamento.peso = request.form.get('peso') if equipamento.tipo == 'Extintor' else None
            equipamento.codigo_barras = request.form.get('codigo_barras')
            equipamento.status = request.form.get('status')
            equipamento.observacoes = request.form.get('observacoes')

            # Datas
            equipamento.data_fabricacao = datetime.strptime(request.form.get('data_fabricacao'), '%Y-%m-%d')
            equipamento.data_validade = datetime.strptime(request.form.get('data_validade'), '%Y-%m-%d')

            teste_n2 = request.form.get('teste_hidrostatico_n2')
            equipamento.teste_hidrostatico_n2 = datetime.strptime(teste_n2, '%m/%Y') if teste_n2 else None

            teste_n3 = request.form.get('teste_hidrostatico_n3')
            equipamento.teste_hidrostatico_n3 = datetime.strptime(teste_n3, '%m/%Y') if teste_n3 else None

            equipamento.atualizado_em = datetime.now()
            equipamento.usuario_atualizacao = flask_session.get('usuario')

            # Upload de foto
            if 'foto' in request.files:
                foto = request.files['foto']
                if foto and foto.filename != '':
                    tag = secure_filename(equipamento.tag_equipamento or 'sem_tag')

                    # Cria pasta de destino
                    upload_folder = os.path.join('static', 'images', 'Fotos_Equipamento', tag)
                    os.makedirs(upload_folder, exist_ok=True)

                    # Nome do arquivo: tag.ext
                    ext = os.path.splitext(secure_filename(foto.filename))[1].lower()
                    nome_arquivo = f"{tag}{ext}"
                    caminho_completo = os.path.join(upload_folder, nome_arquivo)
                    foto.save(caminho_completo)

                    # Caminho relativo para o frontend
                    equipamento.foto = os.path.join('images', 'Fotos_Equipamento', tag, nome_arquivo).replace('\\', '/')

            session.commit()

            # Registro de log
            data_formatada = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
            mensagem_log = (
                f"(WEB) O Usuário {usuario} - {nivel_acesso} realizou uma edição de equipamento em {data_formatada}.\n"
                f"  - Equipamento: {equipamento.tag_equipamento}\n"
                f"  - Localização: {equipamento.localizacao}\n"
                f"  - Tipo: {equipamento.tipo}\n"
                f"  - Status: {equipamento.status}\n"
            )

            log_entry = EntradaLog(
                mensagem=mensagem_log,
                usuario=usuario,
                nivel_acesso=nivel_acesso,
                tipo_log="WEB - Equipamento - Edição",
                data_geracao=data_formatada
            )
            session.add(log_entry)
            session.commit()

            return jsonify({'success': 'Equipamento atualizado com sucesso!'})

        except Exception as e:
            session.rollback()
            print('Erro ao atualizar equipamento:', e)
            return jsonify({'error': 'Erro ao atualizar equipamento.'}), 500

    current_user = get_current_user(flask_session, session)
    return render_template('equipamento/edit_equipamento.html', equipamento=equipamento,
                           current_user=current_user, nivel_acesso=nivel_acesso,
                           foto_perfil_url=foto_perfil_url)

@app.route('/equipamentos/novo_equipamento')
@login_required
@permission_required('novo_equipamento')
def novo_equipamento():
    current_user = get_current_user(flask_session, session)
    return render_template('equipamento/novo_equipamento.html', current_user=current_user)

# Métodos de Logs ###################################################################

@app.route('/api/logs', methods=['GET'])
@login_required
def get_logs():
    try:
        selected_user = request.args.get('usuario', 'Todos')
        selected_type = request.args.get('tipo', 'Todos')
        selected_date = request.args.get('data', '')
        
        # Debug logging
        print(f"Received request with params - user: {selected_user}, type: {selected_type}, date: {selected_date}")
        
        query = session.query(EntradaLog)

        # Aplicar filtros
        if selected_user != 'Todos':
            query = query.filter(EntradaLog.usuario == selected_user)
        if selected_type != 'Todos':
            query = query.filter(EntradaLog.tipo_log.like(f'%{selected_type}%'))
        if selected_date:
            try:
                date_obj = datetime.strptime(selected_date, '%Y-%m-%d')
                query = query.filter(func.DATE(EntradaLog.data_geracao) == date_obj.date())
            except ValueError as ve:
                print(f"Date format error: {ve}")
                return jsonify({
                    'success': False,
                    'error': f"Invalid date format. Use YYYY-MM-DD"
                }), 400

        # Ordenação e execução da query
        logs = query.order_by(EntradaLog.id.desc()).all()
        
        logs_data = []
        for log in logs:
            try:
                log_date = log.data_geracao.isoformat() if log.data_geracao else None
            except AttributeError:
                log_date = str(log.data_geracao)
                
            logs_data.append({
                'id': log.id,
                'data_geracao': log_date,
                'usuario': log.usuario,
                'nivel_acesso': log.nivel_acesso,
                'tipo_log': log.tipo_log,
                'mensagem': log.mensagem,
                'detalhes': getattr(log, 'detalhes', None)  # Safe attribute access
            })

        # Obter usuários e tipos de logs únicos
        unique_users = [u[0] for u in session.query(EntradaLog.usuario).distinct().all() or []]
        unique_types = [t[0] for t in session.query(EntradaLog.tipo_log).distinct().all() or []]

        return jsonify({
            'success': True,
            'logs': logs_data,
            'unique_users': unique_users,
            'unique_types': unique_types,
            'stats': {
                'total_logs': len(logs_data),
                'login_logs': query.filter(EntradaLog.tipo_log.like('%Login%')).count(),
                'logout_logs': query.filter(EntradaLog.tipo_log.like('%Logout%')).count(),
                'create_logs': query.filter(EntradaLog.tipo_log.like('%Criar%') | 
                               EntradaLog.tipo_log.like('%Criação%')).count(),
                'update_logs': query.filter(EntradaLog.tipo_log.like('%Atualizar%') | 
                               EntradaLog.tipo_log.like('%Atualização%')).count(),
                'delete_logs': query.filter(EntradaLog.tipo_log.like('%Excluir%') | 
                               EntradaLog.tipo_log.like('%Exclusão%')).count()
            }
        })

    except Exception as e:
        print(f"Error in /api/logs: {str(e)}")
        return jsonify({
            'success': False,
            'error': "Internal server error",
            'details': str(e)
        }), 500
    
@app.route('/auditoria/logs_auditoria')
@login_required
@permission_required('logs_auditoria')
def logs_auditoria():
    current_user = get_current_user(flask_session, session)
    current_date = datetime.now().strftime('%Y-%m-%d')  # Formato compatível com input date
    return render_template('opcoes/logs_auditoria.html', 
                         current_user=current_user, 
                         current_date=current_date)

@app.route('/api/logs/filters', methods=['GET'])
@login_required
def get_logs_filters():
    try:
        unique_users = session.query(EntradaLog.usuario).distinct().all()
        unique_types = session.query(EntradaLog.tipo_log).distinct().all()

        unique_users = [user[0] for user in unique_users]
        unique_types = [type[0] for type in unique_types]

        return jsonify({
            'success': True,
            'unique_users': unique_users,
            'unique_types': unique_types
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

# Métodos de PM ###################################################################

@app.route('/api/manutencao/tickets', methods=['POST'])
@login_required
def criar_novo_ticket():
    try:
        app.logger.info(f"Dados recebidos: {request.form}")
        app.logger.info(f"Files recebidos: {request.files}")        
        # Obter dados do formulário
        if request.is_json:
            data = request.get_json()
        else:
            data = {
                'equipamento_id': request.form.get('equipamento_id'),
                'tipo': request.form.get('tipo'),
                'data_limite': request.form.get('data_limite'),
                'descricao': request.form.get('descricao'),  # Campo mantido como 'descricao'
                'prioridade': request.form.get('prioridade'),
                'responsavel_id': request.form.get('responsavel_id'),
                'inspecao_id': request.form.get('inspecao_id')
            }
        
        # Verificação dos campos obrigatórios
        required_fields = ['equipamento_id', 'tipo', 'data_limite', 'descricao']  # Mantido 'descricao'
        missing_fields = [field for field in required_fields if not data.get(field)]
        
        if missing_fields:
            return jsonify({
                'success': False,
                'error': f'Campos obrigatórios faltando: {", ".join(missing_fields)}'
            }), 400

        # Inicializar dados e obter usuário atual
        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)

        # Validação robusta dos dados
        if not all(field in data for field in required_fields):
            return jsonify({'error': f'Campos obrigatórios faltando: {", ".join(set(required_fields) - set(data.keys()))}'}), 400
        
        # Verificação do equipamento
        equipamento = session.get(Equipamento, data['equipamento_id'])
        if not equipamento:
            return jsonify({'error': 'Equipamento não encontrado'}), 404
        
        # Cálculo de prioridade com fallback
        prioridade = data.get('prioridade')
        if not prioridade:
            criticidade = getattr(equipamento, 'criticidade', 0)
            prioridade = 'Alta' if criticidade >= 7 else 'Média' if criticidade >= 4 else 'Baixa'
        
        # Validação da data
        try:
            data_limite = datetime.strptime(data['data_limite'], '%Y-%m-%d')
        except ValueError:
            return jsonify({'error': 'Formato de data inválido. Use YYYY-MM-DD'}), 400
        
        # Criação do ticket_id único
        ticket_id = f"MNT-{datetime.now().strftime('%Y%m%d%H%M%S')}"
        
        # Criação da manutenção
        nova_manutencao = Manutencao(
            ticket_id=ticket_id,
            equipamento_id=data['equipamento_id'],
            inspecao_id=data.get('inspecao_id'),
            tipo_manutencao=data['tipo'],
            responsavel_id=data.get('responsavel_id'),                
            prioridade=prioridade,
            data_limite=data_limite,
            descricao=data['descricao'],
            usuario_geracao=usuario,
            status='Aberto',
        )
        
        session.add(nova_manutencao)
        session.flush()  # Para obter o ID antes do commit
        
        # Processamento de anexos se existirem
        if 'anexos' in request.files:
            config = load_config()
            for file in request.files.getlist('anexos'):
                if file.filename == '': continue
                
                info = processar_anexo({
                    'filename': secure_filename(file.filename),
                    'file': file,
                    'content_type': file.content_type
                }, ticket_id, config)
                
                novo_anexo = AnexoManutencao(
                    manutencao_id=nova_manutencao.id,
                    nome=info['nome'],
                    caminho=info['url'],
                    tipo=info['tipo'],
                    tamanho=info['tamanho'],
                    data_upload=datetime.fromisoformat(info['data_upload'])
                )
                session.add(novo_anexo)
        
        session.commit()
        
        # Log do sistema
        data_atual = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
        mensagem_log = (
            f"(WEB) O Usuário {usuario} ({nivel_acesso}) criou um novo ticket de manutenção em {data_atual}.\n"
            f"Detalhes do ticket:\n"
            f"  - Ticket ID: {ticket_id}\n"
            f"  - Equipamento: {equipamento.tag_equipamento}\n"
            f"  - Tipo: {data['tipo']}\n"
            f"  - Prioridade: {prioridade}\n"
            f"  - Data Limite: {data_limite.strftime('%d-%m-%Y')}\n"
            f"  - Responsável: {data.get('responsavel_id', 'Não atribuído')}\n"
        )
        
        session.add(EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Manutenção - Novo Ticket",
            data_geracao=data_atual
        ))
        session.commit()
        
        return jsonify({
            'success': True,
            'ticket_id': ticket_id,
            'risco': calcular_nivel_risco(nova_manutencao),
            'message': 'Ticket criado com sucesso!'
        }), 201
        
    except IntegrityError as ie:
        session.rollback()
        app.logger.error(f"Erro de integridade ao criar ticket: {str(ie)}")
        return jsonify({'error': 'Erro de integridade ao criar ticket'}), 500
        
    except Exception as e:
        session.rollback()
        app.logger.error(f"Erro ao criar ticket: {str(e)}")
        return jsonify({'error': f'Erro ao criar ticket: {str(e)}'}), 500

def processar_anexo(anexo_data, ticket_id, config):
    """Processa upload de anexos e os salva na manutenção do equipamento"""
    nome = anexo_data.get('nome', 'arquivo')
    tipo = anexo_data.get('tipo', 'application/octet-stream')

    # Defina a pasta para armazenar os anexos
    upload_folder = config['UPLOAD_FOLDER']
    manutencao_folder = os.path.join(upload_folder, ticket_id)

    # Crie a pasta se não existir
    os.makedirs(manutencao_folder, exist_ok=True)

    # Nome seguro e caminho final
    filename = secure_filename(anexo_data['filename'])
    filepath = os.path.join(manutencao_folder, filename)

    # 🔧 Salvar o arquivo de fato
    file = anexo_data['file']
    file.save(filepath)  # <-- isso estava faltando!

    # Tamanho real após salvar
    tamanho = os.path.getsize(filepath)

    return {
        'nome': nome,
        'tipo': tipo,
        'tamanho': tamanho,
        'url': filepath,
        'data_upload': datetime.now().isoformat()
    }

# Função para salvar os anexos no banco de dados
def salvar_anexos_manutencao(ticket_id, anexos, config):
    manutencao = session.query(Manutencao).filter_by(ticket_id=ticket_id).first()

    if manutencao:
        # Processa e salva os anexos
        anexos_processados = []
        for anexo in anexos:
            anexo_processado = processar_anexo(anexo, config)
            anexos_processados.append({
                'nome': anexo_processado['nome'],
                'caminho': anexo_processado['url'],  # Caminho do arquivo salvo
                'tipo': anexo_processado['tipo'],
                'tamanho': anexo_processado['tamanho'],
                'data_upload': anexo_processado['data_upload']
            })
        
        # Atualiza o campo de anexos da manutenção (agora estamos salvando como uma lista de dicionários)
        if manutencao.anexos:
            manutencao.anexos.extend(anexos_processados)  # Adiciona novos anexos à lista existente
        else:
            manutencao.anexos = anexos_processados  # Se não houver anexos, cria uma nova lista
        
        session.commit()
        return {"message": "Anexos enviados com sucesso"}
    else:
        return {"error": "Manutenção não encontrada"}, 404

# Função para verificar se o arquivo é permitido
def allowed_file(filename):
    config = load_config()
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in config['ALLOWED_EXTENSIONS']

@app.route('/api/manutencao/anexo/<int:anexo_id>', methods=['DELETE'])
@login_required
def delete_anexo_manutencao(anexo_id):
    try:
        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        if not usuario:
            return jsonify({'error': 'Não autenticado'}), 401

        anexo = session.get(AnexoManutencao, anexo_id)
        if not anexo:
            return jsonify({'error': 'Anexo não encontrado'}), 404

        # Verificar se o usuário tem permissão
        manutencao = session.get(Manutencao, anexo.manutencao_id)
        if not manutencao:
            return jsonify({'error': 'Manutenção não encontrada'}), 404

        # Apenas administradores ou o responsável podem deletar
        if nivel_acesso < 3 and (not manutencao.responsavel_id or manutencao.responsavel_id != session.get('user_id')):
            return jsonify({'error': 'Acesso não autorizado'}), 403

        # Remover arquivo físico
        try:
            if os.path.exists(anexo.caminho):
                os.remove(anexo.caminho)
        except Exception as e:
            app.logger.error(f"Erro ao remover arquivo físico: {str(e)}")

        # Remover do banco de dados
        session.delete(anexo)
        session.commit()

        return jsonify({'success': True})

    except Exception as e:
        session.rollback()
        app.logger.error(f"Erro ao deletar anexo: {str(e)}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/manutencao/<ticket_id>/anexos', methods=['POST'])
@login_required
def upload_anexos_manutencao(ticket_id):
    try:
        # Verificar autenticação
        usuario, _, _ = get_current_user_name_and_access_level(flask_session, session)
        if not usuario:
            return jsonify({'error': 'Não autenticado'}), 401

        # Verificar se existem arquivos
        if 'anexos' not in request.files:
            return jsonify({"error": "Nenhum anexo encontrado"}), 400

        files = request.files.getlist('anexos')
        if not files or all(file.filename == '' for file in files):
            return jsonify({"error": "Nenhum arquivo enviado"}), 400

        # Verificar se o ticket existe
        manutencao = session.query(Manutencao).filter_by(ticket_id=ticket_id).first()
        if not manutencao:
            return jsonify({"error": "Manutenção não encontrada"}), 404

        # Configurações
        config = load_config()
        upload_folder = os.path.join(config['UPLOAD_FOLDER'], 'manutencao')
        os.makedirs(upload_folder, exist_ok=True)

        anexos_salvos = []
        for file in files:
            if file and allowed_file(file.filename):
                # Nome seguro para o arquivo
                filename = secure_filename(f"{uuid.uuid4().hex}_{file.filename}")
                filepath = os.path.join(upload_folder, filename)
                
                # Salvar arquivo
                file.save(filepath)

                # Criar registro no banco
                novo_anexo = AnexoManutencao(
                    manutencao_id=manutencao.id,
                    caminho=filepath,
                    nome=file.filename,
                    tipo=file.content_type,
                    tamanho=os.path.getsize(filepath)
                )
                session.add(novo_anexo)
                session.flush()  # Para obter o ID

                anexos_salvos.append({
                    'id': novo_anexo.id,
                    'nome': file.filename,
                    'tipo': file.content_type,
                    'tamanho': os.path.getsize(filepath),
                    'url': f"/uploads/manutencao/{filename}",  # URL acessível
                    'data_upload': novo_anexo.data_upload.isoformat()
                })

        session.commit()

        salvar_anexos_manutencao(ticket_id, anexos_salvos, config)

        return jsonify({
            "success": True,
            "anexos": anexos_salvos,
            "message": f"{len(anexos_salvos)} anexo(s) salvo(s) com sucesso"
        })

    except Exception as e:
        session.rollback()
        app.logger.error(f"Erro ao processar anexos: {str(e)}")
        return jsonify({
            "error": "Erro ao processar anexos",
            "details": str(e)
        }), 500
    
@app.route('/api/manutencao/<int:manutencao_id>/anexos', methods=['GET'])
@login_required
def get_anexos_manutencao(manutencao_id):
    try:
        ensure_config_loaded()
        upload_folder = os.path.abspath(app.config['UPLOAD_FOLDER'])

        anexos = session.query(AnexoManutencao).filter(AnexoManutencao.manutencao_id == manutencao_id).all()
        if not anexos:
            return jsonify([]), 200

        anexos_urls = []
        for anexo in anexos:
            # Verifica se o arquivo existe
            if not os.path.exists(anexo.caminho):
                app.logger.warning(f"Arquivo não encontrado: {anexo.caminho}")
                continue
                
            # Obtém o caminho relativo dentro da pasta de uploads
            rel_path = os.path.relpath(anexo.caminho, upload_folder)
            web_path = rel_path.replace(os.sep, '/')
            
            # Obtém apenas o nome do arquivo e a subpasta (ticket_id)
            path_parts = web_path.split('/')
            if len(path_parts) >= 2:
                relative_url = f"{path_parts[-2]}/{path_parts[-1]}"
            else:
                relative_url = path_parts[-1]

            anexos_urls.append({
                "id": anexo.id,
                "url": f"/uploads/{relative_url}",  # URL acessível via rota /uploads/<path>
                "caminho": anexo.caminho,
                "nome": anexo.nome,
                "tipo": anexo.tipo,
                "tamanho": anexo.tamanho,
                "data_upload": anexo.data_upload.isoformat() if anexo.data_upload else None
            })

        return jsonify(anexos_urls)
    except Exception as e:
        app.logger.error(f'Erro ao buscar anexos da manutenção {manutencao_id}: {str(e)}')
        return jsonify({'erro': 'Erro ao buscar anexos'}), 500
    
@app.route('/api/manutencao/tickets/<int:ticket_id>', methods=['PUT'])
@login_required
def atualizar_ticket(ticket_id):
    try:
        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        data = request.get_json() if request.is_json else request.form.to_dict()
        
        # Verificar se o ticket existe
        manutencao = session.query(Manutencao).filter_by(id=ticket_id).first()
        if not manutencao:
            return jsonify({'error': 'Manutenção não encontrada'}), 404
        
        # Registro de alterações
        alteracoes = {}
        campos_permitidos = ['status', 'prioridade', 'responsavel_id', 'data_limite', 'descricao']
        
        for campo in campos_permitidos:
            if campo in data and getattr(manutencao, campo) != data[campo]:
                alteracoes[campo] = {
                    'de': getattr(manutencao, campo),
                    'para': data[campo]
                }
                setattr(manutencao, campo, data[campo])
        
        # Registrar histórico se houver alterações
        if alteracoes:
            manutencao.historico = (manutencao.historico or []) + [{
                'data': datetime.now().isoformat(),
                'usuario': usuario,
                'alteracoes': alteracoes
            }]

        # Tratamento especial para mudanças de status
        if 'status' in data:
            agora = datetime.now()
            
            if data['status'] == 'Concluído':
                manutencao.data_conclusao = agora
                manutencao.tempo_resolucao = (agora - manutencao.data_abertura).total_seconds() / 3600
                if manutencao.equipamento:
                    manutencao.equipamento.data_ultima_manutencao = agora
                    manutencao.equipamento.status = 'Ativo'
            
            elif data['status'] in ['Aberto', 'Em andamento']:
                if manutencao.equipamento:
                    manutencao.equipamento.data_ultima_manutencao = agora
                    manutencao.equipamento.status = 'Manutenção'
            
            elif data['status'] == 'Cancelado':
                manutencao.data_conclusao = agora
                if manutencao.equipamento:
                    manutencao.equipamento.status = 'Ativo' if manutencao.equipamento.status == 'Manutenção' else manutencao.equipamento.status

        
        # Processamento de anexos se existirem
        if 'anexos' in request.files:
            config = load_config()
            for file in request.files.getlist('anexos'):
                if file.filename == '': continue
                
                info = processar_anexo({
                    'filename': secure_filename(file.filename),
                    'file': file,
                    'content_type': file.content_type
                }, config)
                
                novo_anexo = AnexoManutencao(
                    manutencao_id=manutencao.id,
                    nome=info['nome'],
                    caminho=info['url'],
                    tipo=info['tipo'],
                    tamanho=info['tamanho'],
                    data_upload=datetime.fromisoformat(info['data_upload'])
                )
                session.add(novo_anexo)
        
        session.commit()
        
        # Log do sistema
        data_atual = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
        mensagem_log = (
            f"(WEB) O Usuário {usuario} ({nivel_acesso}) atualizou o ticket {manutencao.ticket_id} em {data_atual}.\n"
            f"Alterações realizadas:\n" +
            "\n".join([f"  - {campo}: de {dados['de']} para {dados['para']}" 
                       for campo, dados in alteracoes.items()])
        )
        
        session.add(EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Manutenção - Atualização de Ticket",
            data_geracao=data_atual
        ))
        session.commit()
        
        return jsonify({
            'success': True,
            'risco': calcular_nivel_risco(manutencao),
            'message': 'Ticket atualizado com sucesso!'
        })
        
    except IntegrityError as ie:
        session.rollback()
        app.logger.error(f"Erro de integridade ao atualizar ticket: {str(ie)}")
        return jsonify({'error': 'Erro de integridade ao atualizar ticket'}), 500
        
    except Exception as e:
        session.rollback()
        app.logger.error(f"Erro ao atualizar ticket: {str(e)}")
        return jsonify({'error': f'Erro ao atualizar ticket: {str(e)}'}), 500

@app.route('/api/manutencao/tickets/<int:ticket_id>', methods=['DELETE'])
@login_required
def deletar_ticket(ticket_id):
    try:
        usuario, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        
        # Verificar nível de acesso
        if nivel_acesso < 3:
            return jsonify({'error': 'Acesso não autorizado'}), 403
                
        # Verificar se o ticket existe
        manutencao = session.get(Manutencao, ticket_id)
        if not manutencao:
            return jsonify({'error': 'Manutenção não encontrada'}), 404
        
        # Registrar informações para o log antes de deletar
        ticket_info = {
            'ticket_id': manutencao.ticket_id,
            'equipamento': manutencao.equipamento.tag_equipamento if manutencao.equipamento else 'N/A',
            'status': manutencao.status
        }
        
        # Deletar o ticket
        session.delete(manutencao)
        session.commit()
        
        # Log do sistema
        data_atual = datetime.now().strftime('%d-%m-%Y %H:%M:%S')
        mensagem_log = (
            f"(WEB) O Usuário {usuario} ({nivel_acesso}) deletou o ticket {ticket_info['ticket_id']} em {data_atual}.\n"
            f"Informações do ticket deletado:\n"
            f"  - Equipamento: {ticket_info['equipamento']}\n"
            f"  - Status: {ticket_info['status']}\n"
        )
        
        session.add(EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Manutenção - Exclusão de Ticket",
            data_geracao=data_atual
        ))
        session.commit()
        
        return jsonify({
            'success': True,
            'message': 'Ticket deletado com sucesso!'
        })
        
    except IntegrityError as ie:
        session.rollback()
        app.logger.error(f"Erro de integridade ao deletar ticket: {str(ie)}")
        return jsonify({'error': 'Erro de integridade ao deletar ticket'}), 500
        
    except Exception as e:
        session.rollback()
        app.logger.error(f"Erro ao deletar ticket: {str(e)}")
        return jsonify({'error': f'Erro ao deletar ticket: {str(e)}'}), 500

@app.route('/api/manutencao/tickets', methods=['GET'])
@login_required
def listar_manutencoes():
    try:
        # Obter parâmetros de filtro da query string
        status = request.args.get('status')
        prioridade = request.args.get('prioridade')
        
        query = session.query(Manutencao)
        
        if status:
            query = query.filter(Manutencao.status == status)
        if prioridade:
            query = query.filter(Manutencao.prioridade == prioridade)
            
        manutencoes = query.order_by(Manutencao.data_limite.asc()).all()
        
        return jsonify({
            'data': [{
                'id': m.id,
                'ticket_id': m.ticket_id,
                'equipamento': {
                    'id': m.equipamento.id,
                    'tag': m.equipamento.tag_equipamento,
                    'nome': m.equipamento.modelo,
                    'localizacao': m.equipamento.localizacao,
                    'criticidade': getattr(m.equipamento, 'criticidade', 0)
                },
                'tipo_manutencao': m.tipo_manutencao,
                'prioridade': m.prioridade,
                'status': m.status,
                'responsavel': {
                    'id': m.responsavel.id,
                    'nome': m.responsavel.nome
                } if m.responsavel else None,
                'data_abertura': m.data_abertura.isoformat(),
                'data_limite': m.data_limite.isoformat(),
                'dias_restantes': (m.data_limite - datetime.now()).days,
                'descricao': m.descricao,
                'risco': calcular_nivel_risco(m)
            } for m in manutencoes]
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/manutencoes_detalhes/<int:manutencao_id>', methods=['GET'])
@login_required
def obter_detalhes_manutencao_individual(manutencao_id):
    try:
        # Verificação de autenticação
        usuario, nivel_acesso, foto_perfil_url = get_current_user_name_and_access_level(flask_session, session)
        if not usuario or usuario == "Desconhecido":
            return jsonify({'error': 'Não autenticado'}), 401

        manutencao = session.get(Manutencao, manutencao_id)
        if manutencao:
            session.refresh(manutencao)

        if not manutencao:
            return jsonify({'error': 'Manutenção não encontrada'}), 404

        equipamento = manutencao.equipamento
        responsavel = manutencao.responsavel

        dados_manutencao = {
            'id': manutencao.id,
            'ticket_id': manutencao.ticket_id,
            'equipamento': {
                'id': equipamento.id if equipamento else None,
                'tag': equipamento.tag_equipamento if equipamento else 'N/A',
                'nome': equipamento.modelo if equipamento else 'N/A',
                'localizacao': equipamento.localizacao if equipamento else 'N/A',
                'criticidade': getattr(equipamento, 'criticidade', 0)
            },
            'tipo_manutencao': manutencao.tipo_manutencao,
            'prioridade': manutencao.prioridade,
            'status': manutencao.status,
            'responsavel': {
                'id': responsavel.id if responsavel else None,
                'nome': responsavel.nome if responsavel else None,
                'foto': responsavel.foto_perfil if responsavel else None
            } if responsavel else None,
            'data_abertura': manutencao.data_abertura.strftime('%Y-%m-%d %H:%M') if manutencao.data_abertura else None,
            'data_limite': manutencao.data_limite.strftime('%Y-%m-%d') if manutencao.data_limite else None,
            'dias_restantes': (manutencao.data_limite - datetime.now()).days if manutencao.data_limite else None,
            'descricao': manutencao.descricao,
            'historico': manutencao.historico or [],
            'anexos': [anexo.to_dict() for anexo in (manutencao.anexos or [])],

            'risco': calcular_nivel_risco(manutencao)
        }

        return jsonify(dados_manutencao)

    except Exception as e:
        app.logger.error(f"Erro ao buscar manutenção {manutencao_id}: {str(e)}\n{traceback.format_exc()}")
        return jsonify({'error': 'Erro interno ao buscar manutenção'}), 500
 
@app.route('/api/manutencao/<int:manutencao_id>', methods=['PUT'])
@login_required
def atualizar_manutencao(manutencao_id):
    data = request.get_json()
    usuario, _, _ = get_current_user_name_and_access_level(flask_session, session)
    
    # Tenta recuperar a manutenção com versão ou timestamp
    manutencao = session.query(Manutencao).filter_by(id=manutencao_id).with_for_update().first()
    if not manutencao:
        return jsonify({'error': 'Manutenção não encontrada'}), 404
    
    alteracoes = {}
    for campo in ['status', 'prioridade', 'responsavel_id', 'data_limite', 'descricao']:
        if campo in data and getattr(manutencao, campo) != data[campo]:
            alteracoes[campo] = {
                'de': getattr(manutencao, campo),
                'para': data[campo]
            }
            setattr(manutencao, campo, data[campo])

    try:
        session.commit()
        return jsonify({'success': True, 'risco': calcular_nivel_risco(manutencao)})
    except StaleDataError:
        session.rollback()
        return jsonify({'error': 'Erro: Dados desatualizados. Tente novamente.'}), 409
    except Exception as e:
        session.rollback()
        app.logger.error(f"Erro ao atualizar manutenção {manutencao_id}: {str(e)}")
        return jsonify({'error': 'Erro ao atualizar manutenção'}), 500
    
@app.route('/api/listar_equipamentos_cadastrados')
@login_required
def listar_equipamentos_cadastrados():
    try:
        equipamentos = session.query(Equipamento).all()
        return jsonify({
            'data': [{ 
                'id': e.id,
                'tag_equipamento': e.tag_equipamento,
                'modelo': e.modelo,
                'tipo': e.tipo,  # Para compatibilidade                
                'localizacao': e.localizacao,
                'criticidade': getattr(e, 'criticidade', 0)
            } for e in equipamentos]
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/listar_usuarios_cadastrados', methods=['GET'])
@login_required
def listar_usuarios_cadastrados():
    try:
        # Query básica apenas com os campos necessários
        usuarios = session.query(
            Usuario.id,
            Usuario.nome,
            Usuario.primeiro_nome,
            Usuario.sobrenome,
            Usuario.matricula
        ).order_by(Usuario.nome).all()

        # Formatação simplificada dos dados
        usuarios_data = [{
            'id': u.id,
            'nome_completo': u.nome if u.nome else f"{u.primeiro_nome or ''} {u.sobrenome or ''}".strip(),
            'matricula': u.matricula
        } for u in usuarios]

        return jsonify({
            'success': True,
            'data': usuarios_data
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500
    
def calcular_nivel_risco(manutencao):
    """Calcula o nível de risco de uma manutenção com base em múltiplos fatores"""
    if manutencao.status == 'Concluído':
        return 0
        
    risco = 0
    
    # Fator prioridade - garantindo valor padrão se não existir
    fatores_prioridade = {'Alta': 3, 'Média': 2, 'Baixa': 1}
    risco += fatores_prioridade.get(manutencao.prioridade, 1)
    
    # Fator criticidade do equipamento - convertendo para float se necessário
    if manutencao.equipamento and hasattr(manutencao.equipamento, 'criticidade'):
        try:
            criticidade = float(getattr(manutencao.equipamento, 'criticidade', 0))
            risco += criticidade / 2
        except (TypeError, ValueError):
            risco += 0  # Se não puder converter, ignora este fator
    
    # Fator tempo restante - garantindo que data_limite existe e é datetime
    if hasattr(manutencao, 'data_limite') and manutencao.data_limite:
        try:
            dias_restantes = (manutencao.data_limite - datetime.now()).days
            if dias_restantes < 0:
                risco += 5  # Atrasado
            elif dias_restantes < 3:
                risco += 3
            elif dias_restantes < 7:
                risco += 1
        except TypeError:
            risco += 0  # Se não puder calcular dias, ignora este fator
    
    # Fator tipo de manutenção - garantindo valor padrão
    tipo = getattr(manutencao, 'tipo_manutencao', '')
    if tipo == 'Corretiva':
        risco += 2
    
    # Normalizar para escala 0-10
    risco = min(max(float(risco), 0), 10)  # Garantindo que risco é float
    
    return round(risco, 1)

@app.route('/opcoes/gestao_manutencao')
@login_required
def abre_gestao_manutencao():
    current_user = get_current_user(flask_session, session)
    current_date = datetime.now().strftime('%Y-%m-%d')
    
    # Adicionar os selects básicos diretamente no template
    usuarios = session.query(Usuario).order_by(Usuario.nome).all()
    
    return render_template('opcoes/gestao_manutencao.html', 
                         current_user=current_user, 
                         current_date=current_date,
                         usuarios=usuarios)

# Métodos de Relatórios de Inspeções ###################################################################

@app.route('/relatorios/relatorio_inspecoes')
@login_required
@permission_required('relatorio_inspecoes')
def relatorio_inspecoes():
    current_user = get_current_user(flask_session, session)
    return render_template('relatorios/relatorio_inspecoes.html', current_user=current_user)

@app.route('/api/inspecoes_relatorios', methods=['GET'])
@login_required
def get_inspecoes_relatorios():
    try:
        # Parâmetros de paginação
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 15))
        
        # Parâmetros de filtro
        search_term = request.args.get('searchTerm', '').strip()
        status_filter = request.args.get('status', '').strip()
        equipment_type = request.args.get('equipmentType', '').strip()
        responsible = request.args.get('responsible', '').strip()
        start_date = request.args.get('startDate', '').strip()
        end_date = request.args.get('endDate', '').strip()

        if page < 1 or limit < 1:
            return jsonify({'error': 'Parâmetros de paginação inválidos.'}), 400

        query = (
            session.query(Registro, Equipamento)
            .join(Equipamento, Registro.equipamento_id == Equipamento.id)
        )

        # Aplicar filtros
        if search_term:
            query = query.filter(
                or_(
                    Equipamento.tag_equipamento.ilike(f'%{search_term}%'),
                    Registro.localizacao.ilike(f'%{search_term}%'),
                    Registro.responsavel.ilike(f'%{search_term}%'),
                    Registro.numero_inspecao.ilike(f'%{search_term}%')
                )
            )

        if status_filter:
            if status_filter not in ['Pendente', 'Finalizada', 'Cancelada', 'Atrasada']:
                return jsonify({'error': 'Status de inspeção inválido.'}), 400
            query = query.filter(Registro.status_inspecao == status_filter)

        if equipment_type:
            query = query.filter(Equipamento.tipo == equipment_type)

        if responsible:
            query = query.filter(Registro.responsavel == responsible)

        if start_date and end_date:
            try:
                start_date = datetime.strptime(start_date, '%Y-%m-%d')
                end_date = datetime.strptime(end_date, '%Y-%m-%d')
                query = query.filter(Registro.data_registro.between(start_date, end_date))
            except ValueError:
                return jsonify({'error': 'Formato de data inválido. Use YYYY-MM-DD.'}), 400

        total_inspecoes = query.count()
        resultados = query.offset((page - 1) * limit).limit(limit).all()

        inspecoes_data = [
            {
                'id': registro.id,
                'numero_inspecao': registro.numero_inspecao,
                'equipamento_id': registro.equipamento_id,
                'tag_equipamento': equipamento.tag_equipamento,
                'tipo_equipamento': equipamento.tipo,
                'localizacao': registro.localizacao or equipamento.localizacao,
                'responsavel': registro.responsavel,
                'criticidade_inspecao': registro.criticidade_inspecao,                  
                'data_registro': registro.data_registro.strftime('%d-%m-%Y'),
                'data_validade_inspecao': registro.data_validade_inspecao.strftime('%d-%m-%Y'),
                'status_inspecao': registro.status_inspecao,
                'status_equipamento': registro.status_equipamento,
                'observacoes': registro.observacoes,
                'motivo_acao': registro.motivo_acao
            }
            for registro, equipamento in resultados
        ]

        return jsonify({
            'data': inspecoes_data,
            'total': total_inspecoes,
            'page': page,
            'limit': limit
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/inspecoes/contagem', methods=['GET'])
@login_required
def get_inspecoes_contagem():
    try:
        status_filter = request.args.get('status', '').strip()
        
        query = session.query(Registro)
        
        if status_filter:
            if status_filter not in ['Pendente', 'Finalizada', 'Cancelada', 'Atrasada']:
                return jsonify({'error': 'Status de inspeção inválido.'}), 400
            query = query.filter(Registro.status_inspecao == status_filter)
        
        count = query.count()
        
        return jsonify({
            'status': status_filter or 'Total',
            'total': count
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/usuarios/responsaveis', methods=['GET'])
@login_required
def get_usuarios_responsaveis():
    try:
        usuarios = session.query(Usuario.id, Usuario.nome, Usuario.matricula)\
            .outerjoin(StatusUsuario, Usuario.id == StatusUsuario.usuario_id)\
            .filter(or_(StatusUsuario.status == 'Acesso Liberado', StatusUsuario.id == None))\
            .all()
        
        responsaveis = [{
            'id': usuario.id,
            'nome': usuario.nome,
            'matricula': usuario.matricula
        } for usuario in usuarios]
        
        return jsonify(responsaveis)
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/equipamentos_relatorios/<int:equipamento_id>', methods=['GET'])
@login_required
def get_equipamento_detalhes_relatorios(equipamento_id):
    try:
        equipamento = session.query(Equipamento).filter_by(id=equipamento_id).first()
        if not equipamento:
            return jsonify({'error': 'Equipamento não encontrado'}), 404

        # Calcular status com base na validade
        status = equipamento.status
        if status == "Ativo" and equipamento.data_validade < datetime.now().date():
            status = "Vencido"

        equipamento_data = {
            'id': equipamento.id,
            'tag_equipamento': equipamento.tag_equipamento,
            'localizacao': equipamento.localizacao,
            'tipo': equipamento.tipo,
            'classe': equipamento.classe,
            'data_fabricacao': equipamento.data_fabricacao.strftime('%Y-%m-%d'),  # Formato ISO
            'peso': equipamento.peso,
            'teste_hidrostatico_n2': equipamento.teste_hidrostatico_n2.strftime('%Y-%m-%d') if equipamento.teste_hidrostatico_n2 else None,
            'teste_hidrostatico_n3': equipamento.teste_hidrostatico_n3.strftime('%Y-%m-%d') if equipamento.teste_hidrostatico_n3 else None,
            'data_validade': equipamento.data_validade.strftime('%Y-%m-%d'),
            'status': status,
            'codigo_barras': equipamento.codigo_barras,
            'data_ultima_inspecao': equipamento.data_ultima_inspecao.strftime('%Y-%m-%d') if equipamento.data_ultima_inspecao else None,
            'data_proxima_inspecao': equipamento.data_proxima_inspecao.strftime('%Y-%m-%d') if equipamento.data_proxima_inspecao else None,
        }

        return jsonify(equipamento_data)

    except Exception as e:
        print(f"Erro ao buscar equipamento: {str(e)}")  # Log para debug
        return jsonify({'error': str(e)}), 500

# Métodos de Relatórios de Equipamento ###################################################################

@app.route('/relatorios/relatorio_equipamentos')
@login_required
@permission_required('relatorio_equipamentos')
def relatorio_equipamentos():
    current_user = get_current_user(flask_session, session)
    return render_template('relatorios/relatorio_equipamentos.html', current_user=current_user)

@app.route('/api/equipamentos_relatorios', methods=['GET'])
@login_required
def get_equipamentos_cadastrados_relatorios():
    try:
        # Parâmetros de busca e filtros
        search_term = request.args.get('searchTerm', '').strip()
        status_filter = request.args.get('status', '').strip()
        type_filter = request.args.get('type', '').strip()
        location_filter = request.args.get('location', '').strip()
        start_date = request.args.get('startDate', '').strip()
        end_date = request.args.get('endDate', '').strip()
        
        # Parâmetros de paginação
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 15))
        
        if page < 1 or limit < 1:
            return jsonify({'error': 'Parâmetros de paginação inválidos.'}), 400

        query = session.query(Equipamento)

        # Aplicar filtros
        if search_term:
            query = query.filter(
                or_(
                    Equipamento.tag_equipamento.ilike(f'%{search_term}%'),
                    Equipamento.localizacao.ilike(f'%{search_term}%'),
                    Equipamento.tipo.ilike(f'%{search_term}%'),
                    Equipamento.codigo_barras.ilike(f'%{search_term}%')
                )
            )

        if status_filter:
            valid_statuses = ['Ativo', 'Manutenção', 'Descartado', 'Reserva', 'Vencido']
            if status_filter not in valid_statuses:
                return jsonify({'error': 'Status de equipamento inválido.'}), 400
            query = query.filter(Equipamento.status == status_filter)

        if type_filter:
            query = query.filter(Equipamento.tipo == type_filter)

        if location_filter:
            query = query.filter(Equipamento.localizacao == location_filter)

        if start_date and end_date:
            try:
                start_date = datetime.strptime(start_date, '%Y-%m-%d')
                end_date = datetime.strptime(end_date, '%Y-%m-%d')
                query = query.filter(
                    Equipamento.data_validade.between(start_date, end_date)
                )
            except ValueError:
                return jsonify({'error': 'Formato de data inválido. Use YYYY-MM-DD.'}), 400

        # Contagem total antes da paginação
        total_equipamentos = query.count()

        # Aplicar paginação
        equipamentos = query.offset((page - 1) * limit).limit(limit).all()

        # Formatar dados para resposta
        equipamentos_data = [
            {
                'id': equipamento.id,
                'tag_equipamento': equipamento.tag_equipamento,
                'localizacao': equipamento.localizacao,
                'tipo': equipamento.tipo,
                'classe': equipamento.classe,
                'data_fabricacao': equipamento.data_fabricacao.strftime('%d-%m-%Y') if equipamento.data_fabricacao else None,
                'peso': equipamento.peso,
                'teste_hidrostatico_n2': equipamento.teste_hidrostatico_n2.strftime('%d-%m-%Y') if equipamento.teste_hidrostatico_n2 else None,
                'teste_hidrostatico_n3': equipamento.teste_hidrostatico_n3.strftime('%d-%m-%Y') if equipamento.teste_hidrostatico_n3 else None,
                'data_validade': equipamento.data_validade.strftime('%d-%m-%Y') if equipamento.data_validade else None,
                'status': equipamento.status,
                'codigo_barras': equipamento.codigo_barras,
                'data_ultima_inspecao': equipamento.data_ultima_inspecao.strftime('%d-%m-%Y') if equipamento.data_ultima_inspecao else None,
                'data_proxima_inspecao': equipamento.data_proxima_inspecao.strftime('%d-%m-%Y') if equipamento.data_proxima_inspecao else None
            }
            for equipamento in equipamentos
        ]

        return jsonify({
            'data': equipamentos_data,
            'total': total_equipamentos,
            'page': page,
            'limit': limit,
            'total_pages': (total_equipamentos + limit - 1) // limit
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/equipamentos/contagem_status', methods=['GET'])
@login_required
def get_contagem_status_equipamentos():
    try:
        # Consulta para contar equipamentos por status
        contagem = session.query(
            Equipamento.status,
            func.count(Equipamento.id).label('quantidade')
        ).group_by(Equipamento.status).all()

        # Converter para dicionário
        resultado = {status: quantidade for status, quantidade in contagem}

        # Adicionar contagem de equipamentos vencidos
        hoje = datetime.now().date()
        vencidos = session.query(func.count(Equipamento.id)).filter(
            Equipamento.data_validade < hoje,
            Equipamento.status != 'Descartado'
        ).scalar()
        
        resultado['Vencido'] = vencidos

        return jsonify(resultado)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/equipamentos/tipos', methods=['GET'])
@login_required
def get_tipos_equipamentos():
    try:
        # Consulta para obter tipos únicos de equipamentos
        tipos = session.query(
            Equipamento.tipo
        ).distinct().order_by(Equipamento.tipo).all()

        # Extrair apenas os valores
        tipos_lista = [tipo[0] for tipo in tipos if tipo[0]]

        return jsonify(tipos_lista)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/equipamentos/localizacoes', methods=['GET'])
@login_required
def get_localizacoes_equipamentos():
    try:
        # Consulta para obter localizações únicas de equipamentos
        localizacoes = session.query(
            Equipamento.localizacao
        ).distinct().order_by(Equipamento.localizacao).all()

        # Extrair apenas os valores
        localizacoes_lista = [loc[0] for loc in localizacoes if loc[0]]

        return jsonify(localizacoes_lista)

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
# Métodos de Usuários ###################################################################

@app.route('/usuarios/usuarios_cadastrados')
@login_required
@permission_required('usuarios_cadastrados')
def usuarios_cadastrados():
    current_user = get_current_user(flask_session, session)
    allowed_menus = get_allowed_menus_for_user(current_user)    
    return render_template('opcoes/usuarios/usuarios_cadastrados.html', current_user=current_user, allowed_menus=allowed_menus)

@app.route('/api/usuarios/<int:usuario_id>/inspecoes', methods=['GET'])
@login_required
def get_historico_inspecoes_responsavel(usuario_id):
    usuario = session.query(Usuario).filter_by(id=usuario_id).first()
    if not usuario:
        return jsonify({'error': 'Usuário não encontrado'}), 404

    inspecoes = session.query(Registro).filter_by(responsavel=usuario.nome).all()

    inspecoes_data = [
        {
            'id': inspecao.id,
            'numero_inspecao': inspecao.numero_inspecao,            
            'equipamento': inspecao.equipamento_id,
            'tipo': inspecao.tipo,
            'classe': inspecao.classe,
            'localizacao': inspecao.localizacao,
            'status': inspecao.status_equipamento,
            'responsavel': inspecao.responsavel,
            'data_inspecao': inspecao.data_inspecao.strftime('%d/%m/%Y') if inspecao.data_inspecao else '',
            'validade': inspecao.data_validade_inspecao.strftime('%d/%m/%Y') if inspecao.data_validade_inspecao else '',
            'observacoes': inspecao.observacoes
        }
        for inspecao in inspecoes
    ]

    return jsonify(inspecoes_data)

@app.route('/api/usuarios', methods=['GET'])
@login_required
def get_usuarios():
    try:
        # Parâmetros do DataTables
        draw = request.args.get('draw', type=int)
        start = request.args.get('start', type=int)
        length = request.args.get('length', type=int)
        search_value = request.args.get('search[value]', '').strip()
        
        # Filtros adicionais
        status_filter = request.args.get('status', '').strip()
        setor_filter = request.args.get('setor', '').strip()
        nivel_acesso_filter = request.args.get('nivel_acesso', '').strip()
        
        # Ordenação
        order_column_index = request.args.get('order[0][column]', default=0, type=int)
        order_direction = request.args.get('order[0][dir]', 'asc')
        
        # Mapeamento de colunas
        columns = [
            'nome',
            'matricula',
            'setor',
            'nivel_acesso',
            'status',
            'id'  # Coluna de ações
        ]
        order_column = columns[order_column_index] if order_column_index < len(columns) else 'nome'

        # Construção da query
        query = session.query(Usuario)

        # Aplicar filtro de pesquisa geral
        if search_value:
            query = query.filter(
                or_(
                    Usuario.nome.ilike(f'%{search_value}%'),
                    Usuario.matricula.ilike(f'%{search_value}%'),
                    Usuario.email.ilike(f'%{search_value}%'),
                    Usuario.usuario.ilike(f'%{search_value}%')
                )
            )

        # Aplicar filtros adicionais
        if status_filter:
            query = query.join(StatusUsuario).filter(StatusUsuario.status == status_filter)
            
        if setor_filter:
            query = query.filter(Usuario.setor.ilike(f'%{setor_filter}%'))
            
        if nivel_acesso_filter:
            query = query.filter(Usuario.nivel_acesso.ilike(f'%{nivel_acesso_filter}%'))

        # Ordenação
        if order_direction == 'desc':
            query = query.order_by(desc(getattr(Usuario, order_column)))
        else:
            query = query.order_by(getattr(Usuario, order_column))

        # Contagem total antes da paginação
        total_records = query.count()

        # Paginação
        usuarios = query.offset(start).limit(length).all()

        # Formatar dados
        usuarios_data = []
        for usuario in usuarios:
            # Obtém o status mais recente do usuário
            status_usuario = session.query(StatusUsuario).filter_by(usuario_id=usuario.id).order_by(StatusUsuario.status.desc()).first()
            status = status_usuario.status if status_usuario else 'Acesso Liberado'

            usuarios_data.append({
                'id': usuario.id,
                'nome': f"{usuario.primeiro_nome} {usuario.sobrenome}",
                'nome_completo': usuario.nome,
                'matricula': usuario.matricula,
                'setor': usuario.setor,
                'email': usuario.email,
                'usuario': usuario.usuario,
                'nivel_acesso': usuario.nivel_acesso,
                'status': status,
                'DT_RowId': f'usuario_{usuario.id}'  # Para referência no frontend
            })

        return jsonify({
            'draw': draw,
            'recordsTotal': total_records,
            'recordsFiltered': total_records,
            'data': usuarios_data
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    
@app.route('/api/usuarios/<int:usuario_id>', methods=['GET'])
@login_required
def get_usuario_detalhes(usuario_id):
    usuario = session.query(Usuario).filter_by(id=usuario_id).first()
    if not usuario:
        return jsonify({'error': 'Usuário não encontrado'}), 404
    
    # Obtém o status mais recente do usuário
    status_usuario = session.query(StatusUsuario).filter_by(usuario_id=usuario_id).order_by(StatusUsuario.status.desc()).first()
    status = status_usuario.status if status_usuario else 'Acesso Liberado'  # 'Ativo' ou qualquer valor padrão desejado

    usuario_data = {
        'id': usuario.id,
        'nome': usuario.nome,
        'email': usuario.email,
        'matricula': usuario.matricula,
        'setor': usuario.setor,                
        'usuario': usuario.usuario,
        'nivel_acesso': usuario.nivel_acesso,
        'pergunta_seguranca': usuario.pergunta_seguranca,
        'status': status
    }

    return jsonify(usuario_data)

@app.route('/usuarios/edit_usuario/<int:usuario_id>', methods=['GET', 'POST'])
@login_required
@permission_required('editar_usuario')
def edit_usuario(usuario_id):
    usuario = session.query(Usuario).filter_by(id=usuario_id).first()
    if not usuario:
        return jsonify({'error': 'Usuário não encontrado'}), 404

    usuario_logado, nivel_acesso, foto_perfil_url = get_current_user_name_and_access_level(flask_session, session)

    if request.method == 'POST':
        try:
            # Validação dos campos obrigatórios
            required_fields = ['primeiro_nome', 'sobrenome', 'email', 'matricula', 'setor', 'nivel_acesso', 'status']
            missing_fields = [field for field in required_fields if field not in request.form or not request.form[field].strip()]
            
            if missing_fields:
                return jsonify({
                    'error': 'Campos obrigatórios faltando',
                    'missing_fields': missing_fields
                }), 400

            # Valores antigos para log
            valores_antigos = {
                'primeiro_nome': usuario.primeiro_nome,
                'nome_meio': usuario.nome_meio,
                'sobrenome': usuario.sobrenome,
                'nome': usuario.nome,
                'email': usuario.email,
                'matricula': usuario.matricula,
                'setor': usuario.setor,
                'usuario': usuario.usuario,
                'nivel_acesso': usuario.nivel_acesso,
                'status': usuario.status.status if usuario.status else 'Acesso Liberado',  # Corrigido aqui
            }

            # Atualiza os campos básicos
            usuario.primeiro_nome = request.form['primeiro_nome'].strip()
            usuario.nome_meio = request.form.get('nome_meio', '').strip()
            usuario.sobrenome = request.form['sobrenome'].strip()
            
            # Gera o nome completo
            nome_completo = f"{usuario.primeiro_nome} {usuario.nome_meio} {usuario.sobrenome}".replace("  ", " ").strip()
            usuario.nome = nome_completo
            
            # Atualiza outros campos
            usuario.email = request.form['email'].strip()
            usuario.matricula = request.form['matricula'].strip()
            usuario.setor = request.form['setor'].strip()
            usuario.nivel_acesso = request.form['nivel_acesso']
            
            # Atualiza senha se fornecida
            nova_senha = request.form.get('senha', '').strip()
            if nova_senha:
                if len(nova_senha) < 6:
                    return jsonify({
                        'error': 'A senha deve ter pelo menos 6 caracteres',
                        'field': 'senha'
                    }), 400
                senha_criptografada = bcrypt.hashpw(nova_senha.encode('utf-8'), bcrypt.gensalt())                
                usuario.senha = senha_criptografada
            
            # Atualiza status - FORMA CORRETA PARA RELAÇÃO ONE-TO-ONE
            novo_status = request.form['status']
            if usuario.status:
                usuario.status.status = novo_status  # Corrigido aqui
            else:
                novo_status_usuario = StatusUsuario(status=novo_status, usuario=usuario)
                session.add(novo_status_usuario)
                # Não precisa atribuir a usuario.status pois o back_populates já cuida disso
            
            session.commit()

            # Registra mudanças no log
            mudancas = []
            for campo, valor_antigo in valores_antigos.items():
                valor_novo = getattr(usuario, campo) if campo != 'status' else (usuario.status.status if usuario.status else 'Acesso Liberado')  # Corrigido aqui
                if str(valor_novo) != str(valor_antigo):
                    mudancas.append(f"{campo}: {valor_antigo} -> {valor_novo}")

            if mudancas:
                data_atual = datetime.now()
                data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

                mensagem_log = (
                    f"(WEB) O Usuário {usuario_logado} - {nivel_acesso} realizou uma edição de usuário em {data_formatada}.\n"
                    f"Alterações:\n" + "\n".join(mudancas)
                )

                log_entry = EntradaLog(
                    mensagem=mensagem_log,
                    usuario=usuario_logado,
                    nivel_acesso=nivel_acesso,
                    tipo_log="WEB - Usuário - Edição",
                    data_geracao=data_formatada
                )

                session.add(log_entry)
                session.commit()

            return jsonify({
                'success': 'Usuário atualizado com sucesso!',
                'redirect': '/usuarios/usuarios_cadastrados'
            }), 200
            
        except Exception as e:
            session.rollback()
            return jsonify({
                'error': f'Erro ao atualizar usuário: {str(e)}'
            }), 500
        
    # Método GET
    current_user = get_current_user(flask_session, session)
    return render_template('opcoes/usuarios/edit_usuario.html', 
                         usuario=usuario, 
                         current_user=current_user)

@app.route('/api/usuarios/<int:usuario_id>', methods=['DELETE'])
@login_required
@permission_required('excluir_usuario')
def delete_usuario(usuario_id):
    # Busca o usuário a ser excluído
    usuario = session.query(Usuario).filter_by(id=usuario_id).first()
    if not usuario:
        return jsonify({'error': 'Usuário não encontrado'}), 404

    # Obtém o usuário atual (quem está realizando a exclusão) e seu nível de acesso
    usuario_atual, nivel_acesso, foto_perfil_url = get_current_user_name_and_access_level(flask_session, session)
    
    # Verifica se o usuário está tentando bloquear a si mesmo
    if usuario.usuario == usuario_atual:
        return jsonify({'error': 'Você não pode excluir seu próprio usuário!'}), 400

    data_atual = datetime.now()
    data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')
            
    mensagem_log = (
        f"(WEB) O Usuário {usuario_atual} ({nivel_acesso}) realizou uma exclusão de usuário em {data_formatada}.\n"
    )

    log_entry = EntradaLog(
        mensagem=mensagem_log,
        usuario=usuario_atual,
        nivel_acesso=nivel_acesso,
        tipo_log="WEB - Usuário - Exclusão",
        data_geracao=data_formatada
    )

    try:
        # Deletar todos os status do usuário
        session.query(StatusUsuario).filter_by(usuario_id=usuario.id).delete()
        
        # Deletar o usuário
        session.delete(usuario)
        
        # Adicionar a entrada de log e commit
        session.add(log_entry)
        session.commit()

        return jsonify({'message': 'Usuário apagado com sucesso'})
    except Exception as e:
        session.rollback()
        return jsonify({'error': str(e)}), 500

@app.route('/api/usuarios/<int:usuario_id>/bloquear', methods=['POST'])
@login_required
@permission_required('bloquear_usuario')
def bloquear_usuario(usuario_id):
    usuario_alvo = session.query(Usuario).filter_by(id=usuario_id).first()
    if not usuario_alvo:
        return jsonify({'error': 'Usuário não encontrado'}), 404

    usuario_logado, nivel_acesso, foto_perfil_url = get_current_user_name_and_access_level(flask_session, session)
    
    # Verifica se o usuário está tentando bloquear a si mesmo
    if usuario_alvo.usuario == usuario_logado:
        return jsonify({'error': 'Você não pode bloquear seu próprio usuário!'}), 400

    data_atual = datetime.now()
    data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')
    
    # Obter o motivo da requisição
    data = request.get_json()
    motivo = data.get('motivo', 'Motivo não informado')
    
    mensagem_log = (
        f"(WEB) O Usuário {usuario_logado} ({nivel_acesso}) bloqueou o acesso do usuário ID {usuario_id} em {data_formatada}.\n"
        f"Motivo: {motivo}\n"
    )

    log_entry = EntradaLog(
        mensagem=mensagem_log,
        usuario=usuario_logado,
        nivel_acesso=nivel_acesso,
        tipo_log="WEB - Usuário - Bloqueio",
        data_geracao=data_formatada
    )

    try:
        # Atualiza o status do usuário - FORMA CORRETA PARA RELAÇÃO ONE-TO-ONE
        novo_status = "Acesso Bloqueado"
        
        if usuario_alvo.status:  # Se já existe um status
            usuario_alvo.status.status = novo_status
        else:  # Se não existe, cria um novo
            novo_status_usuario = StatusUsuario(status=novo_status, usuario=usuario_alvo)
            session.add(novo_status_usuario)
            # Não precisa atribuir a usuario_alvo.status pois o back_populates já cuida disso

        # Adiciona a entrada de log e commit
        session.add(log_entry)
        session.commit()

        return jsonify({'success': 'Usuário bloqueado com sucesso'}), 200
    except Exception as e:
        session.rollback()
        return jsonify({'error': f'Erro ao bloquear usuário: {str(e)}'}), 500
    
@app.route('/api/usuarios/<int:usuario_id>/desbloquear', methods=['POST'])
@login_required
@permission_required('desbloquear_usuario')
def desbloquear_usuario(usuario_id):
    usuario_alvo = session.query(Usuario).filter_by(id=usuario_id).first()
    if not usuario_alvo:
        return jsonify({'error': 'Usuário não encontrado'}), 404

    usuario_logado, nivel_acesso, foto_perfil_url = get_current_user_name_and_access_level(flask_session, session)
    
    data_atual = datetime.now()
    data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')
    
    # Obter o motivo da requisição
    data = request.get_json()
    motivo = data.get('motivo', 'Motivo não informado')
    
    mensagem_log = (
        f"(WEB) O Usuário {usuario_logado} ({nivel_acesso}) desbloqueou o acesso do usuário ID {usuario_id} em {data_formatada}.\n"
        f"Motivo: {motivo}\n"
    )

    log_entry = EntradaLog(
        mensagem=mensagem_log,
        usuario=usuario_logado,
        nivel_acesso=nivel_acesso,
        tipo_log="WEB - Usuário - Desbloqueio",
        data_geracao=data_formatada
    )

    try:
        # Atualiza o status do usuário - FORMA CORRETA PARA RELAÇÃO ONE-TO-ONE
        novo_status = "Acesso Liberado"
        
        if usuario_alvo.status:  # Se já existe um status
            usuario_alvo.status.status = novo_status
        else:  # Se não existe, cria um novo
            novo_status_usuario = StatusUsuario(status=novo_status, usuario=usuario_alvo)
            session.add(novo_status_usuario)
            # Não precisa atribuir a usuario_alvo.status pois o back_populates já cuida disso

        # Reseta as tentativas de login
        usuario_alvo.tentativas_login = 0

        # Adiciona a entrada de log e commit
        session.add(log_entry)
        session.commit()

        return jsonify({'success': 'Usuário desbloqueado com sucesso'}), 200
    except Exception as e:
        session.rollback()
        return jsonify({'error': f'Erro ao desbloquear usuário: {str(e)}'}), 500
      
@app.route('/api/usuario', methods=['POST'])
@login_required
def registra_novo_usuario():
    try:
        # Obter dados do formulário
        primeiro_nome = request.form.get('primeiro_nome', '').strip()
        nome_meio = request.form.get('nome_meio', '').strip()
        sobrenome = request.form.get('sobrenome', '').strip()
        nome = request.form.get('nome', '').strip()
        email = request.form.get('email', '').strip()
        matricula = request.form.get('matricula', '').strip()
        setor = request.form.get('setor', '').strip()
        usuario_nome = request.form.get('usuario', '').strip()
        senha = request.form.get('senha', '').strip()
        senha_confirmacao = request.form.get('confirme_senha', '').strip()
        pergunta_seguranca = request.form.get('pergunta_seguranca', '').strip()
        resposta_seguranca = request.form.get('resposta_seguranca', '').strip()
        nivel_acesso = request.form.get('nivel_acesso', '').strip()

        # Validações básicas
        required_fields = {
            'primeiro_nome': primeiro_nome,
            'sobrenome': sobrenome,
            'nome': nome,
            'email': email,
            'matricula': matricula,
            'setor': setor,
            'usuario': usuario_nome,
            'senha': senha,
            'confirme_senha': senha_confirmacao,
            'pergunta_seguranca': pergunta_seguranca,
            'resposta_seguranca': resposta_seguranca,
            'nivel_acesso': nivel_acesso
        }

        missing_fields = [field for field, value in required_fields.items() if not value]
        if missing_fields:
            return jsonify({
                'error': 'Campos obrigatórios faltando',
                'missing_fields': missing_fields
            }), 400

        if senha != senha_confirmacao:
            return jsonify({'error': 'As senhas não correspondem.'}), 400

        # Validação da política de senhas
        if len(senha) < 8:
            return jsonify({'error': 'A senha deve ter pelo menos 8 caracteres.'}), 400

        if not re.search(r'[A-Z]', senha):
            return jsonify({'error': 'A senha deve conter pelo menos uma letra maiúscula.'}), 400

        if not re.search(r'[a-z]', senha):
            return jsonify({'error': 'A senha deve conter pelo menos uma letra minúscula.'}), 400

        if not re.search(r'[!@#$%^&*(),.?":{}|<>]', senha):
            return jsonify({'error': 'A senha deve conter pelo menos um símbolo.'}), 400

        if session.query(Usuario).filter_by(email=email).first() or session.query(Usuario).filter_by(usuario=usuario_nome).first():
            return jsonify({'error': 'Email ou usuário já estão em uso.'}), 400

        senha_criptografada = bcrypt.hashpw(senha.encode('utf-8'), bcrypt.gensalt())

        usuario_atual, nivel_acesso, foto_perfil_url = get_current_user_name_and_access_level(flask_session, session)

        novo_usuario = Usuario(
            primeiro_nome=primeiro_nome,
            nome_meio=nome_meio,
            sobrenome=sobrenome,
            nome=nome,
            email=email, 
            matricula=matricula, 
            setor=setor,                 
            usuario=usuario_nome, 
            senha=senha_criptografada, 
            nivel_acesso=nivel_acesso,
            pergunta_seguranca=pergunta_seguranca, 
            resposta_seguranca=resposta_seguranca,
            usuario_geracao=usuario_atual,
        )

        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

        mensagem_log = (
            f"(WEB) O Usuário {usuario_atual} ({nivel_acesso}) cadastrou um novo usuário em {data_atual.strftime('%d-%m-%Y %H:%M:%S')}.\n"
            f"Nome: {nome}\n"
            f"Email: {email}\n"
            f"Matricula: {matricula}\n"  
            f"Setor: {setor}\n"             
            f"Usuário: {usuario_nome}\n"
            f"Nível de Acesso: {nivel_acesso}\n"
        )

        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario_atual,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Usuário - Novo Cadastro",
            data_geracao=data_formatada
        )

        session.add(log_entry)
        session.add(novo_usuario)
        session.commit()

        return jsonify({'success': 'Usuário cadastrado com sucesso.'})

    except Exception as e:
        # Log do erro no servidor
        print(f"Erro ao cadastrar usuário: {str(e)}")
        # Retorna uma resposta JSON mesmo em caso de erro
        return jsonify({
            'error': 'Erro interno no servidor',
            'details': str(e)
        }), 500

@app.route('/usuarios/novo_usuario')
@login_required
@permission_required('novo_usuario')
def novo_usuario():
    current_user = get_current_user(flask_session, session)
    return render_template('opcoes/usuarios/novo_usuario.html', current_user=current_user)

# Métodos de Tipos de Equipamento ###################################################################

@app.route('/novo_tipo_equipamento')
@login_required
@permission_required('novo_tipo_equipamento')
def novo_tipo_equipamento():
    current_user = get_current_user(flask_session, session)   
    return render_template('opcoes/tipos_equipamentos/cadastro_tipos_equipamentos.html', current_user=current_user)

@app.route('/api/tipo_equipamento', methods=['POST'])
@login_required
@permission_required('novo_tipo_equipamento')
def registra_novo_tipo_equipamento():
    try:
        # Obter dados do formulário
        nome_tipo = request.form.get('nome_tipo', '').strip()
        descricao = request.form.get('descricao', '').strip()
        validade_inspecao = request.form.get('validade_inspecao')
        checklist_items = request.form.get('checklist_items')
        
        # Validações básicas
        if not nome_tipo:
            return jsonify({'error': 'O nome do tipo de equipamento é obrigatório.'}), 400
            
        if not validade_inspecao or not validade_inspecao.isdigit() or int(validade_inspecao) <= 0:
            return jsonify({'error': 'A validade da inspeção deve ser um número positivo de dias.'}), 400
            
        # Verifica se o tipo já existe
        if session.query(TiposEquipamentos).filter(func.lower(TiposEquipamentos.nome) == func.lower(nome_tipo)).first():
            return jsonify({'error': 'Este tipo de equipamento já está cadastrado.'}), 400
            
        # Processa os itens do checklist
        try:
            checklist_data = json.loads(checklist_items) if checklist_items else []
            if not checklist_data:
                return jsonify({'error': 'É necessário adicionar pelo menos um item ao checklist.'}), 400
        except json.JSONDecodeError:
            return jsonify({'error': 'Formato inválido para os itens do checklist.'}), 400
            
        # Obtém informações do usuário atual
        usuario_atual, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        
        # Cria novo tipo de equipamento
        novo_tipo = TiposEquipamentos(
            nome=nome_tipo,
            descricao=descricao,
            validade_inspecao=int(validade_inspecao),
            usuario_geracao=usuario_atual,
            status=True
        )
        
        session.add(novo_tipo)
        session.flush()  # Para obter o ID do novo tipo
        
        # Adiciona os itens do checklist
        for index, item in enumerate(checklist_data, start=1):
            checklist_item = ChecklistTipo(
                nome_item=item['nome_item'],
                obrigatorio=item.get('obrigatorio', True),
                ordem=index,
                tipo_id=novo_tipo.id
            )
            session.add(checklist_item)
        
        # Cria log da operação
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')
        
        mensagem_log = (
            f"(WEB) O Usuário {usuario_atual} ({nivel_acesso}) cadastrou um novo tipo de equipamento em {data_formatada}.\n"
            f"Tipo: {nome_tipo}\n"
            f"Descrição: {descricao}\n"
            f"Itens do Checklist: {len(checklist_data)}\n"
        )
        
        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario_atual,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Tipo Equipamento - Novo Cadastro",
            data_geracao=data_formatada
        )
        
        session.add(log_entry)
        session.commit()
        
        return jsonify({'success': 'Tipo de equipamento cadastrado com sucesso.'})
        
    except Exception as e:
        session.rollback()
        print(f"Erro ao cadastrar tipo de equipamento: {str(e)}")
        return jsonify({
            'error': 'Erro interno no servidor',
            'details': str(e)
        }), 500

@app.route('/tipos_equipamentos_cadastrados')
@login_required
@permission_required('tipos_equipamentos_cadastrados')
def tipos_equipamentos_cadastrados():
    current_user = get_current_user(flask_session, session)
    allowed_menus = get_allowed_menus_for_user(current_user)    
    return render_template('opcoes/tipos_equipamentos/tipos_equipamentos_cadastrados.html', 
                         current_user=current_user, 
                         allowed_menus=allowed_menus)

@app.route('/api/tipos_equipamentos_ativos', methods=['GET'])
@login_required
def listar_tipos_equipamentos_ativos():
    try:
        search_term = request.args.get('search', '').strip()
        
        query = session.query(TiposEquipamentos).filter(TiposEquipamentos.status == True)
        
        if search_term:
            query = query.filter(
                or_(
                    TiposEquipamentos.nome.ilike(f'%{search_term}%'),
                    TiposEquipamentos.descricao.ilike(f'%{search_term}%')
                )
            )
        
        tipos = query.options(joinedload(TiposEquipamentos.checklists)).order_by(TiposEquipamentos.nome).all()
        
        tipos_data = []
        for tipo in tipos:
            tipos_data.append({
                'id': tipo.id,
                'nome': tipo.nome,
                'descricao': tipo.descricao,
                'validade_inspecao': tipo.validade_inspecao,
                'status': tipo.status,
                'data_geracao': tipo.data_geracao.strftime('%d/%m/%Y %H:%M') if tipo.data_geracao else '',
                'usuario_geracao': tipo.usuario_geracao,
                'checklists': [{
                    'id': item.id,
                    'nome_item': item.nome_item,
                    'obrigatorio': item.obrigatorio,
                    'ordem': item.ordem
                } for item in tipo.checklists]
            })
        
        return jsonify(tipos_data)
        
    except Exception as e:
        print(f"Erro ao listar tipos de equipamentos: {str(e)}")
        return jsonify({'error': 'Erro interno no servidor'}), 500

@app.route('/api/tipos_equipamentos', methods=['GET'])
@login_required
def listar_tipos_equipamentos():
    try:
        draw = int(request.args.get('draw', 1))
        start = int(request.args.get('start', 0))
        length = int(request.args.get('length', 10))
        search_value = request.args.get('search[value]', '').strip()
        order_col = request.args.get('order[0][column]', '0')
        order_dir = request.args.get('order[0][dir]', 'asc')

        status = request.args.get('status')
        validade_min = request.args.get('validade_min', type=int)
        validade_max = request.args.get('validade_max', type=int)

        col_map = {
            '0': TiposEquipamentos.nome,
            '1': TiposEquipamentos.status,
            '3': TiposEquipamentos.validade_inspecao,
            '4': TiposEquipamentos.data_geracao,
        }

        query = session.query(TiposEquipamentos).filter()

        # Filtro por busca
        if search_value:
            query = query.filter(
                or_(
                    TiposEquipamentos.nome.ilike(f'%{search_value}%'),
                    TiposEquipamentos.descricao.ilike(f'%{search_value}%')
                )
            )

        # Filtros extras
        if status in ['true', 'false']:
            query = query.filter(TiposEquipamentos.status == (status == 'true'))

        if validade_min is not None:
            query = query.filter(TiposEquipamentos.validade_inspecao >= validade_min)

        if validade_max is not None:
            query = query.filter(TiposEquipamentos.validade_inspecao <= validade_max)

        total_records = session.query(TiposEquipamentos).count()
        filtered_records = query.count()

        # Ordenação
        sort_col = col_map.get(order_col, TiposEquipamentos.nome)
        sort_func = asc if order_dir == 'asc' else desc
        query = query.order_by(sort_func(sort_col))

        # Paginação
        tipos = query.offset(start).limit(length).options(joinedload(TiposEquipamentos.checklists)).all()

        # Construção da resposta
        tipos_data = []
        for tipo in tipos:
            tipos_data.append({
                'id': tipo.id,
                'nome': tipo.nome,
                'descricao': tipo.descricao,
                'validade_inspecao': tipo.validade_inspecao,
                'status': tipo.status,
                'data_geracao': tipo.data_geracao.strftime('%d/%m/%Y %H:%M') if tipo.data_geracao else '',
                'usuario_geracao': tipo.usuario_geracao,
                'checklists': [{
                    'id': item.id,
                    'nome_item': item.nome_item,
                    'obrigatorio': item.obrigatorio,
                    'ordem': item.ordem
                } for item in tipo.checklists]
            })

        return jsonify({
            'draw': draw,
            'recordsTotal': total_records,
            'recordsFiltered': filtered_records,
            'data': tipos_data
        })

    except Exception as e:
        print(f"Erro ao listar tipos de equipamentos: {str(e)}")
        return jsonify({'error': 'Erro interno no servidor'}), 500
    
@app.route('/api/tipos_equipamentos/<int:tipo_id>', methods=['GET'])
@login_required
def obter_tipo_equipamento(tipo_id):
    try:
        tipo = session.get(TiposEquipamentos, tipo_id, options=[joinedload(TiposEquipamentos.checklists)])
        
        if not tipo:
            return jsonify({'error': 'Tipo de equipamento não encontrado'}), 404
            
        tipo_data = {
            'id': tipo.id,
            'nome': tipo.nome,
            'descricao': tipo.descricao,
            'validade_inspecao': tipo.validade_inspecao,
            'status': tipo.status,
            'data_geracao': tipo.data_geracao.strftime('%d/%m/%Y %H:%M') if tipo.data_geracao else '',
            'usuario_geracao': tipo.usuario_geracao,
            'checklists': [{
                'id': item.id,
                'nome_item': item.nome_item,
                'obrigatorio': item.obrigatorio,
                'ordem': item.ordem
            } for item in tipo.checklists]
        }
        
        return jsonify(tipo_data)
        
    except Exception as e:
        print(f"Erro ao obter tipo de equipamento: {str(e)}")
        return jsonify({'error': 'Erro interno no servidor'}), 500
    
@app.route('/api/tipos_equipamentos/<int:tipo_id>/equipamentos', methods=['GET'])
@login_required
def get_equipamentos_tipo(tipo_id):
    try:
        tipo = session.query(TiposEquipamentos).get(tipo_id)
        if not tipo:
            return jsonify({'error': 'Tipo de equipamento não encontrado'}), 404
        
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 15))

        query = session.query(Equipamento).filter(Equipamento.tipo == tipo.nome)

        total_equipamentos = query.count()
        equipamentos = query.offset((page - 1) * limit).limit(limit).all()

        equipamentos_data = [
            {
                'id': equipamento.id,
                'tag_equipamento': equipamento.tag_equipamento,
                'localizacao': equipamento.localizacao,
                'tipo': equipamento.tipo,
                'status': equipamento.status,
                'data_fabricacao': equipamento.data_fabricacao.strftime('%d/%m/%Y') if equipamento.data_fabricacao else '',
                'data_validade': equipamento.data_validade.strftime('%d/%m/%Y') if equipamento.data_validade else ''
            }
            for equipamento in equipamentos
        ]

        return jsonify({
            'data': equipamentos_data,
            'total': total_equipamentos,
            'tipo': {
                'id': tipo.id,
                'nome': tipo.nome
            }
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/tipos_equipamentos/<int:tipo_id>/status', methods=['PUT'])
@login_required
@permission_required('gerenciar_tipos_equipamentos')
def alterar_status_tipo_equipamento(tipo_id):
    try:
        tipo = session.get(TiposEquipamentos, tipo_id)
        if not tipo:
            return jsonify({'error': 'Tipo de equipamento não encontrado'}), 404

        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'Requisição inválida. Esperado JSON no corpo da requisição.'}), 400
        
        motivo = data.get('motivo', 'Motivo não informado')
        
        if 'status' not in data:
            return jsonify({'error': 'Status não informado'}), 400

        novo_status = data['status']

            
        # Verificar se há equipamentos cadastrados se estiver desativando
        if not novo_status:
            equipamentos_count = session.query(Equipamento).filter(
                func.lower(Equipamento.tipo) == tipo.nome.lower()
            ).count()

            if equipamentos_count > 0:
                return jsonify({
                    'error': f'Não é possível desativar este tipo pois existem {equipamentos_count} equipamentos cadastrados'
                }), 400
        
        # Registrar alteração no log
        usuario_atual, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')
        
        mensagem_log = (
            f"(WEB) O Usuário {usuario_atual} ({nivel_acesso}) alterou o status do tipo de equipamento {tipo.nome} "
            f"para {'ATIVO' if novo_status else 'INATIVO'} em {data_formatada}.\n"
            f"Motivo: {motivo}\n"
        )
        
        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario_atual,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Tipo Equipamento - Alteração de Status",
            data_geracao=data_formatada
        )
        
        tipo.status = novo_status
        session.add(log_entry)
        session.commit()
        
        return jsonify({
            'success': f'Tipo de equipamento {tipo.nome} {"ativado" if novo_status else "desativado"} com sucesso.'
        })

    except Exception as e:
        session.rollback()
        print(f"Erro ao alterar status do tipo de equipamento: {str(e)}")
        return jsonify({'error': 'Erro interno no servidor'}), 500

@app.route('/editar_tipo_equipamento/<int:tipo_id>', methods=['GET', 'POST'])
@login_required
@permission_required('editar_tipo_equipamento')
def editar_tipo_equipamento(tipo_id):
    tipo = session.query(TiposEquipamentos).options(joinedload(TiposEquipamentos.checklists)).get(tipo_id)
    if not tipo:
        return jsonify({'error': 'Tipo de equipamento não encontrado'}), 404

    usuario_logado, nivel_acesso, foto_perfil_url = get_current_user_name_and_access_level(flask_session, session)

    if request.method == 'POST':
        try:
            # Validação dos campos obrigatórios
            required_fields = ['nome', 'validade_inspecao']
            missing_fields = [field for field in required_fields if field not in request.form or not request.form[field].strip()]
            
            if missing_fields:
                return jsonify({
                    'error': 'Campos obrigatórios faltando',
                    'missing_fields': missing_fields
                }), 400

            # Valida validade_inspecao
            try:
                validade_inspecao = int(request.form['validade_inspecao'])
                if validade_inspecao <= 0:
                    return jsonify({'error': 'A validade da inspeção deve ser um número positivo de dias.'}), 400
            except ValueError:
                return jsonify({'error': 'A validade da inspeção deve ser um número válido.'}), 400

            # Valores antigos para log
            valores_antigos = {
                'nome': tipo.nome,
                'descricao': tipo.descricao,
                'validade_inspecao': tipo.validade_inspecao,
                'status': tipo.status
            }

            # Atualiza os campos
            tipo.nome = request.form['nome'].strip()
            tipo.descricao = request.form.get('descricao', '').strip()
            tipo.validade_inspecao = validade_inspecao
            tipo.status = request.form.get('status', 'off') == 'on'
            
            # Atualiza itens do checklist
            checklist_items = []
            for key, value in request.form.items():
                if key.startswith('checklist_items[') and key.endswith('][nome_item]'):
                    prefix = key.split('[')[1].split(']')[0]
                    nome_item = value.strip()
                    obrigatorio = request.form.get(f'checklist_items[{prefix}][obrigatorio]', 'off') == 'on'
                    
                    if nome_item:  # Só adiciona se tiver nome
                        checklist_items.append({
                            'nome_item': nome_item,
                            'obrigatorio': obrigatorio
                        })
            
            # Remove itens antigos e adiciona novos
            tipo.checklists = []
            for idx, item in enumerate(checklist_items, start=1):
                checklist_item = ChecklistTipo(
                    nome_item=item['nome_item'],
                    obrigatorio=item['obrigatorio'],
                    ordem=idx,
                    tipo_id=tipo.id
                )
                tipo.checklists.append(checklist_item)
            
            session.commit()

            # Registra mudanças no log
            mudancas = []
            for campo, valor_antigo in valores_antigos.items():
                valor_novo = getattr(tipo, campo)
                if str(valor_novo) != str(valor_antigo):
                    mudancas.append(f"{campo}: {valor_antigo} -> {valor_novo}")
            
            # Adiciona mudanças no checklist
            mudancas_checklist = []
            if len(checklist_items) != len(valores_antigos.get('checklists', [])):
                mudancas_checklist.append(f"Itens do checklist alterados: {len(valores_antigos.get('checklists', []))} -> {len(checklist_items)}")

            if mudancas or mudancas_checklist:
                data_atual = datetime.now()
                data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

                mensagem_log = (
                    f"(WEB) O Usuário {usuario_logado} - {nivel_acesso} realizou uma edição de tipo de equipamento em {data_formatada}.\n"
                    f"Alterações:\n" + "\n".join(mudancas + mudancas_checklist)
                )

                log_entry = EntradaLog(
                    mensagem=mensagem_log,
                    usuario=usuario_logado,
                    nivel_acesso=nivel_acesso,
                    tipo_log="WEB - Tipo Equipamento - Edição",
                    data_geracao=data_formatada
                )

                session.add(log_entry)
                session.commit()

            return jsonify({
                'success': 'Tipo de equipamento atualizado com sucesso!',
                'redirect': '/tipos_equipamentos_cadastrados'
            }), 200
            
        except Exception as e:
            session.rollback()
            return jsonify({
                'error': f'Erro ao atualizar tipo de equipamento: {str(e)}'
            }), 500
        
    # Método GET
    current_user = get_current_user(flask_session, session)
    checklists_json = json.dumps([
        {
            'id': item.id,
            'nome_item': item.nome_item,
            'obrigatorio': item.obrigatorio
        } for item in tipo.checklists
    ])

    return render_template('opcoes/tipos_equipamentos/editar_tipo_equipamento.html', 
                        tipo=tipo, 
                        checklists_json=Markup(checklists_json),
                        current_user=current_user)

# Métodos de Setores ###################################################################

@app.route('/novo_setor')
@login_required
@permission_required('novo_setor')
def novo_setor():
    current_user = get_current_user(flask_session, session)
    return render_template('opcoes/setores/cadastro_setores.html', current_user=current_user)

@app.route('/api/setor', methods=['POST'])
@login_required
def registra_novo_setor():
    try:
        # Obter dados do formulário
        nome_setor = request.form.get('nome_setor', '').strip()
        centro_custo = request.form.get('centro_custo', '').strip()

        # Validações básicas
        required_fields = {
            'nome_setor': nome_setor,
            'centro_custo': centro_custo
        }

        missing_fields = [field for field, value in required_fields.items() if not value]
        if missing_fields:
            return jsonify({
                'error': 'Campos obrigatórios faltando',
                'missing_fields': missing_fields
            }), 400

        # Verifica se o setor já existe
        if session.query(Setores).filter_by(nome_setor=nome_setor).first():
            return jsonify({'error': 'Este setor já está cadastrado.'}), 400

        # Obtém informações do usuário atual
        usuario_atual, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)

        # Cria novo setor
        novo_setor = Setores(
            nome_setor=nome_setor,
            centro_custo=centro_custo,
            data_geracao=datetime.now(),
            usuario_geracao=usuario_atual
        )

        # Cria log da operação
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

        mensagem_log = (
            f"(WEB) O Usuário {usuario_atual} ({nivel_acesso}) cadastrou um novo setor em {data_atual.strftime('%d-%m-%Y %H:%M:%S')}.\n"
            f"Setor: {nome_setor}\n"
            f"Centro de Custo: {centro_custo}\n"
        )

        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario_atual,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Setor - Novo Cadastro",
            data_geracao=data_formatada
        )

        session.add(log_entry)
        session.add(novo_setor)
        session.commit()

        return jsonify({'success': 'Setor cadastrado com sucesso.'})

    except Exception as e:
        # Log do erro no servidor
        print(f"Erro ao cadastrar setor: {str(e)}")
        # Retorna uma resposta JSON mesmo em caso de erro
        return jsonify({
            'error': 'Erro interno no servidor',
            'details': str(e)
        }), 500

@app.route('/setores_cadastrados')
@login_required
@permission_required('setores_cadastrados')
def setores_cadastrados():
    current_user = get_current_user(flask_session, session)
    allowed_menus = get_allowed_menus_for_user(current_user)    
    return render_template('opcoes/setores/setores_cadastrados.html', current_user=current_user, allowed_menus=allowed_menus)

@app.route('/api/setores', methods=['GET'])
@login_required
def listar_setores():
    try:
        # Parâmetros do DataTables
        draw = request.args.get('draw', type=int)
        start = request.args.get('start', type=int)
        length = request.args.get('length', type=int)
        search_value = request.args.get('search[value]', '').strip()
        
        # Filtros adicionais
        centro_custo_filter = request.args.get('centro_custo', '').strip()
        
        # Ordenação
        order_column_index = request.args.get('order[0][column]', default=0, type=int)
        order_direction = request.args.get('order[0][dir]', 'asc')
        
        # Mapeamento de colunas
        columns = [
            'nome_setor',
            'centro_custo',
            'data_geracao',
            'id'  # Coluna de ações
        ]
        order_column = columns[order_column_index] if order_column_index < len(columns) else 'nome_setor'

        # Construção da query
        query = session.query(Setores)

        # Aplicar filtro de pesquisa geral
        if search_value:
            query = query.filter(
                or_(
                    Setores.nome_setor.ilike(f'%{search_value}%'),
                    Setores.centro_custo.ilike(f'%{search_value}%'),
                    Setores.usuario_geracao.ilike(f'%{search_value}%')
                )
            )

        # Aplicar filtros adicionais
        if centro_custo_filter:
            query = query.filter(Setores.centro_custo.ilike(f'%{centro_custo_filter}%'))
            
        # Ordenação
        if order_direction == 'desc':
            query = query.order_by(desc(getattr(Setores, order_column)))
        else:
            query = query.order_by(getattr(Setores, order_column))

        # Contagem total antes da paginação
        total_records = query.count()

        # Paginação
        setores = query.offset(start).limit(length).all()

        # Formatar dados
        setores_data = []
        for setor in setores:
            try:
                data_formatada = datetime.strptime(setor.data_geracao, '%d-%m-%Y %H:%M:%S').strftime('%d/%m/%Y') if setor.data_geracao else ''
            except ValueError:
                data_formatada = setor.data_geracao  # Em caso de erro, exibe como está

            setores_data.append({ 
                'id': setor.id,
                'nome_setor': setor.nome_setor,
                'centro_custo': setor.centro_custo,
                'data_geracao': data_formatada,
                'usuario_geracao': setor.usuario_geracao,
                'DT_RowId': f'setor_{setor.id}'
            })

        return jsonify({
            'draw': draw,
            'recordsTotal': total_records,
            'recordsFiltered': total_records,
            'data': setores_data
        })

    except Exception as e:
        print(f"Erro ao listar setores: {str(e)}")
        return jsonify({'error': 'Erro interno no servidor'}), 500

@app.route('/api/setores_selecao', methods=['GET'])
def listar_setores_tela_login():
    try:
        # Verifica se é uma requisição do DataTables
        is_datatable = 'draw' in request.args
        
        if is_datatable:
            # Código existente para DataTables
            draw = request.args.get('draw', type=int)
            start = request.args.get('start', type=int)
            length = request.args.get('length', type=int)
            search_value = request.args.get('search[value]', '').strip()
            order_column_index = request.args.get('order[0][column]', default=0, type=int)
            order_direction = request.args.get('order[0][dir]', 'asc')
            
            columns = ['nome_setor', 'centro_custo', 'data_geracao', 'id']
            order_column = columns[order_column_index] if order_column_index < len(columns) else 'nome_setor'
            
            query = session.query(Setores)
            
            if search_value:
                query = query.filter(
                    or_(
                        Setores.nome_setor.ilike(f'%{search_value}%'),
                        Setores.centro_custo.ilike(f'%{search_value}%'),
                        Setores.usuario_geracao.ilike(f'%{search_value}%')
                    )
                )
            
            if order_direction == 'desc':
                query = query.order_by(desc(getattr(Setores, order_column)))
            else:
                query = query.order_by(getattr(Setores, order_column))
            
            total_records = query.count()
            setores = query.offset(start).limit(length).all()
            
            setores_data = []
            for setor in setores:
                try:
                    data_formatada = datetime.strptime(setor.data_geracao, '%d-%m-%Y %H:%M:%S').strftime('%d/%m/%Y') if setor.data_geracao else ''
                except ValueError:
                    data_formatada = setor.data_geracao
                
                setores_data.append({ 
                    'id': setor.id,
                    'nome_setor': setor.nome_setor,
                    'centro_custo': setor.centro_custo,
                    'data_geracao': data_formatada,
                    'usuario_geracao': setor.usuario_geracao,
                    'DT_RowId': f'setor_{setor.id}'
                })
            
            return jsonify({
                'draw': draw,
                'recordsTotal': total_records,
                'recordsFiltered': total_records,
                'data': setores_data
            })
        else:
            # Retorno simplificado para o frontend de cadastro
            setores = session.query(Setores).order_by(Setores.nome_setor).all()
            setores_data = [{
                'id': setor.id,
                'nome_setor': setor.nome_setor,
                'centro_custo': setor.centro_custo
            } for setor in setores]
            
            return jsonify(setores_data)

    except Exception as e:
        print(f"Erro ao listar setores: {str(e)}")
        return jsonify({'error': 'Erro interno no servidor'}), 500

@app.route('/api/setores/<int:setor_id>/inspecoes', methods=['GET'])
@login_required
def get_historico_inspecoes_setor(setor_id):
    try:
        setor = session.get(Setores, setor_id)
        if not setor:
            return jsonify({'error': 'Setor não encontrado'}), 404


        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 15))

        query = (
            session.query(Registro, Equipamento)
            .outerjoin(Equipamento, Registro.equipamento_id == Equipamento.id)
            .filter(Registro.localizacao == setor.nome_setor)
        )

        total_inspecoes = query.count()
        resultados = query.offset((page - 1) * limit).limit(limit).all()

        inspecoes_data = [
            {
                'id': registro.id,
                'numero_inspecao': registro.numero_inspecao,
                'equipamento_id': registro.equipamento_id,
                'tag_equipamento': equipamento.tag_equipamento if equipamento else 'N/A',
                'tipo_equipamento': equipamento.tipo if equipamento else 'N/A',
                'localizacao': registro.localizacao,
                'responsavel': registro.responsavel,
                'criticidade_inspecao': registro.criticidade_inspecao,                  
                'data_inspecao': registro.data_inspecao.strftime('%d/%m/%Y') if registro.data_inspecao else '',
                'data_validade': registro.data_validade_inspecao.strftime('%d/%m/%Y') if registro.data_validade_inspecao else '',
                'status_inspecao': registro.status_inspecao,
                'status_equipamento': registro.status_equipamento,
                'observacoes': registro.observacoes
            }
            for registro, equipamento in resultados
        ]

        return jsonify({
            'data': inspecoes_data,
            'total': total_inspecoes,
            'setor': {
                'id': setor.id,
                'nome': setor.nome_setor
            }
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/setores/<int:setor_id>/equipamentos', methods=['GET'])
@login_required
def get_equipamentos_setor(setor_id):
    try:
        setor = session.query(Setores).filter_by(id=setor_id).first()
        if not setor:
            return jsonify({'error': 'Setor não encontrado'}), 404
        
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 15))

        query = session.query(Equipamento).filter(Equipamento.localizacao == setor.nome_setor)

        total_equipamentos = query.count()
        equipamentos = query.offset((page - 1) * limit).limit(limit).all()

        equipamentos_data = [
            {
                'id': equipamento.id,
                'tag_equipamento': equipamento.tag_equipamento,
                'localizacao': equipamento.localizacao,
                'tipo': equipamento.tipo,
                'classe': equipamento.classe,
                'data_fabricacao': equipamento.data_fabricacao.strftime('%d-%m-%Y'),
                'peso': equipamento.peso,
                'teste_hidrostatico_n2': equipamento.teste_hidrostatico_n2.strftime('%d-%m-%Y') if equipamento.teste_hidrostatico_n2 else None,
                'teste_hidrostatico_n3': equipamento.teste_hidrostatico_n3.strftime('%d-%m-%Y') if equipamento.teste_hidrostatico_n3 else None,
                'data_validade': equipamento.data_validade.strftime('%d-%m-%Y'),
                'status': equipamento.status,
                'codigo_barras': equipamento.codigo_barras,
                'data_ultima_inspecao': equipamento.data_ultima_inspecao.strftime('%d-%m-%Y') if equipamento.data_ultima_inspecao else None,
                'data_proxima_inspecao': equipamento.data_proxima_inspecao.strftime('%d-%m-%Y') if equipamento.data_proxima_inspecao else None
            }
            for equipamento in equipamentos
        ]

        return jsonify({
            'data': equipamentos_data,
            'total': total_equipamentos,
            'setor': {
                'id': setor.id,
                'nome': setor.nome_setor
            }
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/setores/<int:setor_id>/responsaveis', methods=['GET'])
@login_required
def get_responsaveis_setor(setor_id):
    try:
        print(f"Buscando responsáveis para o setor ID: {setor_id}")  # Log de depuração
        
        setor = session.query(Setores).filter_by(id=setor_id).first()
        if not setor:
            print("Setor não encontrado")  # Log de depuração
            return jsonify({'error': 'Setor não encontrado'}), 404
        
        page = int(request.args.get('page', 1))
        limit = int(request.args.get('limit', 15))

        query = session.query(Usuario).filter(Usuario.setor == setor.nome_setor)

        total_usuarios = query.count()
        usuarios = query.offset((page - 1) * limit).limit(limit).all()

        usuarios_data = [
            {
                'id': usuario.id,
                'nome': f"{usuario.primeiro_nome} {usuario.sobrenome}",
                'nome_completo': usuario.nome,
                'matricula': usuario.matricula,       
                'setor': usuario.setor,             
                'email': usuario.email,
                'usuario': usuario.usuario,
                'nivel_acesso': usuario.nivel_acesso,
            }
            for usuario in usuarios
        ]

        return jsonify({
            'data': usuarios_data,
            'total': total_usuarios
        })

    except Exception as e:
        print(f"Erro: {str(e)}")  # Log de depuração
        return jsonify({'error': str(e)}), 500
    
@app.route('/api/setores/<int:setor_id>', methods=['GET'])
@login_required
def obter_setor(setor_id):
    try:
        setor = session.get(Setores, setor_id)
        
        if not setor:
            return jsonify({'error': 'Setor não encontrado'}), 404
            
        setor_data = {
            'id': setor.id,
            'nome_setor': setor.nome_setor,
            'centro_custo': setor.centro_custo,
            'data_geracao': setor.data_geracao,
            'usuario_geracao': setor.usuario_geracao
        }
        
        return jsonify(setor_data)
        
    except Exception as e:
        print(f"Erro ao obter setor: {str(e)}")
        return jsonify({'error': 'Erro interno no servidor'}), 500

@app.route('/api/setores/<int:setor_id>', methods=['DELETE'])
@login_required
@permission_required('excluir_setor')
def excluir_setor(setor_id):
    try:
        setor = session.get(Setores, setor_id)
        
        if not setor:
            return jsonify({'error': 'Setor não encontrado'}), 404
            
        # Verificar se o setor está em uso antes de excluir
        em_uso = session.query(Usuario).filter_by(setor=setor.nome_setor).first()
        if em_uso:
            return jsonify({
                'error': 'Este setor não pode ser excluído pois está vinculado a um ou mais usuários'
            }), 400
            
        # Obter dados para o log
        usuario_atual, nivel_acesso, _ = get_current_user_name_and_access_level(flask_session, session)
        motivo = request.json.get('motivo', 'Motivo não informado')
        
        # Criar log da operação
        data_atual = datetime.now()
        data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')
        
        mensagem_log = (
            f"(WEB) O Usuário {usuario_atual} ({nivel_acesso}) excluiu o setor {setor.nome_setor} em {data_atual.strftime('%d-%m-%Y %H:%M:%S')}.\n"
            f"Motivo: {motivo}\n"
            f"Centro de Custo: {setor.centro_custo}\n"
        )
        
        log_entry = EntradaLog(
            mensagem=mensagem_log,
            usuario=usuario_atual,
            nivel_acesso=nivel_acesso,
            tipo_log="WEB - Setor - Exclusão",
            data_geracao=data_formatada
        )
        
        session.add(log_entry)
        session.delete(setor)
        session.commit()
        
        return jsonify({'message': 'Setor excluído com sucesso'})
        
    except Exception as e:
        session.rollback()
        print(f"Erro ao excluir setor: {str(e)}")
        return jsonify({'error': 'Erro interno no servidor'}), 500

@app.route('/editar_setor/<int:setor_id>', methods=['GET', 'POST'])
@login_required
@permission_required('editar_setor')
def editar_setor(setor_id):
    setor = session.query(Setores).filter_by(id=setor_id).first()
    if not setor:
        return jsonify({'error': 'Setor não encontrado'}), 404

    usuario_logado, nivel_acesso, foto_perfil_url = get_current_user_name_and_access_level(flask_session, session)

    if request.method == 'POST':
        try:
            # Validação dos campos obrigatórios
            required_fields = ['nome_setor', 'centro_custo']
            missing_fields = [field for field in required_fields if field not in request.form or not request.form[field].strip()]
            
            if missing_fields:
                return jsonify({
                    'error': 'Campos obrigatórios faltando',
                    'missing_fields': missing_fields
                }), 400

            # Valores antigos para log
            valores_antigos = {
                'nome_setor': setor.nome_setor,
                'centro_custo': setor.centro_custo
            }

            # Atualiza os campos
            setor.nome_setor = request.form['nome_setor'].strip()
            setor.centro_custo = request.form['centro_custo'].strip()
            
            session.commit()

            # Registra mudanças no log
            mudancas = []
            for campo, valor_antigo in valores_antigos.items():
                valor_novo = getattr(setor, campo)
                if str(valor_novo) != str(valor_antigo):
                    mudancas.append(f"{campo}: {valor_antigo} -> {valor_novo}")

            if mudancas:
                data_atual = datetime.now()
                data_formatada = data_atual.strftime('%d-%m-%Y %H:%M:%S')

                mensagem_log = (
                    f"(WEB) O Usuário {usuario_logado} - {nivel_acesso} realizou uma edição de setor em {data_formatada}.\n"
                    f"Alterações:\n" + "\n".join(mudancas)
                )

                log_entry = EntradaLog(
                    mensagem=mensagem_log,
                    usuario=usuario_logado,
                    nivel_acesso=nivel_acesso,
                    tipo_log="WEB - Setor - Edição",
                    data_geracao=data_formatada
                )

                session.add(log_entry)
                session.commit()

            return jsonify({
                'success': 'Setor atualizado com sucesso!',
                'redirect': '/setores_cadastrados'
            }), 200
            
        except Exception as e:
            session.rollback()
            return jsonify({
                'error': f'Erro ao atualizar setor: {str(e)}'
            }), 500
        
    # Método GET
    current_user = get_current_user(flask_session, session)
    return render_template('opcoes/setores/editar_setor.html', 
                         setor=setor, 
                         current_user=current_user)

# Métodos de Notificações ###################################################################

def obter_notificacoes(db_session):
    registros = db_session.query(Registro).filter(Registro.status_inspecao == "Pendente").all()
    today = date.today()
    limite_vencimento = today + timedelta(days=30)
    notificacoes = []

    for registro in registros:
        if registro.data_validade_inspecao is not None:
            dias_para_vencer = (registro.data_validade_inspecao - today).days

            if registro.data_validade_inspecao < today:
                mensagem = f"A Inspeção do(a) {registro.tipo} localizado(a) em {registro.localizacao} está vencido(a) desde {registro.data_validade_inspecao}"
                notificacoes.append({
                    "mensagem": mensagem,
                    "tipo": "vencido",
                    "localizacao": registro.localizacao,
                    "tipo": registro.tipo,
                    "data_validade_inspecao": registro.data_validade_inspecao,
                    "data_inspecao": registro.data_inspecao
                })
            elif today < registro.data_validade_inspecao <= limite_vencimento:
                mensagem = f"A Inspeção do(a) {registro.tipo} localizado(a) em {registro.localizacao} vence em {registro.data_validade_inspecao}. Faltam aproximadamente {dias_para_vencer} dias."
                notificacoes.append({
                    "mensagem": mensagem,
                    "tipo": "proximo",
                    "localizacao": registro.localizacao,
                    "tipo": registro.tipo,
                    "data_validade_inspecao": registro.data_validade_inspecao,
                    "data_inspecao": registro.data_inspecao
                })

        if registro.data_inspecao is not None:
            dias_desde_manutencao = (today - registro.data_inspecao).days
            if dias_desde_manutencao > 365:  # Supondo que a manutenção deve ser feita anualmente
                mensagem = f"A Inspeção do(a) {registro.tipo} localizado em {registro.localizacao} do tipo {registro.tipo} não foi mantido há mais de 1 ano."
                notificacoes.append({
                    "mensagem": mensagem,
                    "tipo": "manutencao_atrasada",
                    "localizacao": registro.localizacao,
                    "tipo": registro.tipo,
                    "data_validade_inspecao": registro.data_validade_inspecao,
                    "data_inspecao": registro.data_inspecao
                })

    return notificacoes

# Rota para notificações de inspeções pendentes
@app.route('/notificacoes')
@login_required
@permission_required('notificacoes')
def notificacoes():
    current_user = get_current_user(flask_session, session)
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        # Se for uma requisição AJAX, retorna as notificações em JSON
        notificacoes = obter_notificacoes(session)
        return jsonify(notificacoes)
    else:
        # Caso contrário, renderiza a página de notificações
        return render_template('notificacoes/notificacoes_inspecoes.html', current_user=current_user)

def obter_notificacoes(db_session):
    today = date.today()
    limite_vencimento = today + timedelta(days=30)
    notificacoes = []

    registros_pendentes = (
        db_session.query(Registro, Equipamento)
        .join(Equipamento, Registro.equipamento_id == Equipamento.id)
        .filter(
            Registro.status_inspecao == "Pendente",
            Registro.data_validade_inspecao <= limite_vencimento
        )
        .all()
    )

    for registro, equipamento in registros_pendentes:
        dias_para_vencer = (registro.data_validade_inspecao - today).days

        if registro.data_validade_inspecao < today:
            mensagem = f"A Inspeção #{registro.numero_inspecao}, realizada para o equipamento TAG: {equipamento.tag_equipamento}, no dia {registro.data_inspecao}, localizado em {equipamento.localizacao} está vencido desde {registro.data_validade_inspecao}."
            notificacoes.append({
                "mensagem": mensagem,
                "tipo": "vencido",
                "localizacao": equipamento.localizacao,
                "tipo_equipamento": equipamento.tipo,
                "data_inspecao": registro.data_inspecao.strftime('%Y-%m-%d'),                
                "data_validade": registro.data_validade_inspecao.strftime('%Y-%m-%d'),
                "tag_equipamento": equipamento.tag_equipamento,
                "equipamento_id": equipamento.id,
                "numero_inspecao": registro.numero_inspecao,  # Adicionando o número da inspeção
                "criticidade_inspecao": registro.criticidade_inspecao,                
                "status_inspecao": registro.status_inspecao
            })
        elif today < registro.data_validade_inspecao <= limite_vencimento:
            mensagem = f"A Inspeção #{registro.numero_inspecao}, realizada para o equipamento TAG: {equipamento.tag_equipamento}, no dia {registro.data_inspecao}, localizado em {equipamento.localizacao} vence em {registro.data_validade_inspecao}. Faltam aproximadamente {dias_para_vencer} dias."
            notificacoes.append({
                "mensagem": mensagem,
                "tipo": "proximo",
                "localizacao": equipamento.localizacao,
                "tipo_equipamento": equipamento.tipo,
                "data_inspecao": registro.data_inspecao.strftime('%Y-%m-%d'),                
                "data_validade": registro.data_validade_inspecao.strftime('%Y-%m-%d'),
                "tag_equipamento": equipamento.tag_equipamento,
                "equipamento_id": equipamento.id,
                "numero_inspecao": registro.numero_inspecao,  # Adicionando o número da inspeção
                "criticidade_inspecao": registro.criticidade_inspecao,                 
                "status_inspecao": registro.status_inspecao
            })

    return notificacoes

# Rota para notificações de equipamentos
@app.route('/notificacoes_equipamentos')
@login_required
@permission_required('notificacoes_equipamentos')
def notificacoes_equipamentos():
    current_user = get_current_user(flask_session, session)
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        # Se for uma requisição AJAX, retorna as notificações em JSON
        notificacoes = obter_notificacoes_equipamentos(session)
        return jsonify(notificacoes)
    else:
        # Caso contrário, renderiza a página de notificações de equipamentos
        return render_template('notificacoes/notificacoes_equipamentos.html', current_user=current_user)
 
# Função para obter notificações de equipamentos
def obter_notificacoes_equipamentos(db_session):
    equipamentos = db_session.query(Equipamento).all()
    today = date.today()
    limite_vencimento = today + timedelta(days=30)
    notificacoes = []

    for equipamento in equipamentos:
        if equipamento.data_validade is not None:
            dias_para_vencer = (equipamento.data_validade - today).days

            if equipamento.data_validade < today:
                mensagem = f"O Equipamento {equipamento.tag_equipamento} localizado em {equipamento.localizacao} do tipo {equipamento.tipo} está vencido desde {equipamento.data_validade}"
                notificacoes.append({
                    "mensagem": mensagem,
                    "tipo": "vencido",
                    "localizacao": equipamento.localizacao,
                    "tipo_equipamento": equipamento.tipo,
                    "data_validade": equipamento.data_validade.strftime('%Y-%m-%d'),
                    "tag_equipamento": equipamento.tag_equipamento,                                   
                })
            elif today < equipamento.data_validade <= limite_vencimento:
                mensagem = f"O Equipamento {equipamento.tag_equipamento} localizado em {equipamento.localizacao} do tipo {equipamento.tipo} vence em {equipamento.data_validade}. Faltam aproximadamente {dias_para_vencer} dias."
                notificacoes.append({
                    "mensagem": mensagem,
                    "tipo": "proximo",
                    "localizacao": equipamento.localizacao,
                    "tipo_equipamento": equipamento.tipo,
                    "data_validade": equipamento.data_validade.strftime('%Y-%m-%d'),
                    "tag_equipamento": equipamento.tag_equipamento,               
                })

    return notificacoes

# Métodos de Notificações Unificadas
def obter_notificacoes_unificadas(db_session):
    today = date.today()
    limite_vencimento = today + timedelta(days=30)
    notificacoes = []

    # Notificações de Inspeções
    registros_pendentes = (
        db_session.query(Registro, Equipamento)
        .join(Equipamento, Registro.equipamento_id == Equipamento.id)
        .filter(
            Registro.status_inspecao == "Pendente",
            Registro.data_validade_inspecao <= limite_vencimento
        )
        .all()
    )

    for registro, equipamento in registros_pendentes:
        dias_para_vencer = (registro.data_validade_inspecao - today).days
        tipo = "vencido" if registro.data_validade_inspecao < today else "proximo"
        
        notificacoes.append({
            "tipo": "inspecao",
            "subtipo": tipo,
            "titulo": f"Inspeção {'vencida' if tipo == 'vencido' else 'próxima do vencimento'}",
            "mensagem": f"Inspeção #{registro.numero_inspecao} para equipamento {equipamento.tag_equipamento}",
            "detalhes": {
                "localizacao": equipamento.localizacao,
                "tipo_equipamento": equipamento.tipo,
                "data_inspecao": registro.data_inspecao.strftime('%d/%m/%Y'),
                "data_validade": registro.data_validade_inspecao.strftime('%d/%m/%Y'),
                "dias_restantes": dias_para_vencer if tipo == "proximo" else 0,
                "status": registro.status_inspecao
            },
            "prioridade": 1 if tipo == "vencido" else 2,
            "timestamp": datetime.now().isoformat(),
            "url": f"/inspecao/detalhes/{registro.id}"
        })

    # Notificações de Equipamento
    equipamentos = db_session.query(Equipamento).filter(
        Equipamento.data_validade <= limite_vencimento
    ).all()

    for equipamento in equipamentos:
        if equipamento.data_validade:
            dias_para_vencer = (equipamento.data_validade - today).days
            tipo = "vencido" if equipamento.data_validade < today else "proximo"
            
            notificacoes.append({
                "tipo": "equipamento",
                "subtipo": tipo,
                "titulo": f"Equipamento {'vencido' if tipo == 'vencido' else 'próximo do vencimento'}",
                "mensagem": f"Equipamento {equipamento.tag_equipamento} ({equipamento.tipo})",
                "detalhes": {
                    "localizacao": equipamento.localizacao,
                    "tipo_equipamento": equipamento.tipo,
                    "data_validade": equipamento.data_validade.strftime('%d/%m/%Y'),
                    "dias_restantes": dias_para_vencer if tipo == "proximo" else 0,
                    "status": "Vencido" if tipo == "vencido" else "Em dia"
                },
                "prioridade": 1 if tipo == "vencido" else 3,
                "timestamp": datetime.now().isoformat(),
                "url": f"/equipamentos/editar/{equipamento.id}"
            })

    # Ordenar por prioridade (vencidos primeiro) e depois por timestamp
    notificacoes.sort(key=lambda x: (x['prioridade'], x['timestamp']))
    
    return notificacoes

# Rota para notificações unificadas
@app.route('/api/notificacoes')
@login_required
def notificacoes_api():
    notificacoes = obter_notificacoes_unificadas(session)
    return jsonify({
        "total": len(notificacoes),
        "notificacoes": notificacoes
    })

# Rota para marcar notificação como lida
@app.route('/api/notificacoes/marcar_lida', methods=['POST'])
@login_required
def marcar_notificacao_lida():
    data = request.get_json()
    # Implementar lógica para marcar como lida no banco de dados
    return jsonify({"success": True})

def run_server():
    """Iniciar o servidor Flask com base na configuração carregada."""
    config = load_config()
    ip = config.get("SERVER_IP", "127.0.0.1")
    port = config.get("SERVER_PORT", 5000)
    app.run(host=ip, port=port, debug=True)
 
if __name__ == '__main__':
    run_server()
    # Inicializar configurações quando o app iniciar
    init_config()    